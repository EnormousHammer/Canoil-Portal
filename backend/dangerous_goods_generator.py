"""
Dangerous Goods Declaration Generator
Automatically detects dangerous goods items and generates filled Word documents
"""

from docx import Document
from docx.shared import Pt
from typing import Dict, Any, Optional
import os
from datetime import datetime
import re

# Dangerous Goods Product Mapping
# Template paths are relative to this file's directory
def _get_template_path(filename):
    """Get absolute path to template file"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(current_dir, 'templates', 'dangerous_goods', filename)

DANGEROUS_GOODS_PRODUCTS = {
    'REOLUBE TURBOFLUID 32B GT': {
        'template': _get_template_path('REOLUBE TURBOFLUID 32B GT.docx'),
        'match_patterns': [
            '32B GT',           # In description
            '32BGT',            # In item code (without space)
            'REOL32BGT',        # Full item code prefix
            'TURBOFLUID 32B',   # Alternative description
            'CC 32BGT',         # Alternative code format
            'CC32BGT'           # Alternative code format (no space)
        ]
    },
    'REOLUBE TURBOFLUID 46B': {
        'template': _get_template_path('REOLUBE TURBOFLUID 46B.docx'),
        'match_patterns': [
            '46B',              # In description or code
            'REOL46B',          # Full item code prefix
            'TURBOFLUID 46B',   # Alternative description
            'CC 46B',           # Alternative code format
            'CC46B',            # Alternative code format (no space)
            'REOL46BGWCONSIGN'  # GW Consignment version (same product)
        ]
    },
    'REOLUBE TURBOFLUID 46XC': {
        'template': _get_template_path('REOLUBE TURBOFLUID 46XC.docx'),
        'match_patterns': [
            '46XC',             # In description or code
            'REOL46XC',         # Full item code prefix
            'TURBOFLUID 46XC',  # Alternative description
            'CC 46XC',          # Alternative code format
            'CC46XC'            # Alternative code format (no space)
        ]
    }
}


def detect_dangerous_goods(items: list) -> list:
    """
    Detect dangerous goods items from SO items
    Returns list of (combined_items, product_name, template_path)
    Groups items by product type and sums quantities
    """
    print(f"\n🔍 DANGEROUS GOODS DETECTION - Checking {len(items)} items...")
    
    # Group items by product type
    product_groups = {}
    
    for item in items:
        item_code = str(item.get('item_code', '')).upper().strip()
        description = str(item.get('description', '')).upper().strip()
        
        print(f"   Checking: {item_code} | {description}")
        
        # Skip non-product items
        if any(keyword in description for keyword in ['PALLET', 'FREIGHT', 'BROKERAGE', 'CHARGE']):
            print(f"      ⏭️  Skipping (non-product item)")
            continue
        
        # Check against each dangerous goods product
        for product_name, config in DANGEROUS_GOODS_PRODUCTS.items():
            matched = False
            matched_pattern = None
            
            for pattern in config['match_patterns']:
                pattern_upper = pattern.upper().strip()
                
                # Check if pattern is IN the item code or description
                if pattern_upper in item_code or pattern_upper in description:
                    matched = True
                    matched_pattern = pattern
                    break
            
            if matched:
                # Group items by product type
                if product_name not in product_groups:
                    product_groups[product_name] = {
                        'items': [],
                        'template': config['template'],
                        'total_quantity': 0
                    }
                
                product_groups[product_name]['items'].append(item)
                product_groups[product_name]['total_quantity'] += item.get('quantity', 0)
                
                print(f"      🔴 MATCH: {product_name} (pattern: '{matched_pattern}')")
                print(f"         Quantity: {item.get('quantity')} {item.get('unit', '')}")
                break  # Stop checking other product types once matched
    
    # Convert to list format with combined quantities
    dangerous_items = []
    for product_name, group_data in product_groups.items():
        # Create a combined item with total quantity
        first_item = group_data['items'][0]
        combined_item = {
            **first_item,
            'quantity': group_data['total_quantity'],  # TOTAL quantity from all matching items
            'original_items': group_data['items']  # Keep track of original items
        }
        
        dangerous_items.append({
            'item': combined_item,
            'product_name': product_name,
            'template': group_data['template']
        })
        
        print(f"   ✅ Combined: {product_name} - TOTAL {group_data['total_quantity']} units")
    
    return dangerous_items


def format_address_for_dg(address_dict: Dict[str, Any], company_name: str = '') -> str:
    """
    Format address for dangerous goods declaration
    Returns multi-line formatted address string
    """
    lines = []
    
    # Company name
    company = address_dict.get('company') or address_dict.get('company_name') or company_name
    if company:
        lines.append(company)
    
    # Contact person
    contact = address_dict.get('attention') or address_dict.get('attn') or address_dict.get('contact_person')
    if contact:
        lines.append(f"Attn: {contact}")
    
    # Street address
    street = address_dict.get('address') or address_dict.get('street1') or address_dict.get('street_address')
    if street:
        lines.append(street)
    
    # City, Province, Postal Code
    city_parts = []
    if address_dict.get('city'):
        city_parts.append(address_dict['city'])
    if address_dict.get('province') or address_dict.get('state'):
        city_parts.append(address_dict.get('province') or address_dict.get('state'))
    if address_dict.get('postal_code'):
        city_parts.append(address_dict.get('postal_code'))
    
    if city_parts:
        lines.append(', '.join(city_parts))
    
    # Country
    if address_dict.get('country'):
        lines.append(address_dict['country'])
    
    return '\n'.join(lines)


def fill_dangerous_goods_declaration(
    template_path: str,
    so_data: Dict[str, Any],
    item: Dict[str, Any],
    email_shipping: Dict[str, Any] = None
) -> Document:
    """
    Fill dangerous goods declaration Word document with SO data
    """
    print(f"\n📋 Filling Dangerous Goods Declaration...")
    print(f"   Template: {os.path.basename(template_path)}")
    
    # Load template
    doc = Document(template_path)
    
    # Get addresses (support both field name formats)
    shipping_addr = so_data.get('shipping_address', {}) or so_data.get('ship_to', {})
    billing_addr = so_data.get('billing_address', {}) or so_data.get('sold_to', {})
    customer_name = so_data.get('customer_name', '')
    
    print(f"   DEBUG DG FILL: SO data keys: {list(so_data.keys())}")
    print(f"   DEBUG DG FILL: Shipping address: {shipping_addr}")
    print(f"   DEBUG DG FILL: Billing address: {billing_addr}")
    print(f"   DEBUG DG FILL: Customer name: {customer_name}")
    print(f"   DEBUG DG FILL: Item: {item}")
    
    # Format addresses
    consignee_text = format_address_for_dg(shipping_addr, customer_name)
    buyer_text = format_address_for_dg(billing_addr, customer_name)
    
    print(f"   DEBUG DG FILL: Consignee text: {consignee_text}")
    print(f"   DEBUG DG FILL: Buyer text: {buyer_text}")
    
    # Format destination
    dest_parts = []
    if shipping_addr.get('city'):
        dest_parts.append(shipping_addr['city'])
    if shipping_addr.get('province') or shipping_addr.get('state'):
        dest_parts.append(shipping_addr.get('province') or shipping_addr.get('state'))
    if shipping_addr.get('country'):
        dest_parts.append(shipping_addr['country'])
    destination = ', '.join(dest_parts) if dest_parts else ''
    
    # Calculate quantity for the dangerous goods item
    quantity = item.get('quantity', 0)
    unit = item.get('unit', 'DRUM').upper()
    description = item.get('description', '')
    
    # Extract weight from description (e.g., "247KG", "230KG", "17KG")
    weight_match = re.search(r'(\d+)\s*KG', description.upper())
    weight_str = f"{weight_match.group(1)}KG" if weight_match else ""
    
    # Format quantity text (e.g., "2 – 247KG Drums")
    if 'DRUM' in unit:
        if weight_str:
            quantity_text = f"{quantity} – {weight_str} Drum" if quantity == 1 else f"{quantity} – {weight_str} Drums"
        else:
            quantity_text = f"{quantity} – Drum" if quantity == 1 else f"{quantity} – Drums"
    elif 'PAIL' in unit:
        if weight_str:
            quantity_text = f"{quantity} – {weight_str} Pail" if quantity == 1 else f"{quantity} – {weight_str} Pails"
        else:
            quantity_text = f"{quantity} – Pail" if quantity == 1 else f"{quantity} – Pails"
    else:
        quantity_text = f"{quantity} – {unit}" if quantity == 1 else f"{quantity} – {unit}s"
    
    # Get PO number for reference
    po_number = so_data.get('po_number', '') or so_data.get('order_details', {}).get('po_number', '')
    
    print(f"   Consignee: {consignee_text[:50]}...")
    print(f"   Buyer: {buyer_text[:50]}...")
    print(f"   Destination: {destination}")
    print(f"   Quantity: {quantity_text}")
    print(f"   PO#: {po_number}")
    
    # Fill the Word document table - PRESERVE FORMATTING
    # The template has ONE large table with merged cells
    if doc.tables:
        table = doc.tables[0]
        
        try:
            # Row 0: PO# - Fill Shipper's Reference Number (search all cells, fill once)
            po_filled = False
            if po_number:
                for row_idx, row in enumerate(table.rows):
                    if po_filled:
                        break
                    for cell in row.cells:
                        if po_filled:
                            break
                        for paragraph in cell.paragraphs:
                            # Look for the P.O.# field more broadly
                            if ('P.O.#' in paragraph.text or 'PO#' in paragraph.text or 'Reference Number' in paragraph.text) and not po_filled:
                                # Clear and rebuild with PO number
                                paragraph.clear()
                                run1 = paragraph.add_run("Shipper's Reference Number: P.O.#")
                                run1.bold = True
                                run1.font.name = 'Times New Roman'
                                run1.font.size = Pt(11)
                                run2 = paragraph.add_run(po_number)
                                run2.bold = False
                                run2.font.name = 'Times New Roman'
                                run2.font.size = Pt(9)  # PO# must be size 9
                                # Ensure no line breaks in PO cell
                                paragraph.paragraph_format.keep_together = True
                                print(f"   ✅ Filled PO#: {po_number} in row {row_idx}")
                                po_filled = True
                                break
                
                if not po_filled:
                    print(f"   ⚠️  Warning: Could not find PO# field to fill")
                    print(f"   DEBUG: Searching for PO# field in document...")
                    # Debug: print all text in table to see what's there
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    print(f"     Row {row_idx}, Cell {cell_idx}: '{para.text}'")
            
            # Row 1: Consignee - Find paragraph with "Consignee:" and add address after it
            consignee_cell = table.rows[1].cells[0]
            for paragraph in consignee_cell.paragraphs:
                if 'Consignee:' in paragraph.text:
                    # Clear paragraph completely
                    paragraph.clear()
                    # Add label
                    run1 = paragraph.add_run('Consignee:\n')
                    run1.bold = False
                    run1.font.name = 'Times New Roman'
                    run1.font.size = Pt(11)
                    # Add address text (not bold, Times New Roman)
                    # If address is more than 4 lines, use size 9 to prevent layout issues
                    consignee_lines = consignee_text.split('\n')
                    address_font_size = Pt(9) if len(consignee_lines) > 4 else Pt(11)
                    run2 = paragraph.add_run(consignee_text)
                    run2.bold = False
                    run2.font.name = 'Times New Roman'
                    run2.font.size = address_font_size
                    break
            
            # Row 1: Buyer - Only fill if DIFFERENT from consignee
            # ORDER-INDEPENDENT comparison: Extract key elements as SETS
            def extract_address_signature(addr):
                """
                Extract address signature: order-independent set of key elements
                Returns set of (numbers, words, postal_code)
                """
                if not addr:
                    return set()
                
                import re
                addr_upper = addr.upper()
                
                # Extract postal code first (so we can remove it from words)
                postal = ''
                postal_pattern = r'[A-Z]\d[A-Z]\s?\d[A-Z]\d|\d{5}(?:-?\d{4})?'
                postal_match = re.search(postal_pattern, addr_upper)
                if postal_match:
                    postal = postal_match.group().replace(' ', '').replace('-', '')
                    # Remove postal from address to avoid duplicate matching
                    addr_upper = addr_upper.replace(postal_match.group(), '')
                
                # Remove all punctuation
                for char in [':', ',', '.', '-', '#', '(', ')', '[', ']', '/', '\\', '\n']:
                    addr_upper = addr_upper.replace(char, ' ')
                
                # Split into words
                words = addr_upper.split()
                
                # Separate numbers and text
                numbers = set()
                text_words = set()
                
                for word in words:
                    word = word.strip()
                    if not word:
                        continue
                    if word.isdigit():
                        numbers.add(word)
                    else:
                        # Skip common noise words
                        noise = ['CANADA', 'USA', 'UNITED', 'STATES', 'ATTENTION', 'ATTN', 'CO', 'THE']
                        if word not in noise:
                            text_words.add(word)
                
                # Create signature: set of all key elements
                signature = set()
                signature.add(('POSTAL', postal))
                for num in numbers:
                    signature.add(('NUM', num))
                for word in text_words:
                    signature.add(('WORD', word))
                
                return signature
            
            consignee_sig = extract_address_signature(consignee_text)
            buyer_sig = extract_address_signature(buyer_text)
            
            # Calculate similarity: what % of elements match?
            if consignee_sig and buyer_sig:
                common = consignee_sig & buyer_sig  # Intersection
                total = consignee_sig | buyer_sig    # Union
                similarity = len(common) / len(total) if total else 0
            else:
                similarity = 0
            
            # Consider same if 80%+ elements match (allows minor differences like "Dept" position)
            addresses_are_same = similarity >= 0.80
            
            print(f"   DEBUG: Consignee has {len(consignee_sig)} elements, Buyer has {len(buyer_sig)} elements")
            print(f"   DEBUG: Common elements: {len(consignee_sig & buyer_sig)}, Similarity: {similarity:.1%}")
            print(f"   DEBUG: Are they the same address? {addresses_are_same} (threshold: 80%)")
            
            buyer_cell = table.rows[1].cells[7]
            if buyer_text and not addresses_are_same:
                # Addresses are DIFFERENT - fill Buyer field
                for paragraph in buyer_cell.paragraphs:
                    if 'Buyer:' in paragraph.text:
                        paragraph.clear()
                        run1 = paragraph.add_run('Buyer:\n')
                        run1.bold = False
                        run1.font.name = 'Times New Roman'
                        run1.font.size = Pt(11)
                        # If address is more than 4 lines, use size 9 to prevent layout issues
                        buyer_lines = buyer_text.split('\n')
                        address_font_size = Pt(9) if len(buyer_lines) > 4 else Pt(11)
                        run2 = paragraph.add_run(buyer_text)
                        run2.bold = False
                        run2.font.name = 'Times New Roman'
                        run2.font.size = address_font_size
                        break
                print(f"   ✅ Buyer is DIFFERENT - filled separately")
            else:
                # Addresses are SAME - leave Buyer empty
                print(f"   ℹ️  Buyer same as Consignee - left empty")
            
            # Row 5-6: Destination - Replace only the text after "Destination:"
            for row_idx in [5, 6]:
                for cell in table.rows[row_idx].cells:
                    for paragraph in cell.paragraphs:
                        if 'Destination:' in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run(f'Destination: {destination}')
                            run.bold = False
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            break
            
            # Row 11: Quantity - Find the cell with quantity placeholder and replace
            for cell in table.rows[11].cells:
                for paragraph in cell.paragraphs:
                    # Match various patterns: "– 230KG Drums", "– 247KG Drum", "KG Drum", etc.
                    if ('KG Drum' in paragraph.text or 'KG DRUM' in paragraph.text or 
                        '– Drum' in paragraph.text or 'Pail' in paragraph.text):
                        # Clear and add quantity with weight (BOLD, Times New Roman)
                        paragraph.clear()
                        run = paragraph.add_run(quantity_text)
                        run.bold = True  # Quantity should be BOLD
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        print(f"   ✅ Filled quantity: {quantity_text}")
                        break
            
            print("   ✅ Word document filled successfully (formatting preserved)")
            
        except Exception as e:
            print(f"   ⚠️ Warning: Error filling some fields: {e}")
            print(f"   Document will be saved with partial fill")
    
    # NOTE: Signature space handled externally in template - don't modify structure
    return doc


def generate_dangerous_goods_declarations(
    so_data: Dict[str, Any],
    items: list,
    email_shipping: Dict[str, Any] = None
) -> Dict[str, Any]:
    """
    Main function: Detect dangerous goods and generate all required declarations
    Also finds latest SDS and COFA documents for each DG item
    
    Returns dictionary with:
    {
        'dg_forms': [(filepath, filename), ...],
        'sds_files': [(filepath, filename, product_name), ...],
        'cofa_files': [(filepath, filename, product_name, batch), ...]
    }
    """
    print("="*80)
    print("DANGEROUS GOODS DECLARATION GENERATOR")
    print("="*80)
    
    # Detect dangerous goods
    dangerous_items = detect_dangerous_goods(items)
    
    if not dangerous_items:
        print("✅ No dangerous goods detected in this shipment")
        return {'dg_forms': [], 'sds_files': [], 'cofa_files': []}
    
    print(f"\n🔴 Found {len(dangerous_items)} dangerous goods item(s)")
    
    generated_files = []
    sds_files = []
    cofa_files = []
    so_number = so_data.get('so_number', 'UNKNOWN')
    
    # Import document finder
    try:
        from document_finder import find_documents_for_dg_item
    except ImportError:
        print("⚠️  Warning: document_finder module not available, SDS/COFA search disabled")
        find_documents_for_dg_item = None
    
    # Generate a declaration for each dangerous goods item - with individual error handling
    for idx, dg_item in enumerate(dangerous_items, 1):
        try:
            item = dg_item['item']
            product_name = dg_item['product_name']
            template_path = dg_item['template']
            
            # Check if template file exists
            if not os.path.exists(template_path):
                print(f"❌ Template file not found: {template_path}")
                print(f"   Skipping DG form generation for {product_name}")
                continue
            
            # Get product code and batch number from item
            item_code = item.get('item_code', '')
            description = item.get('description', '')
            batch_number = item.get('batch_number') or item.get('batch')
            
            print(f"\n📦 Processing DG Item {idx}/{len(dangerous_items)}: {product_name}")
            print(f"   Item Code: {item_code}")
            print(f"   Batch: {batch_number or 'Not provided'}")
            print(f"   Template: {template_path}")
            
            # Fill the document
            doc = fill_dangerous_goods_declaration(template_path, so_data, item, email_shipping)
            
            # Save the filled document
            output_dir = "generated_documents"
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"DG_Declaration_SO{so_number}_{product_name.replace(' ', '_')}_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            doc.save(filepath)
            
            print(f"   ✅ DG Form Generated: {filename}")
            generated_files.append((filepath, filename))
            
        except Exception as dg_gen_err:
            print(f"   ❌ FAILED to generate DG form for {product_name}: {dg_gen_err}")
            import traceback
            traceback.print_exc()
            # Continue to next DG item even if this one fails
            continue
        
        # Find SDS and COFA documents - with error handling
        if find_documents_for_dg_item:
            try:
                # Try to import Google Drive service to pass to document finder
                google_drive_svc = None
                try:
                    from app import google_drive_service
                    google_drive_svc = google_drive_service
                except:
                    pass
                
                docs = find_documents_for_dg_item(item_code or product_name, batch_number, google_drive_svc)
                
                if docs.get('sds'):
                    sds_path = docs['sds']
                    sds_filename = os.path.basename(sds_path)
                    sds_files.append((sds_path, sds_filename, product_name))
                    print(f"   ✅ SDS Found: {sds_filename}")
                else:
                    print(f"   ⚠️  SDS Not Found for {product_name}")
                
                if docs.get('cofa'):
                    cofa_path = docs['cofa']
                    cofa_filename = os.path.basename(cofa_path)
                    cofa_files.append((cofa_path, cofa_filename, product_name, batch_number))
                    print(f"   ✅ COFA Found: {cofa_filename}")
                else:
                    if batch_number:
                        print(f"   ⚠️  COFA Not Found for {product_name} batch {batch_number}")
                    else:
                        print(f"   ⏭️  COFA Search Skipped (no batch number)")
            except Exception as doc_find_err:
                print(f"   ⚠️  Document search failed for {product_name}: {doc_find_err}")
                # Continue anyway - DG form already generated
    
    print("\n" + "="*80)
    print(f"✅ COMPLETED:")
    print(f"   - {len(generated_files)} DG Declaration(s)")
    print(f"   - {len(sds_files)} SDS file(s) found")
    print(f"   - {len(cofa_files)} COFA file(s) found")
    print("="*80)
    
    return {
        'dg_forms': generated_files,
        'sds_files': sds_files,
        'cofa_files': cofa_files
    }


def save_dangerous_goods_declaration(doc: Document, so_number: str, product_name: str) -> tuple:
    """
    Save dangerous goods declaration and return filepath and filename
    """
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"DG_Declaration_SO{so_number}_{product_name.replace(' ', '_')}_{timestamp}.docx"
    
    output_dir = "generated_documents"
    os.makedirs(output_dir, exist_ok=True)
    
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath, filename

