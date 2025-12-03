"""
Delivery Note Generator for Axel France Orders
Automatically generates filled Word documents for Axel France shipments
"""

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional
import os
from datetime import datetime
import re


def _get_template_path():
    """Get absolute path to delivery note template file"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(current_dir, 'templates', 'delivery_note', 'delivery_note_template.docx')


def is_axel_france_order(so_data: Dict[str, Any]) -> bool:
    """
    Check if the order is for Axel France
    Returns True if customer name contains 'AXEL' (case-insensitive)
    """
    customer_name = so_data.get('customer_name', '')
    shipping_company = ''
    
    # Check shipping address company name
    shipping_addr = so_data.get('shipping_address', {}) or so_data.get('ship_to', {})
    if shipping_addr:
        shipping_company = shipping_addr.get('company', '') or shipping_addr.get('company_name', '')
    
    # Also check billing address
    billing_addr = so_data.get('billing_address', {}) or so_data.get('sold_to', {})
    if billing_addr:
        billing_company = billing_addr.get('company', '') or billing_addr.get('company_name', '')
    else:
        billing_company = ''
    
    # Check all possible sources for "AXEL"
    all_names = f"{customer_name} {shipping_company} {billing_company}".upper()
    
    is_axel = 'AXEL' in all_names
    
    print(f"🔍 Axel France Detection:")
    print(f"   Customer Name: {customer_name}")
    print(f"   Shipping Company: {shipping_company}")
    print(f"   Billing Company: {billing_company}")
    print(f"   Is Axel France: {is_axel}")
    
    return is_axel


def format_shipping_address(so_data: Dict[str, Any]) -> str:
    """
    Format the full shipping address for DELIVER TO section
    """
    shipping_addr = so_data.get('shipping_address', {}) or so_data.get('ship_to', {})
    customer_name = so_data.get('customer_name', '')
    
    lines = []
    
    # Company name
    company = shipping_addr.get('company') or shipping_addr.get('company_name') or customer_name
    if company:
        lines.append(company)
    
    # Street address (could be multi-line)
    street = shipping_addr.get('address') or shipping_addr.get('street') or shipping_addr.get('street1')
    if street:
        # Handle multi-line addresses
        street_lines = street.replace('\r\n', '\n').replace('\r', '\n').split('\n')
        lines.extend([s.strip() for s in street_lines if s.strip()])
    
    # City, Province/State, Postal Code
    city = shipping_addr.get('city', '')
    province = shipping_addr.get('province') or shipping_addr.get('state', '')
    postal = shipping_addr.get('postal_code') or shipping_addr.get('postal', '')
    
    city_line_parts = []
    if city:
        city_line_parts.append(city)
    if province:
        city_line_parts.append(province)
    if postal:
        city_line_parts.append(postal)
    
    if city_line_parts:
        lines.append(' '.join(city_line_parts))
    
    # Country
    country = shipping_addr.get('country', '')
    if country:
        lines.append(country)
    
    return '\n'.join(lines)


def get_company_name(so_data: Dict[str, Any]) -> str:
    """
    Get just the company name for DELIVER TO / SHIP TO rows
    """
    shipping_addr = so_data.get('shipping_address', {}) or so_data.get('ship_to', {})
    customer_name = so_data.get('customer_name', '')
    
    company = shipping_addr.get('company') or shipping_addr.get('company_name') or customer_name
    return company or 'AXEL France'


def format_item_quantity(item: Dict[str, Any]) -> str:
    """
    Format quantity with unit (e.g., "6 drums", "2 pails")
    """
    quantity = item.get('quantity', 0)
    unit = item.get('unit', 'units').lower()
    
    # Normalize common unit names
    if 'drum' in unit:
        unit_display = 'drum' if quantity == 1 else 'drums'
    elif 'pail' in unit:
        unit_display = 'pail' if quantity == 1 else 'pails'
    elif 'case' in unit:
        unit_display = 'case' if quantity == 1 else 'cases'
    elif 'box' in unit:
        unit_display = 'box' if quantity == 1 else 'boxes'
    elif 'kg' in unit or 'kilo' in unit:
        unit_display = 'kg'
    elif 'lb' in unit or 'pound' in unit:
        unit_display = 'lbs'
    else:
        unit_display = unit if quantity == 1 else f"{unit}s"
    
    return f"{quantity} {unit_display}"


def get_item_code_short(item: Dict[str, Any]) -> str:
    """
    Get shortened item code for display (e.g., "MOV LL 0", "MOV LL 2")
    """
    item_code = item.get('item_code', '') or item.get('item_no', '')
    description = item.get('description', '')
    
    # If item code exists, use it
    if item_code:
        # Shorten very long codes
        if len(item_code) > 15:
            return item_code[:15]
        return item_code
    
    # Otherwise try to extract from description
    # Look for patterns like "MOV LL 0", "REOL32BGT", etc.
    match = re.search(r'^([A-Z0-9\s]+)', description.upper())
    if match:
        code = match.group(1).strip()[:15]
        return code
    
    return description[:15] if description else 'ITEM'


def get_item_description_short(item: Dict[str, Any]) -> str:
    """
    Get shortened description (e.g., "Lubricating grease")
    """
    description = item.get('description', '')
    
    # Common mappings for known products
    desc_upper = description.upper()
    
    if 'MOV' in desc_upper and ('GREASE' in desc_upper or 'LL' in desc_upper):
        return 'Lubricating grease'
    elif 'REOLUBE' in desc_upper or 'TURBOFLUID' in desc_upper:
        return 'Turbine oil'
    elif 'GREASE' in desc_upper:
        return 'Lubricating grease'
    elif 'OIL' in desc_upper:
        return 'Lubricating oil'
    
    # Default: use first 20 chars of description
    if len(description) > 20:
        return description[:20] + '...'
    return description or 'Product'


def filter_product_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Filter out non-product items (pallets, freight, charges, etc.)
    Only keep actual products
    """
    product_items = []
    
    for item in items:
        description = str(item.get('description', '')).upper()
        item_code = str(item.get('item_code', '')).upper()
        
        # Skip non-product items
        skip_keywords = ['PALLET', 'FREIGHT', 'BROKERAGE', 'CHARGE', 'SHIPPING', 
                        'HANDLING', 'DELIVERY', 'SURCHARGE', 'TAX', 'FEE']
        
        is_product = True
        for keyword in skip_keywords:
            if keyword in description or keyword in item_code:
                is_product = False
                break
        
        if is_product and item.get('quantity', 0) > 0:
            product_items.append(item)
    
    return product_items


def fill_delivery_note(
    so_data: Dict[str, Any],
    items: List[Dict[str, Any]],
    booking_number: str = ''
) -> Document:
    """
    Fill delivery note Word document with SO data
    
    Args:
        so_data: Sales order data
        items: List of items from the SO
        booking_number: Optional booking number (from user input)
    
    Returns:
        Filled Document object
    """
    template_path = _get_template_path()
    
    print(f"\n📋 Filling Delivery Note...")
    print(f"   Template: {template_path}")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Delivery note template not found: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Apply professional font (Calibri) to entire document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Also set the font for the style's element (for compatibility)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    
    # Add "Canada" to the Canoil DELIVERY FROM address
    # Find "Georgetown, ON L7G 4R7" and add "Canada" after it
    from docx.oxml import OxmlElement as OxmlElem
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if 'Georgetown' in text and 'L7G 4R7' in text:
            # Found the Georgetown line - insert Canada after it
            canada_p = OxmlElem('w:p')
            canada_r = OxmlElem('w:r')
            canada_t = OxmlElem('w:t')
            canada_t.text = 'Canada'
            canada_r.append(canada_t)
            canada_p.append(canada_r)
            paragraph._element.addnext(canada_p)
            print("   ✅ Added 'Canada' to Canoil address")
            break
    
    # Get data for filling
    customer_name = get_company_name(so_data)
    shipping_address = format_shipping_address(so_data)
    po_number = so_data.get('po_number', '') or so_data.get('order_details', {}).get('po_number', '')
    
    # Filter to only product items
    product_items = filter_product_items(items)
    
    print(f"   Customer: {customer_name}")
    print(f"   Shipping Address: {shipping_address}")
    print(f"   PO#: {po_number}")
    print(f"   Booking#: {booking_number or '(empty)'}")
    print(f"   Items: {len(product_items)} products")
    
    # Split shipping address into lines
    address_lines = shipping_address.split('\n')
    
    # Track which DELIVER TO: we've seen (first = address header, second = table row header)
    deliver_to_count = 0
    address_filled = False
    
    # First pass: Find the DELIVER TO: header paragraph
    deliver_to_header_para = None
    deliver_to_header_idx = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == 'DELIVER TO:':
            deliver_to_header_para = paragraph
            deliver_to_header_idx = i
            break
    
    # Fill address by inserting paragraphs after DELIVER TO: header
    if deliver_to_header_para is not None:
        # Get the paragraph's parent element
        parent = deliver_to_header_para._element.getparent()
        
        # Find the index of the DELIVER TO: paragraph element
        para_index = list(parent).index(deliver_to_header_para._element)
        
        # Create and insert address paragraphs right after DELIVER TO:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Insert address lines in reverse order (so they end up in correct order)
        for addr_line in reversed(address_lines):
            # Create a new paragraph element
            new_p = OxmlElement('w:p')
            new_r = OxmlElement('w:r')
            new_t = OxmlElement('w:t')
            new_t.text = addr_line
            new_r.append(new_t)
            new_p.append(new_r)
            
            # Insert after the DELIVER TO: paragraph
            deliver_to_header_para._element.addnext(new_p)
        
        address_filled = True
        deliver_to_count = 1
        
        # Find the last address element we inserted
        last_addr_element = deliver_to_header_para._element.getnext()
        for _ in range(len(address_lines) - 1):
            if last_addr_element is not None:
                next_elem = last_addr_element.getnext()
                if next_elem is not None:
                    last_addr_element = next_elem
        
        # Remove ALL empty placeholder paragraphs from template between address and BOOKING# row
        # But keep track so we can add exactly 1 back
        if last_addr_element is not None:
            next_elem = last_addr_element.getnext()
            elements_to_remove = []
            booking_elem = None
            
            while next_elem is not None:
                text_content = ''.join(next_elem.itertext())
                if 'BOOKING#' in text_content:
                    booking_elem = next_elem
                    break
                # Mark all empty paragraphs for removal
                if not text_content.strip():
                    elements_to_remove.append(next_elem)
                next_elem = next_elem.getnext()
            
            # Remove all the empty paragraphs
            for elem in elements_to_remove:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
            
            # Now add exactly 1 blank paragraph for spacing
            blank_p = OxmlElement('w:p')
            last_addr_element.addnext(blank_p)
    
    # Second pass: Find and fill BOOKING# row - PRESERVE ORIGINAL SPACING
    # Also ensure there's spacing before this section
    booking_para_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        
        # Find the row with BOOKING# header (but not values yet)
        if 'BOOKING#' in text and 'SHIP TO:' in text:
            booking_para_idx = idx
            
            # Ensure the paragraph before this one is empty (for spacing)
            if idx > 0:
                prev_para = doc.paragraphs[idx - 1]
                if prev_para.text.strip():
                    # Previous para has content, we need to ensure spacing
                    # The template should already have empty paras, but let's verify
                    pass
            
            paragraph.clear()
            
            # Match the exact layout from the example:
            # "DELIVER TO:                              SHIP TO:                          BOOKING# 26197060."
            run1 = paragraph.add_run('DELIVER TO:')
            run1.bold = True
            
            run2 = paragraph.add_run('                              ')  # Spaces to SHIP TO
            
            run3 = paragraph.add_run('SHIP TO:')
            run3.bold = True
            
            run4 = paragraph.add_run('                          ')  # Spaces to BOOKING#
            
            run5 = paragraph.add_run('BOOKING# ')
            run5.bold = True
            
            if booking_number:
                run6 = paragraph.add_run(f'{booking_number}.')
                run6.bold = True
            break
    
    # Ensure there are 2 empty paragraphs before the BOOKING# row for spacing
    if booking_para_idx and booking_para_idx >= 2:
        # Check paragraphs before BOOKING# row
        # Find where the address ends and ensure blank lines after
        for idx in range(booking_para_idx - 1, max(0, booking_para_idx - 3), -1):
            para = doc.paragraphs[idx]
            # If this para has content (part of address), the ones after should be blank
            if para.text.strip() and 'DELIVER TO' not in para.text:
                # This is likely the last line of address (France)
                # Make sure the next para(s) before BOOKING# are blank for spacing
                for blank_idx in range(idx + 1, booking_para_idx):
                    if blank_idx < len(doc.paragraphs):
                        doc.paragraphs[blank_idx].clear()  # Ensure it's blank for spacing
                break
    
    # Third pass: Find and fill the company names + PO# row - MATCH EXAMPLE LAYOUT
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Find the line with PO# (the values line, not the BOOKING# header line)
        if 'PO#' in text and 'BOOKING#' not in text:
            paragraph.clear()
            
            # Match the exact layout from the example:
            # "AXEL France                                  AXEL France                     PO# 27863."
            run1 = paragraph.add_run(f'{customer_name}')
            
            # Calculate spaces to maintain column alignment (position ~45)
            spaces1 = ' ' * max(1, 42 - len(customer_name))
            run2 = paragraph.add_run(spaces1)
            
            run3 = paragraph.add_run(f'{customer_name}')
            
            # Spaces before PO# (position ~85)
            spaces2 = ' ' * max(1, 37 - len(customer_name))
            run4 = paragraph.add_run(spaces2)
            
            run5 = paragraph.add_run('PO# ')
            run5.bold = True
            
            if po_number:
                run6 = paragraph.add_run(f'{po_number}.')
                run6.bold = True
            break
    
    # Fill the table with items
    if doc.tables:
        table = doc.tables[0]
        
        # Table structure:
        # Row 0: Headers (ITEM | DESCRIPTION | ORDERED | DELIVERED | OUTSTANDING)
        # Row 1+: Item rows
        
        # First, clear existing item rows (keep header)
        # The template has empty rows we'll fill
        
        # Find the header row index
        header_row_idx = 0
        for row_idx, row in enumerate(table.rows):
            cell_text = ''.join([cell.text for cell in row.cells]).upper()
            if 'ITEM' in cell_text and 'DESCRIPTION' in cell_text and 'ORDERED' in cell_text:
                header_row_idx = row_idx
                break
        
        print(f"   Header row index: {header_row_idx}")
        
        # Fill item rows starting after header
        for item_idx, item in enumerate(product_items):
            row_idx = header_row_idx + 1 + item_idx
            
            item_code = get_item_code_short(item)
            description = get_item_description_short(item)
            quantity = format_item_quantity(item)
            
            print(f"   Item {item_idx + 1}: {item_code} | {description} | {quantity}")
            
            if row_idx < len(table.rows):
                # Use existing row
                row = table.rows[row_idx]
                
                # Fill cells
                if len(row.cells) >= 5:
                    # ITEM column
                    row.cells[0].text = item_code
                    for para in row.cells[0].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                    
                    # DESCRIPTION column
                    row.cells[1].text = description
                    for para in row.cells[1].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                    
                    # ORDERED column
                    row.cells[2].text = quantity
                    for para in row.cells[2].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                    
                    # DELIVERED and OUTSTANDING columns - leave empty
                    row.cells[3].text = ''
                    row.cells[4].text = ''
            else:
                # Need to add a new row
                new_row = table.add_row()
                
                new_row.cells[0].text = item_code
                new_row.cells[1].text = description
                new_row.cells[2].text = quantity
                new_row.cells[3].text = ''
                new_row.cells[4].text = ''
                
                # Apply formatting
                for cell in new_row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
    
    print("   ✅ Delivery Note filled successfully")
    return doc


def generate_delivery_note(
    so_data: Dict[str, Any],
    items: List[Dict[str, Any]],
    booking_number: str = ''
) -> Dict[str, Any]:
    """
    Main function: Generate delivery note for Axel France order
    
    Returns dictionary with:
    {
        'success': True/False,
        'filepath': path to generated file,
        'filename': filename,
        'error': error message if any
    }
    """
    print("="*80)
    print("DELIVERY NOTE GENERATOR - AXEL FRANCE")
    print("="*80)
    
    # Check if this is an Axel France order
    if not is_axel_france_order(so_data):
        print("⏭️  Not an Axel France order - skipping delivery note")
        return {
            'success': False,
            'error': 'Not an Axel France order',
            'skipped': True
        }
    
    try:
        # Fill the document
        doc = fill_delivery_note(so_data, items, booking_number)
        
        # Save the document
        output_dir = "generated_documents"
        os.makedirs(output_dir, exist_ok=True)
        
        so_number = so_data.get('so_number', 'UNKNOWN')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"DeliveryNote_SO{so_number}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        
        print(f"\n✅ Delivery Note Generated: {filename}")
        print("="*80)
        
        return {
            'success': True,
            'filepath': filepath,
            'filename': filename
        }
        
    except Exception as e:
        print(f"\n❌ Error generating delivery note: {e}")
        import traceback
        traceback.print_exc()
        print("="*80)
        
        return {
            'success': False,
            'error': str(e)
        }


# For testing
if __name__ == '__main__':
    # Test data
    test_so_data = {
        'so_number': '3100',
        'customer_name': 'AXEL France',
        'po_number': '27863',
        'shipping_address': {
            'company': 'AXEL France',
            'address': '30 rue de Pied de Fond\nZI St Liguaire CS 98821',
            'city': 'Noirt F-79000 Niort',
            'country': 'France'
        }
    }
    
    test_items = [
        {'item_code': 'MOV LL 2', 'description': 'Lubricating grease', 'quantity': 2, 'unit': 'drum'},
        {'item_code': 'MOV LL 0', 'description': 'Lubricating grease', 'quantity': 12, 'unit': 'drum'},
    ]
    
    result = generate_delivery_note(test_so_data, test_items, '83203363')
    print(result)

