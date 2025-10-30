#!/usr/bin/env python3
"""
Detailed BOL Template Analysis
Deep dive into the actual BOL template structure to understand exact field mappings
"""

from docx import Document
import os

def analyze_bol_template_detailed():
    """Detailed analysis of the BOL template"""
    
    template_path = r"G:\Shared drives\IT_Automation\Automating Roles\Logistics\Bill Of Landing\Canoil Skeletal BOL Template.docx"
    
    print("🔍 DETAILED BOL TEMPLATE ANALYSIS")
    print("=" * 60)
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return
    
    try:
        doc = Document(template_path)
        print(f"✅ Template loaded successfully")
        print(f"📊 Document structure: {len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs")
        
        # Focus on the critical tables
        analyze_address_table(doc.tables[2] if len(doc.tables) > 2 else None)
        analyze_items_table(doc.tables[4] if len(doc.tables) > 4 else None)
        
    except Exception as e:
        print(f"❌ Error analyzing template: {e}")

def analyze_address_table(table):
    """Analyze Table 2 - Address and Date information"""
    if not table:
        print("❌ Address table not found")
        return
        
    print(f"\n📋 TABLE 2 - ADDRESS & DATE ANALYSIS:")
    print(f"   Structure: {len(table.rows)} rows × {len(table.columns)} columns")
    
    for row_idx, row in enumerate(table.rows):
        print(f"\n   ROW {row_idx}:")
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            
            # Check for highlighting
            has_highlighting = check_cell_highlighting(cell)
            highlight_marker = " 🟡" if has_highlighting else ""
            
            # Identify field types
            field_type = identify_field_type(cell_text, row_idx, col_idx)
            
            if cell_text:
                print(f"     Col {col_idx}: {repr(cell_text[:100])}{highlight_marker}")
                if field_type:
                    print(f"              → {field_type}")
            else:
                print(f"     Col {col_idx}: [EMPTY]{highlight_marker}")

def analyze_items_table(table):
    """Analyze Table 4 - Items and details"""
    if not table:
        print("❌ Items table not found")
        return
        
    print(f"\n📋 TABLE 4 - ITEMS & DETAILS ANALYSIS:")
    print(f"   Structure: {len(table.rows)} rows × {len(table.columns)} columns")
    
    # Focus on key rows
    key_rows = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 14]  # Headers, items, batch, dangerous goods, totals
    
    for row_idx in key_rows:
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            print(f"\n   ROW {row_idx}:")
            
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # Check for highlighting
                has_highlighting = check_cell_highlighting(cell)
                highlight_marker = " 🟡" if has_highlighting else ""
                
                # Identify what should go in this cell
                field_purpose = identify_items_field_purpose(cell_text, row_idx, col_idx)
                
                if cell_text:
                    print(f"     Col {col_idx}: {repr(cell_text[:80])}{highlight_marker}")
                    if field_purpose:
                        print(f"              → {field_purpose}")
                else:
                    print(f"     Col {col_idx}: [EMPTY]{highlight_marker}")
                    if field_purpose:
                        print(f"              → {field_purpose}")

def check_cell_highlighting(cell):
    """Check if a cell has highlighting"""
    try:
        # Check for run-level highlighting
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color:
                    return True
        
        # Check for cell shading
        if hasattr(cell, '_tc') and hasattr(cell._tc, 'tcPr'):
            shd_elements = cell._tc.tcPr.xpath('.//w:shd')
            if shd_elements:
                return True
                
        return False
    except:
        return False

def identify_field_type(text, row, col):
    """Identify what type of field this is in the address table"""
    text_lower = text.lower()
    
    if row == 0:
        if 'date' in text_lower and '.' in text:
            return "DATE FIELD - needs user-selected ship date"
        elif 'pro number' in text_lower:
            return "PRO NUMBER FIELD - optional"
    elif row == 3:
        if 'consignee' in text_lower:
            return "CONSIGNEE NAME - needs customer name"
    elif row == 4:
        if 'street' in text_lower and '.' in text:
            return "CONSIGNEE ADDRESS - needs shipping address"
    elif row == 5:
        if 'city' in text_lower and '.' in text:
            return "CITY/PROVINCE - needs parsed city/province"
        elif 'postal' in text_lower and '.' in text:
            return "POSTAL CODE - needs parsed postal/zip"
    
    return None

def identify_items_field_purpose(text, row, col):
    """Identify what should go in each items table cell"""
    text_lower = text.lower()
    
    if row == 0:  # Headers
        return f"HEADER: {text}"
    elif row == 1:  # Main item row
        if col == 0 and '8' in text:
            return "PIECES - total quantity of all items"
        elif col == 1 and 'product' in text_lower:
            return "PRODUCT DESCRIPTION - formatted item descriptions"
        elif col == 2 and 'weight' in text_lower:
            return "WEIGHT - total weight from email/SO"
    elif row == 2:  # Batch row
        if col == 1 and 'batch' in text_lower:
            return "BATCH NUMBER - from email parsing"
    elif row == 4:  # Dangerous goods
        if col == 1 and 'dangerous' in text_lower:
            return "DANGEROUS GOODS - auto-detect Reolube"
    elif row == 5:  # Class 9 or PO
        if col == 1 and ('class 9' in text_lower or 'po#' in text_lower):
            return "CLASS 9 (if Reolube) OR PO# (if not dangerous)"
    elif row == 6:  # UN# or SO#
        if col == 1 and ('un#' in text_lower or 'so #' in text_lower):
            return "UN# 3082 (if Reolube 46XC) OR SO# (if not dangerous)"
    elif row == 7:  # Packing Group or Ref#
        if col == 1 and ('packing group' in text_lower or 'ref #' in text_lower):
            return "PACKING GROUP III (if Reolube 46XC) OR REF# (if not dangerous)"
    elif row == 9:  # Skid sizing
        if col == 1 and 'skid' in text_lower:
            return "SKID SIZING - from email pallet info"
    elif row == 14:  # Totals
        if col == 0 and 'total' in text_lower:
            return "TOTAL PIECES - sum of all quantities"
        elif col == 2 and 'total weight' in text_lower:
            return "TOTAL WEIGHT - total weight summary"
    
    return None

if __name__ == "__main__":
    analyze_bol_template_detailed()
