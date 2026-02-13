from flask import Flask, jsonify, request, send_file, Response
from flask_cors import CORS
from flask_compress import Compress
import os
import json
import base64
import re
import time
from datetime import datetime, timedelta
import glob
from dotenv import load_dotenv
import requests
import PyPDF2
import pdfplumber
from docx import Document
from enterprise_analytics import EnterpriseAnalytics
import sys
import io
import requests
import zipfile
import csv
import xml.etree.ElementTree as ET

# Ensure console logging works on Windows with Unicode characters (emojis, symbols)
# Only wrap if not already wrapped to avoid closing file handles
try:
    if hasattr(sys.stdout, "buffer") and not isinstance(sys.stdout, io.TextIOWrapper):
        # Store original to prevent closure issues
        original_stdout = sys.stdout
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    if hasattr(sys.stderr, "buffer") and not isinstance(sys.stderr, io.TextIOWrapper):
        # Store original to prevent closure issues
        original_stderr = sys.stderr
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)
except Exception:
    # Never crash the app just because logging reconfiguration failed
    pass

# Load environment variables
load_dotenv()

# MPS (Master Production Schedule) constants
MPS_SHEET_ID = '1zAOY7ngP2mLVi-W_FL9tsPiKDPqbU6WEUmrrTDeKygw'
MPS_CSV_URL = f'https://docs.google.com/spreadsheets/d/{MPS_SHEET_ID}/export?format=csv'

# Import logistics automation module
try:
    from .logistics_automation import logistics_bp
    LOGISTICS_AVAILABLE = True
except ImportError:
    try:
        from logistics_automation import logistics_bp
        LOGISTICS_AVAILABLE = True
    except ImportError as e:
        print(f"Logistics module not available: {e}")
        LOGISTICS_AVAILABLE = False

# Import purchase requisition service
try:
    from .purchase_requisition_service import pr_service
    PR_SERVICE_AVAILABLE = True
except ImportError:
    try:
        from purchase_requisition_service import pr_service
        PR_SERVICE_AVAILABLE = True
    except ImportError as e:
        print(f"Purchase Requisition service not available: {e}")
        PR_SERVICE_AVAILABLE = False

# BOL HTML module removed - using the one in logistics_automation.py
BOL_HTML_AVAILABLE = False

def safe_float(value):
    """Safely convert value to float, handling commas and None values"""
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        # Remove commas and convert
        cleaned = value.replace(',', '').strip()
        if cleaned == '' or cleaned == 'None':
            return 0.0
        try:
            return float(cleaned)
        except ValueError:
            return 0.0
    return 0.0

def extract_so_data_from_docx(docx_path):
    """Extract Sales Order data from Word document - MUCH EASIER than PDF!"""
    try:
        so_data = {
            'so_number': '',
            'customer_name': '',
            'order_date': '',
            'total_amount': 0.0,
            'items': [],
            'status': '',
            'raw_text': ''
        }
        
        # Open Word document
        doc = Document(docx_path)
        
        # Extract all text
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Extract from tables (much more reliable in Word!)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                full_text += row_text + "\n"
        
        so_data['raw_text'] = full_text
        
        # Extract SO number from filename or text
        filename = os.path.basename(docx_path)
        so_match = re.search(r'salesorder[_\s]*(\d+)', filename, re.IGNORECASE)
        if so_match:
            so_data['so_number'] = so_match.group(1)
        
        # Extract from text if not in filename
        if not so_data['so_number']:
            so_text_match = re.search(r'(?:SO|Order No\.?|Sales Order)[:\s#]*(\d+)', full_text, re.IGNORECASE)
            if so_text_match:
                so_data['so_number'] = so_text_match.group(1)
        
        # Extract customer name (easier in Word!)
        customer_patterns = [
            r'Sold To:\s*([^\n\r|]+)',
            r'Customer:\s*([^\n\r|]+)',
            r'Bill To:\s*([^\n\r|]+)',
            r'Ship To:\s*([^\n\r|]+)',
            r'Company:\s*([^\n\r|]+)'
        ]
        
        for pattern in customer_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                customer_name = match.group(1).strip()
                # Clean up
                if ' - CDN$' in customer_name:
                    customer_name = customer_name.split(' - CDN$')[0].strip()
                if '|' in customer_name:
                    customer_name = customer_name.split('|')[0].strip()
                so_data['customer_name'] = customer_name
                break
        
        # Extract order date
        date_patterns = [
            r'(?:Order Date|Date):\s*([^\n\r|]+)',
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(\d{4}[/-]\d{1,2}[/-]\d{1,2})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                so_data['order_date'] = match.group(1).strip()
                break
        
        # Extract total amount
        amount_patterns = [
            r'Total:\s*\$?([\d,]+\.?\d*)',
            r'Amount Due:\s*\$?([\d,]+\.?\d*)',
            r'Grand Total:\s*\$?([\d,]+\.?\d*)',
            r'\$([\d,]+\.?\d*)'
        ]
        
        for pattern in amount_patterns:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            if matches:
                amounts = [safe_float(amount) for amount in matches]
                so_data['total_amount'] = max(amounts) if amounts else 0.0
                break
        
        # Extract items from tables (MUCH better in Word!)
        for table in doc.tables:
            headers = [cell.text.lower() for cell in table.rows[0].cells]
            if any(word in ' '.join(headers) for word in ['item', 'description', 'product']):
                for row in table.rows[1:]:  # Skip header
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data if cell):  # Non-empty row
                        so_data['items'].append(' | '.join(row_data))
        
        # Determine status from path
        if 'cancelled' in docx_path.lower():
            so_data['status'] = 'Cancelled'
        elif 'completed' in docx_path.lower() or 'closed' in docx_path.lower():
            so_data['status'] = 'Completed'
        elif 'production' in docx_path.lower() or 'scheduled' in docx_path.lower():
            so_data['status'] = 'In Production'
        else:
            so_data['status'] = 'Unknown'
        
        return so_data
        
    except Exception as e:
        print(f"Error extracting SO data from {docx_path}: {str(e)}")
        return None

def extract_addresses_from_layout(pdf_path):
    """
    Extract Sold To and Ship To addresses using word positions.
    Uses x-coordinates to separate left column (Sold To) from right column (Ship To).
    Also extracts batch number and MO number from the far right.
    """
    sold_to_lines = []
    ship_to_lines = []
    batch_number = ''
    mo_number = ''
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                words = page.extract_words()
                if not words:
                    continue
                
                # Find the "Sold To:" and "Ship To:" positions
                sold_to_x = None
                ship_to_x = None
                address_start_y = None
                
                for w in words:
                    text_lower = w['text'].lower()
                    if text_lower == 'sold' or text_lower == 'sold:':
                        sold_to_x = w['x0']
                        address_start_y = w['top']
                    elif text_lower == 'ship' and ship_to_x is None:
                        # Only set ship_to_x if it's after sold_to (the header, not "Ship Date")
                        if sold_to_x and w['x0'] > sold_to_x + 100:
                            ship_to_x = w['x0']
                
                if not sold_to_x or not ship_to_x or not address_start_y:
                    continue
                
                # Calculate column boundaries
                # Left column (Sold To): x0 < midpoint
                # Right column (Ship To): x0 >= midpoint but < far right
                midpoint = (sold_to_x + ship_to_x) / 2 + 50  # A bit past the midpoint
                far_right = 400  # MO/Batch area starts around x=400
                
                # Find address end (Business No line or Item No line)
                address_end_y = address_start_y + 150  # Default: 150 pixels down
                for w in words:
                    text_upper = w['text'].upper()
                    if text_upper in ['BUSINESS', 'ITEM'] and w['top'] > address_start_y:
                        address_end_y = w['top']
                        break
                
                # Group words by approximate Y position (same line)
                left_lines = {}
                right_lines = {}
                
                for w in words:
                    # Only process words in the address area
                    if w['top'] < address_start_y or w['top'] > address_end_y:
                        continue
                    
                    # Skip headers
                    text_lower = w['text'].lower()
                    if text_lower in ['sold', 'to:', 'ship', 'to']:
                        continue
                    
                    # Extract batch number and MO from far right
                    if w['x0'] > far_right:
                        text_upper = w['text'].upper()
                        if 'BATCH' in text_upper or 'WH' in text_upper:
                            # Look for batch number pattern
                            import re
                            if re.match(r'WH\d+[A-Z]\d+', text_upper):
                                batch_number = w['text']
                        elif text_upper == 'MO':
                            # Next word should be the MO number
                            pass
                        elif re.match(r'^\d{4}$', w['text']):
                            mo_number = w['text']
                        continue
                    
                    # Round Y to group words on same line (within 3 pixels)
                    y_key = round(w['top'] / 3) * 3
                    
                    if w['x0'] < midpoint:
                        # Left column (Sold To)
                        if y_key not in left_lines:
                            left_lines[y_key] = []
                        left_lines[y_key].append((w['x0'], w['text']))
                    else:
                        # Right column (Ship To)
                        if y_key not in right_lines:
                            right_lines[y_key] = []
                        right_lines[y_key].append((w['x0'], w['text']))
                
                # Sort and join words on each line
                for y in sorted(left_lines.keys()):
                    line_words = sorted(left_lines[y], key=lambda x: x[0])
                    line_text = ' '.join(w[1] for w in line_words).strip()
                    if line_text and line_text.upper() not in ['SOLD TO:', 'SOLD TO', 'TO:']:
                        sold_to_lines.append(line_text)
                
                for y in sorted(right_lines.keys()):
                    line_words = sorted(right_lines[y], key=lambda x: x[0])
                    line_text = ' '.join(w[1] for w in line_words).strip()
                    if line_text and line_text.upper() not in ['SHIP TO:', 'SHIP TO', 'TO:']:
                        # Skip MO references and batch numbers that might slip through
                        if not line_text.startswith('MO ') and 'Batch' not in line_text:
                            ship_to_lines.append(line_text)
                
                # Also look for batch number in a different pattern
                for w in words:
                    if 'Batch' in w['text'] or 'batch' in w['text']:
                        # Find the next word which should be "no:" or the batch number
                        idx = words.index(w)
                        for next_w in words[idx:idx+4]:
                            import re
                            if re.match(r'WH\d+[A-Z]\d+', next_w['text'].upper()):
                                batch_number = next_w['text']
                                break
                    elif w['text'].upper() == 'MO' and not mo_number:
                        idx = words.index(w)
                        for next_w in words[idx+1:idx+3]:
                            if re.match(r'^\d{4}$', next_w['text']):
                                mo_number = next_w['text']
                                break
                
    except Exception as e:
        print(f"⚠️ Layout address extraction failed: {e}")
        import traceback
        traceback.print_exc()
    
    result = {
        'sold_to_raw': '\n'.join(sold_to_lines),
        'ship_to_raw': '\n'.join(ship_to_lines),
        'batch_number': batch_number,
        'mo_number': mo_number
    }
    
    print(f"📍 LAYOUT EXTRACTION:")
    print(f"   Sold To: {result['sold_to_raw'][:100]}...")
    print(f"   Ship To: {result['ship_to_raw'][:100]}...")
    print(f"   Batch: {batch_number}, MO: {mo_number}")
    
    return result

def extract_addresses_from_pdf_tables(pdf_path):
    """
    PRE-EXTRACT Sold To and Ship To addresses from PDF tables.
    The PDF has these in a two-column table - LEFT is Sold To, RIGHT is Ship To.
    This is MORE RELIABLE than trying to parse merged text.
    """
    sold_to_lines = []
    ship_to_lines = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if not tables:
                    continue
                    
                found_address_header = False
                
                for table in tables:
                    if not table:
                        continue
                        
                    for row_idx, row in enumerate(table):
                        if not row:
                            continue
                        
                        # Handle different row lengths
                        left_cell = str(row[0] or '').strip() if len(row) > 0 else ''
                        right_cell = str(row[1] or '').strip() if len(row) > 1 else ''
                        
                        left_upper = left_cell.upper()
                        right_upper = right_cell.upper()
                        
                        # Detect "Sold To" / "Ship To" header row
                        if 'SOLD TO' in left_upper or 'BILL TO' in left_upper:
                            found_address_header = True
                            continue
                        
                        # Stop when we hit item table or other sections
                        if found_address_header:
                            stop_keywords = ['ITEM', 'ORDERED', 'BUSINESS NO', 'QTY', 'UNIT PRICE', 'AMOUNT', 'DESCRIPTION']
                            if any(kw in left_upper for kw in stop_keywords):
                                break
                        
                        # Collect address lines - LEFT = Sold To, RIGHT = Ship To
                        if found_address_header:
                            skip_labels = ['SOLD TO:', 'SHIP TO:', 'BILL TO:', 'SOLD TO', 'SHIP TO', 'BILL TO']
                            
                            if left_cell and left_cell.upper() not in [s.upper() for s in skip_labels]:
                                sold_to_lines.append(left_cell)
                            
                            if right_cell and right_cell.upper() not in [s.upper() for s in skip_labels]:
                                ship_to_lines.append(right_cell)
    except Exception as e:
        print(f"⚠️ Table address extraction failed: {e}")
    
    return {
        'sold_to_raw': '\n'.join(sold_to_lines),
        'ship_to_raw': '\n'.join(ship_to_lines)
    }


def extract_so_data_from_pdf(pdf_path):
    """
    SIMPLE SO PARSER - Just extract what's there, no complex logic
    Based on actual analysis of SO PDFs - handles all formats
    ENHANCED: Uses layout-preserving extraction and table parsing for better address handling
    """
    try:
        print(f"PARSING: {os.path.basename(pdf_path)}")
        
        # FIRST: Try layout-based extraction (uses word positions - most reliable)
        layout_addresses = extract_addresses_from_layout(pdf_path)
        
        # SECOND: Try table-based extraction as fallback
        table_addresses = extract_addresses_from_pdf_tables(pdf_path)
        
        # Use layout addresses if available, otherwise fall back to table addresses
        if layout_addresses.get('sold_to_raw') or layout_addresses.get('ship_to_raw'):
            pre_extracted = layout_addresses
            print(f"📍 Using LAYOUT-based address extraction")
        elif table_addresses.get('sold_to_raw') or table_addresses.get('ship_to_raw'):
            pre_extracted = table_addresses
            print(f"📋 Using TABLE-based address extraction")
        else:
            pre_extracted = {'sold_to_raw': '', 'ship_to_raw': '', 'batch_number': '', 'mo_number': ''}
        
        # Extract text from PDF with layout preservation for better column handling
        with pdfplumber.open(pdf_path) as pdf:
            full_text = ""
            layout_text = ""  # Layout-preserved text for address parsing
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
                # Also get layout-preserved text
                layout_page_text = page.extract_text(layout=True)
                if layout_page_text:
                    layout_text += layout_page_text + "\n"
        
        lines = full_text.split('\n')
        
        # DEBUG: Print all lines that might contain items (have unit keywords)
        print(f"DEBUG SO PARSING: Total lines: {len(lines)}")
        for i, line in enumerate(lines):
            line_upper = line.upper()
            if any(unit in line_upper for unit in ['LITER', 'LITRE', 'DRUM', 'PAIL', 'CASE', 'KEG', 'GALLON', 'TOTE', 'BULK']):
                print(f"DEBUG LINE {i}: '{line}'")
        
        # Initialize result
        so_data = {
            'so_number': '',
            'customer_name': '',
            'customer_address': '',
            'customer_phone': '',
            'customer_email': '',
            'order_date': '',
            'due_date': '',
            'po_number': '',
            'terms': '',
            'subtotal': 0.0,
            'tax': 0.0,
            'total_amount': 0.0,
            'items': [],
            'status': '',
            'raw_text': full_text,
            'billing_address': {},
            'shipping_address': {},
            'business_number': '',
            'line_references': [],
            'special_instructions': '',
            'file_info': {
                'filename': os.path.basename(pdf_path),
                'size_bytes': os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0
            },
            'sold_to': {
                'company_name': '',
                'contact_person': '',
                'address': '',
                'phone': '',
                'email': ''
            },
            'ship_to': {
                'company_name': '',
                'contact_person': '',
                'address': '',
                'phone': '',
                'email': ''
            }
        }
        
        # Find the Sold To: Ship To: line
        sold_to_ship_to_line = -1
        for i, line in enumerate(lines):
            if 'Sold To:' in line and 'Ship To:' in line:
                sold_to_ship_to_line = i
                break
        
        if sold_to_ship_to_line >= 0:
            # CLEAN ADDRESS PARSING - Handle TWO-COLUMN PDF layout
            # PDFs have Sold To on left, Ship To on right - pdfplumber extracts them in various ways:
            # 1. Same line with space between: "Company A Company B" 
            # 2. Consecutive duplicate lines: "26 Wooley Road" appears twice
            # 3. Two-column with double-space: "Address1  Address2"
            
            # Stop keywords (case-insensitive) - anything after these is NOT address
            stop_patterns = [
                r'^MO\s+\d+',  # Manufacturing Order reference (e.g., "MO 3804" or "MO 3448, Ship with...")
                r'^Line\s+\d+:',  # Line references (e.g., "Line 1: MO 3656")
                r'business\s+no', r'item\s+no',
                r'subtotal', r'hst', r'total\s+amount',
                r'terms:', r'\d+\s+drum', r'\d+\s+pail', r'\d+\s+gallon', r'\d+\s+case', r'\d+\s+keg'
            ]
            stop_regex = re.compile('|'.join(stop_patterns), re.IGNORECASE)
            
            # Collect raw address lines - but SKIP consecutive duplicates
            raw_address_lines = []
            prev_line = None
            for i in range(sold_to_ship_to_line + 1, min(sold_to_ship_to_line + 20, len(lines))):
                line = lines[i].strip()
                if not line:
                    continue
                # Stop if entire line matches stop pattern
                if stop_regex.search(line):
                    break
                # Skip consecutive duplicates (two-column PDF artifact)
                if line == prev_line:
                    prev_line = line
                    continue
                raw_address_lines.append(line)
                prev_line = line
            
            def deduplicate_merged_text(text):
                """
                Handle PDF two-column merging that causes duplicates on same line.
                E.g., "ECO Undercoating LLC ECO Undercoating LLC" -> "ECO Undercoating LLC"
                E.g., "Georgia Western Inc. Georgia Western Inc." -> "Georgia Western Inc."
                E.g., "Mil-Comm Products Company Inc Pflaumer·warehouse" -> split into two companies
                """
                if not text:
                    return text, text
                
                # Method 1: Check for double-space separation (two different values)
                if '  ' in text:
                    parts = [p.strip() for p in re.split(r'\s{2,}', text) if p.strip()]
                    if len(parts) == 2:
                        if parts[0] == parts[1]:
                            # Same value duplicated
                            return parts[0], parts[0]
                        else:
                            # Two different values (Sold To / Ship To)
                            return parts[0], parts[1]
                
                # Method 2: Check if the text is the same phrase repeated with single space
                # E.g., "ECO Undercoating LLC ECO Undercoating LLC"
                words = text.split()
                if len(words) >= 2 and len(words) % 2 == 0:
                    half = len(words) // 2
                    first_half = ' '.join(words[:half])
                    second_half = ' '.join(words[half:])
                    if first_half == second_half:
                        return first_half, first_half
                    # Check if they're different (Sold To vs Ship To company)
                    # Look for company suffixes to validate
                    company_suffixes = ['Inc', 'LLC', 'Ltd', 'Corp', 'Co', 'Company', 'Limited']
                    first_is_company = any(s.lower() in first_half.lower() for s in company_suffixes)
                    second_is_company = any(s.lower() in second_half.lower() for s in company_suffixes)
                    if first_is_company and second_is_company:
                        return first_half, second_half
                
                # Method 3: Check for special separator like · (middle dot) or known patterns
                # E.g., "Mil-Comm Products Company Inc Pflaumer·warehouse"
                if '·' in text:
                    # Split at the middle dot
                    parts = text.split('·')
                    if len(parts) == 2:
                        # The part before · might have the company name concatenated
                        # Look for company suffix to find the split point
                        before_dot = parts[0].strip()
                        after_dot = parts[1].strip()
                        
                        # Find company suffix in before_dot
                        company_suffixes = ['Inc', 'LLC', 'Ltd', 'Corp', 'Co', 'Company', 'Limited']
                        for suffix in company_suffixes:
                            if suffix in before_dot:
                                # Find where the suffix ends
                                suffix_idx = before_dot.rfind(suffix)
                                if suffix_idx > 0:
                                    sold_to = before_dot[:suffix_idx + len(suffix)].strip()
                                    ship_to_prefix = before_dot[suffix_idx + len(suffix):].strip()
                                    ship_to = (ship_to_prefix + ' ' + after_dot).strip()
                                    return sold_to, ship_to
                        
                        # Fallback: just split at ·
                        return before_dot, after_dot
                
                # No duplication detected
                return text, text
            
            def extract_phone(text):
                """Extract phone number from text"""
                match = re.search(r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})', text)
                return match.group(1) if match else ''
            
            def clean_line(text):
                """Remove phone numbers, batch info, and other non-address data"""
                if not text:
                    return ''
                # Remove phone numbers
                text = re.sub(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', '', text)
                # Remove batch/pull info
                text = re.sub(r',?\s*pull\s+from.*$', '', text, flags=re.IGNORECASE)
                text = re.sub(r',?\s*batch\s*#?\s*\S+.*$', '', text, flags=re.IGNORECASE)
                # Remove N/A
                text = re.sub(r'\bN/?A\b', '', text, flags=re.IGNORECASE)
                # Clean up
                text = re.sub(r'\s+', ' ', text).strip()
                return text
            
            # Detect if we have alternating lines (different Sold To and Ship To addresses)
            # vs duplicated lines (same address for both)
            # Pattern 1: Duplicated - lines repeat consecutively (already filtered out above)
            # Pattern 2: Alternating - Sold To and Ship To on alternating lines
            
            # Check if addresses are different by looking at the structure
            # If we have an odd number of address lines after company, it's likely alternating
            is_alternating = False
            if len(raw_address_lines) > 3:
                # Check if lines look like they alternate (different content)
                # Heuristic: if line 2 and line 3 are very different (one is name, one is address)
                if len(raw_address_lines) >= 3:
                    line1 = raw_address_lines[1] if len(raw_address_lines) > 1 else ''
                    line2 = raw_address_lines[2] if len(raw_address_lines) > 2 else ''
                    # If one looks like a street address and the other looks like a name
                    is_street1 = bool(re.search(r'\d+\s+\w+', line1))  # Has street number
                    is_street2 = bool(re.search(r'\d+\s+\w+', line2))  # Has street number
                    is_name1 = bool(re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+$', line1))  # Looks like a name
                    is_name2 = bool(re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+$', line2))  # Looks like a name
                    
                    # If one is street and other is name, or both are streets but different
                    if (is_street1 and is_name2) or (is_name1 and is_street2):
                        is_alternating = True
                    elif is_street1 and is_street2 and line1 != line2:
                        # Both are streets but different - could be alternating
                        is_alternating = True
            
            # Parse addresses
            sold_to_parts = []
            ship_to_parts = []
            phone = ''
            sold_to_company = ''
            ship_to_company = ''
            ship_to_contact = ''
            detected_country = ''  # Track the country from raw address lines
            is_pickup_order = False  # Track if this is a "PICK UP" order (no shipping needed)
            
            # Helper to detect and clean "PICK UP" from text
            def clean_pickup_instruction(text):
                """Remove PICK UP instruction from text and return (cleaned_text, is_pickup)"""
                if not text:
                    return text, False
                # Patterns for pickup instructions
                pickup_patterns = [
                    r'\s*-\s*PICK\s*UP\s*$',  # "- PICK UP" at end
                    r'\s*PICK\s*UP\s*-\s*',   # "PICK UP -" in middle
                    r'^\s*PICK\s*UP\s*$',     # Just "PICK UP"
                    r'\s*\(\s*PICK\s*UP\s*\)',  # "(PICK UP)"
                ]
                is_pickup = False
                cleaned = text
                for pattern in pickup_patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        is_pickup = True
                        cleaned = re.sub(pattern, '', text, flags=re.IGNORECASE).strip()
                return cleaned, is_pickup
            
            for idx, line in enumerate(raw_address_lines):
                # Extract phone first
                line_phone = extract_phone(line)
                if line_phone and not phone:
                    phone = line_phone
                
                # Clean the line
                clean = clean_line(line)
                if not clean:
                    continue
                
                # Deduplicate merged text
                sold_part, ship_part = deduplicate_merged_text(clean)
                
                if idx == 0:
                    # First line is company name(s)
                    # ONLY clean PICK UP from Sold To - we want to keep it as a note for Ship To
                    sold_to_company, _ = clean_pickup_instruction(sold_part)  # Always clean from Sold To
                    
                    # For Ship To, check if it has PICK UP and store as instruction
                    ship_to_company_raw, ship_pickup = clean_pickup_instruction(ship_part)
                    
                    if ship_pickup:
                        is_pickup_order = True
                        print(f"📦 PICK UP ORDER DETECTED - Will show as instruction under Ship To")
                    
                    # Check if ship_part is a contact person (no company suffix) or a company
                    company_suffixes = ['Inc', 'LLC', 'Ltd', 'Corp', 'Co.', 'Company', 'Limited', 'warehouse', 'Warehouse']
                    if sold_to_company != ship_to_company_raw and ship_to_company_raw:
                        if any(s.lower() in ship_to_company_raw.lower() for s in company_suffixes):
                            ship_to_company = ship_to_company_raw.replace('·', ' ').strip()
                        else:
                            # It's a contact person for Ship To
                            ship_to_company = sold_to_company
                            ship_to_contact = ship_to_company_raw
                    else:
                        # Use cleaned company name (without PICK UP) for Ship To company
                        ship_to_company = ship_to_company_raw if ship_to_company_raw else sold_to_company
                else:
                    # Capture country but don't add to address parts
                    if clean.lower() in ['usa', 'u.s.a.', 'u.s.', 'us', 'united states']:
                        detected_country = 'USA'
                        continue
                    elif clean.lower() in ['canada', 'ca']:
                        detected_country = 'Canada'
                        continue
                    elif clean.lower() in ['japan', 'jp']:
                        detected_country = 'Japan'
                        continue
                    elif clean.lower() in ['puerto rico', 'pr']:
                        detected_country = 'Puerto Rico'
                        continue
                    
                    if is_alternating and sold_part == ship_part:
                        # Alternating pattern: odd lines (1, 3, 5...) go to Sold To
                        # Even lines (2, 4, 6...) go to Ship To
                        # But first check if this looks like a contact name
                        is_contact_name = bool(re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+$', clean))
                        
                        if is_contact_name and not ship_to_contact:
                            ship_to_contact = clean
                        elif idx % 2 == 1:  # Odd index (1, 3, 5...)
                            sold_to_parts.append(clean)
                        else:  # Even index (2, 4, 6...)
                            ship_to_parts.append(clean)
                    else:
                        # Same address for both (duplicated pattern)
                        if sold_part not in sold_to_parts:
                            sold_to_parts.append(sold_part)
                        if ship_part != sold_part and ship_part not in ship_to_parts:
                            ship_to_parts.append(ship_part)
            
            # Set company names
            so_data['sold_to']['company_name'] = sold_to_company
            so_data['ship_to']['company_name'] = ship_to_company or sold_to_company
            so_data['ship_to']['contact_person'] = ship_to_contact
            
            # Set phones
            if phone:
                so_data['sold_to']['phone'] = phone
                so_data['ship_to']['phone'] = phone
            
            # Build clean addresses - deduplicate parts AND remove stock comments
            def clean_address_part(part):
                """Clean a single address part - remove stock comments, batch numbers, etc."""
                if not part:
                    return ''
                cleaned = part.strip()
                # Remove stock-related comments (CRITICAL - these should NEVER be in addresses)
                cleaned = re.sub(r',?\s*pull\s+from\s+stock.*$', '', cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r',?\s*pull\s+from.*$', '', cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r',?\s*batch\s*#?\s*\S+.*$', '', cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r',?\s*lot\s*#?\s*\S+.*$', '', cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r',?\s*stock.*$', '', cleaned, flags=re.IGNORECASE)
                # Remove trailing commas and spaces
                cleaned = re.sub(r',\s*$', '', cleaned)
                cleaned = cleaned.strip()
                return cleaned
            
            def build_address(parts):
                # Remove duplicates while preserving order AND clean each part
                seen = set()
                unique_parts = []
                for p in parts:
                    p_clean = clean_address_part(p)  # Clean BEFORE checking duplicates
                    if p_clean and p_clean.lower() not in seen:
                        seen.add(p_clean.lower())
                        unique_parts.append(p_clean)
                # Join and clean the final result one more time (in case cleaning created issues)
                result = ', '.join(unique_parts)
                # Final pass: remove any remaining stock comments that might have slipped through
                result = re.sub(r',?\s*pull\s+from\s+stock.*$', '', result, flags=re.IGNORECASE)
                result = re.sub(r',?\s*batch\s*#?\s*\S+.*$', '', result, flags=re.IGNORECASE)
                result = re.sub(r',\s*,+', ',', result)  # Multiple commas
                result = re.sub(r',\s*$', '', result)  # Trailing comma
                return result.strip()
            
            so_data['sold_to']['address'] = build_address(sold_to_parts)
            
            # Handle PICK UP orders - don't copy billing address for shipping
            if is_pickup_order:
                so_data['ship_to']['address'] = ''  # No shipping address for pickup
                so_data['is_pickup_order'] = True
                print(f"📦 Ship To address cleared - this is a PICKUP order")
            else:
                so_data['ship_to']['address'] = build_address(ship_to_parts) or build_address(sold_to_parts)
                so_data['is_pickup_order'] = False
            
            # Store detected country for later use
            if detected_country:
                so_data['detected_country'] = detected_country
        
        # Set customer name
        so_data['customer_name'] = so_data['sold_to']['company_name']
        
        # Parse other fields
        for line in lines:
            line = line.strip()
            
            # SO Number
            if line.startswith('Order No.:'):
                so_data['so_number'] = line.split('Order No.:')[1].strip()
                
            # Dates
            elif line.startswith('Date:'):
                so_data['order_date'] = line.split('Date:')[1].strip()
            elif line.startswith('Ship Date:'):
                so_data['ship_date'] = line.split('Ship Date:')[1].strip()
                so_data['due_date'] = so_data['ship_date']  # For compatibility
                
            # Business Number
            elif line.startswith('Business No.:'):
                so_data['business_number'] = line.split('Business No.:')[1].strip()
                
            # Items - detect product lines with ANY unit type from product catalog
            # Skip header line and charge lines (Pallet, Freight, Brokerage)
            # Comprehensive unit list based on actual item descriptions in catalog
            elif (re.search(r'\b(DRUM|DRUMS|PAIL|PAILS|GALLON|GAL|CASE|CASES|KEG|KEGS|LITER|LITRE|BOTTLE|BTL|BAG|BAGS|BOX|BOXES|CARTON|CTN|EACH|TOTE|TOTES|TUBE|TUBES|JAR|JARS|TUB|TUBS|CAN|CANS|IBC|ROLL|ROLLS|KG|LB|LBS|BULK)\d*\b', line.upper())
                  and not line.startswith('Item No.')
                  and not line.startswith('Pallet')
                  and not line.startswith('Freight')
                  and not line.startswith('Brokerage')):
                parts = line.split()
                if len(parts) >= 4:
                    # Find the unit type position - comprehensive list from product catalog
                    # CRITICAL: Order by specificity - longer/more specific units first
                    unit_types = [
                        'GALLON', 'CARTON', 'BOTTLE', 'DRUMS', 'PAILS', 'CASES', 'KEGS',
                        'LITER', 'LITRE', 'DRUM', 'PAIL', 'CASE', 'EACH', 'BAGS', 'TOTE',
                        'TOTES', 'TUBES', 'BOXES', 'ROLLS', 'JARS', 'TUBS', 'CANS',
                        'TUBE', 'BAG', 'BOX', 'KEG', 'JAR', 'TUB', 'CAN', 'IBC', 'GAL',
                        'BTL', 'CTN', 'LBS', 'LB', 'KG', 'BULK', 'ROLL'
                    ]
                    unit_idx = -1
                    unit = ''
                    quantity = ''
                    
                    # Search through all parts to find a valid unit with a numeric quantity before it
                    for i, part in enumerate(parts):
                        if i < 1:  # Skip first position - unit needs quantity before it
                            continue
                        part_upper = part.upper()
                        # Check for exact match or starts with unit type
                        for ut in unit_types:
                            if part_upper == ut or part_upper.startswith(ut):
                                # CRITICAL: Only accept if the position before is a number (quantity)
                                # Handle comma-formatted numbers like "1,187"
                                potential_qty = parts[i - 1].replace(',', '')
                                if potential_qty.isdigit():
                                    unit_idx = i
                                    unit = ut  # Normalize to just the unit type
                                    quantity = potential_qty
                                    break
                        if unit_idx >= 0:
                            break
                    
                    if unit_idx < 1 or not quantity:
                        # No valid unit found with numeric quantity before it
                        continue
                    
                    quantity_idx = unit_idx - 1
                    
                    # Item code is everything before quantity
                    item_code = ' '.join(parts[:quantity_idx])
                    
                    # Description is everything after unit until prices
                    desc_start = unit_idx + 1
                    desc_end = len(parts)
                    for j, part in enumerate(parts[desc_start:], desc_start):
                        # Stop at price patterns
                        if re.match(r'^[\d,]+\.\d+$', part) or re.match(r'^(?:US|CDN)?\$[\d,]+', part):
                            desc_end = j
                            break
                    
                    description = ' '.join(parts[desc_start:desc_end])
                    
                    # Find prices - handle commas, currency symbols, and various formats
                    # Pattern 1: Two numbers with commas (e.g., "8,865.00 US$17,730.00" or "8,865.00 17,730.00")
                    price_match = re.search(r'([\d,]+\.\d+)\s+(?:US\$|CDN\$|\$)?([\d,]+\.\d+)', line)
                    unit_price = 0.0
                    amount = 0.0
                    if price_match:
                        unit_price = float(price_match.group(1).replace(',', ''))
                        amount = float(price_match.group(2).replace(',', ''))
                    else:
                        # Pattern 2: Look for any numbers with decimals
                        all_numbers = re.findall(r'([\d,]+\.\d+)', line)
                        if len(all_numbers) >= 2:
                            # Last two numbers are usually unit price and total
                            unit_price = float(all_numbers[-2].replace(',', ''))
                            amount = float(all_numbers[-1].replace(',', ''))
                        elif len(all_numbers) == 1:
                            # Only one number - might be total, calculate unit price from quantity
                            amount = float(all_numbers[0].replace(',', ''))
                            try:
                                qty = int(quantity)
                                if qty > 0:
                                    unit_price = amount / qty
                            except:
                                pass
                    
                    try:
                        item = {
                            'item_code': item_code,
                            'description': description,
                            'quantity': int(quantity),
                            'unit': unit,
                            'unit_price': unit_price,
                            'amount': amount,
                            'total_price': amount,  # Frontend expects total_price
                            'price': unit_price  # Alias for compatibility
                        }
                        # Only add real product items, not charges
                        if item_code not in ['Pallet', 'Freight', 'Brokerage']:
                            so_data['items'].append(item)
                            print(f"DEBUG ITEM PARSED: code='{item_code}', desc='{description}', qty={quantity}, unit={unit}")
                    except ValueError:
                        # Skip items that don't parse correctly
                        continue
                    
            # Subtotal - handle currency prefixes like "US$" or "CDN$"
            elif line.startswith('Subtotal:'):
                # Pattern: "Subtotal: US$5,570.00" or "Subtotal: 5,570.00"
                subtotal_match = re.search(r'Subtotal:\s*(?:US\$|CDN\$|\$)?([\d,]+\.?\d*)', line)
                if subtotal_match:
                    so_data['subtotal'] = float(subtotal_match.group(1).replace(',', ''))
            
            # Tax (HST/GST)
            elif ('HST' in line or 'GST' in line) and re.search(r'([\d,]+\.?\d*)', line):
                tax_match = re.search(r'(?:HST|GST)\s*([\d,]+\.?\d*)', line)
                if tax_match:
                    so_data['tax'] = float(tax_match.group(1).replace(',', ''))
                    
            # Terms
            elif line.startswith('Terms:'):
                so_data['terms'] = line.split('Terms:')[1].strip()
                
            # Comment line - extract PO AND Total Amount
            # Pattern: "Comment: PO 329438 Total Amount US$5,570.00"
            elif 'Comment:' in line:
                # Extract PO from Comment line
                if not so_data['po_number']:
                    po_match = re.search(r'\bPO\s*[:#]?\s*(\d+[-\d]*)', line, re.IGNORECASE)
                    if po_match:
                        so_data['po_number'] = po_match.group(1)
                        print(f"  📋 Found PO in Comment: {so_data['po_number']}")
                # Also extract Total Amount if present in Comment line
                if 'Total Amount' in line and so_data['total_amount'] == 0:
                    total_match = re.search(r'Total Amount\s*(?:US\$|CDN\$|\$)?([\d,]+\.?\d*)', line)
                    if total_match:
                        so_data['total_amount'] = float(total_match.group(1).replace(',', ''))
                        
            # Total Amount - handle currency prefixes (standalone line)
            elif 'Total Amount' in line and so_data['total_amount'] == 0:
                # Pattern: "Total Amount US$5,570.00" or "Total Amount 5,570.00"
                total_match = re.search(r'Total Amount\s*(?:US\$|CDN\$|\$)?([\d,]+\.?\d*)', line)
                if total_match:
                    so_data['total_amount'] = float(total_match.group(1).replace(',', ''))
                    
            # Generic PO pattern (for lines without Comment)
            elif re.search(r'\bP\.?O\.?\s*[:#]?\s*\d', line, re.IGNORECASE) and not so_data['po_number']:
                po_match = re.search(r'\bP\.?O\.?\s*[:#]?\s*(\d+[-\d]*)', line, re.IGNORECASE)
                if po_match:
                    so_data['po_number'] = po_match.group(1)
                    print(f"  📋 Found PO: {so_data['po_number']}")
        
        # Create billing and shipping addresses for compatibility
        # First, CLEAN the raw address string by removing garbage (including PICK UP instructions)
        def clean_address_string(address_str):
            """Remove garbage from address string - phone numbers, batch info, N/A, etc."""
            if not address_str:
                return ''
            
            cleaned = address_str
            
            # STEP 0: Remove PICK UP instructions (these are notes, not addresses)
            cleaned = re.sub(r'\s*-\s*PICK\s*UP\s*,?\s*', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r'\s*PICK\s*UP\s*-\s*', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r'^\s*PICK\s*UP\s*,?\s*', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r'\s*\(\s*PICK\s*UP\s*\)\s*', '', cleaned, flags=re.IGNORECASE)
            
            # STEP 1: Remove everything AFTER certain keywords (they indicate end of address)
            # CRITICAL: Remove stock comments - these should NEVER be in addresses
            cleaned = re.sub(r',?\s*pull\s+from\s+stock.*$', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',?\s*pull\s+from.*$', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',?\s*batch\s*#?\s*\S+.*$', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',?\s*lot\s*#?\s*\S+.*$', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',?\s*stock.*$', '', cleaned, flags=re.IGNORECASE)
            
            # STEP 2: Remove phone numbers (various formats)
            cleaned = re.sub(r',?\s*\d{3}[-.\s]+\d{3}[-.\s]+\d{4}\b', '', cleaned)
            cleaned = re.sub(r',?\s*\(\d{3}\)\s*\d{3}[-.\s]?\d{4}', '', cleaned)
            cleaned = re.sub(r',?\s*\b\d{10}\b', '', cleaned)
            
            # STEP 3: Remove N/A values
            cleaned = re.sub(r',?\s*N/?A\s*N/?A', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',?\s*\bN/?A\b', '', cleaned, flags=re.IGNORECASE)
            
            # STEP 4: Remove Attn prefix (but keep the name)
            cleaned = re.sub(r'Attn:?\s*', '', cleaned, flags=re.IGNORECASE)
            
            # STEP 5: Handle country names - only keep at end, remove from middle
            cleaned = re.sub(r',\s*\bUSA\b\s*,', ',', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',\s*\bCanada\b\s*,', ',', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',\s*\bUSA\b\s*$', '', cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r',\s*\bCanada\b\s*$', '', cleaned, flags=re.IGNORECASE)
            
            # STEP 6: Final cleanup
            cleaned = re.sub(r',\s*,+', ',', cleaned)  # Multiple commas
            cleaned = re.sub(r'\s+', ' ', cleaned)  # Multiple spaces
            cleaned = re.sub(r',\s*$', '', cleaned)  # Trailing comma
            cleaned = re.sub(r'^\s*,', '', cleaned)  # Leading comma
            cleaned = cleaned.strip().strip(',').strip()
            
            return cleaned
        
        def parse_address(address_str, contact_person=''):
            """Parse address string into street, city, province, postal_code, country"""
            result = {'street': '', 'city': '', 'province': '', 'postal_code': '', 'country': 'Canada'}
            
            if not address_str:
                return result
            
            # First clean the address
            cleaned = clean_address_string(address_str)
            
            if not cleaned:
                return result
            
            # Province/State full names to 2-letter codes mapping
            province_name_to_code = {
                # Canada
                'alberta': 'AB', 'british columbia': 'BC', 'manitoba': 'MB', 
                'new brunswick': 'NB', 'newfoundland': 'NL', 'newfoundland and labrador': 'NL',
                'nova scotia': 'NS', 'northwest territories': 'NT', 'nunavut': 'NU',
                'ontario': 'ON', 'prince edward island': 'PE', 'quebec': 'QC', 
                'saskatchewan': 'SK', 'yukon': 'YT',
                # US States (common ones)
                'alabama': 'AL', 'alaska': 'AK', 'arizona': 'AZ', 'arkansas': 'AR',
                'california': 'CA', 'colorado': 'CO', 'connecticut': 'CT', 'delaware': 'DE',
                'florida': 'FL', 'georgia': 'GA', 'hawaii': 'HI', 'idaho': 'ID',
                'illinois': 'IL', 'indiana': 'IN', 'iowa': 'IA', 'kansas': 'KS',
                'kentucky': 'KY', 'louisiana': 'LA', 'maine': 'ME', 'maryland': 'MD',
                'massachusetts': 'MA', 'michigan': 'MI', 'minnesota': 'MN', 'mississippi': 'MS',
                'missouri': 'MO', 'montana': 'MT', 'nebraska': 'NE', 'nevada': 'NV',
                'new hampshire': 'NH', 'new jersey': 'NJ', 'new mexico': 'NM', 'new york': 'NY',
                'north carolina': 'NC', 'north dakota': 'ND', 'ohio': 'OH', 'oklahoma': 'OK',
                'oregon': 'OR', 'pennsylvania': 'PA', 'rhode island': 'RI', 'south carolina': 'SC',
                'south dakota': 'SD', 'tennessee': 'TN', 'texas': 'TX', 'utah': 'UT',
                'vermont': 'VT', 'virginia': 'VA', 'washington': 'WA', 'west virginia': 'WV',
                'wisconsin': 'WI', 'wyoming': 'WY'
            }
            
            valid_codes = ['AL','AK','AZ','AR','CA','CO','CT','DE','FL','GA','HI','ID','IL','IN','IA','KS','KY','LA','ME','MD','MA','MI','MN','MS','MO','MT','NE','NV','NH','NJ','NM','NY','NC','ND','OH','OK','OR','PA','RI','SC','SD','TN','TX','UT','VT','VA','WA','WV','WI','WY',  # US
                           'AB','BC','MB','NB','NL','NS','NT','NU','ON','PE','QC','SK','YT']  # Canada
            
            # Extract postal code first and use it to determine country
            # Canadian: A1A 1A1, US: 12345 or 12345-6789
            postal_match = re.search(r'([A-Z]\d[A-Z]\s?\d[A-Z]\d)', cleaned, re.IGNORECASE)
            if postal_match:
                result['postal_code'] = postal_match.group(1).upper()
                result['country'] = 'Canada'  # Canadian postal code format
                cleaned = cleaned.replace(postal_match.group(0), '').strip()
            else:
                # Try US zip
                us_postal_match = re.search(r'\b(\d{5})(-\d{4})?\b', cleaned)
                if us_postal_match:
                    result['postal_code'] = us_postal_match.group(0)
                    result['country'] = 'USA'  # US zip code format
                    cleaned = cleaned.replace(us_postal_match.group(0), '').strip()
            
            # Override country if explicitly mentioned in original string
            if 'USA' in address_str.upper():
                result['country'] = 'USA'
            elif 'CANADA' in address_str.upper():
                result['country'] = 'Canada'
            
            # FIRST: Check for full province/state names BEFORE splitting
            # This handles cases like "Rouyn-Noranda, Quebec" where Quebec is a full name
            province_found = False
            for full_name, code in province_name_to_code.items():
                # Match full province name as a standalone word (case-insensitive)
                pattern = r'\b' + re.escape(full_name) + r'\b'
                if re.search(pattern, cleaned, re.IGNORECASE):
                    result['province'] = code
                    # Remove the full province name from cleaned string
                    cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE).strip()
                    province_found = True
                    print(f"DEBUG ADDRESS: Found full province name '{full_name}' -> '{code}'")
                    break
            
            # Split by comma and clean up
            parts = [p.strip() for p in cleaned.split(',') if p.strip()]
            
            if not parts:
                return result
            
            # If province not found yet, look for 2-letter codes
            if not province_found:
                for i, part in enumerate(reversed(parts)):
                    idx = len(parts) - 1 - i
                    # Look for 2-letter state/province code
                    state_match = re.search(r'\b([A-Z]{2})\b', part.upper())
                    if state_match:
                        # Verify it looks like a state (not random letters)
                        potential_state = state_match.group(1)
                        if potential_state in valid_codes:
                            result['province'] = potential_state
                            # Remove state from part
                            parts[idx] = re.sub(r'\b' + potential_state + r'\b', '', part).strip()
                            break
            
            # Clean up empty parts
            parts = [p.strip() for p in parts if p.strip()]
            
            # Last part is usually city
            if parts:
                result['city'] = parts[-1]
                parts = parts[:-1]
            
            # Remaining parts are street
            if parts:
                # Remove contact person name if it's in the street
                street = ', '.join(parts)
                if contact_person and contact_person in street:
                    street = street.replace(contact_person, '').strip().strip(',').strip()
                result['street'] = street
            
            return result
        
        # Get contact person for filtering
        sold_to_contact = so_data['sold_to'].get('contact_person', '')
        ship_to_contact = so_data['ship_to'].get('contact_person', '')
        
        # CRITICAL FIX: Use pre-extracted addresses (from layout or table extraction)
        # These are more reliable for two-column PDFs
        pre_sold_to = pre_extracted.get('sold_to_raw', '').strip()
        pre_ship_to = pre_extracted.get('ship_to_raw', '').strip()
        pre_batch = pre_extracted.get('batch_number', '').strip()
        pre_mo = pre_extracted.get('mo_number', '').strip()
        
        # Store batch number in so_data
        if pre_batch:
            so_data['batch_number'] = pre_batch
            print(f"📦 Batch number from PDF: {pre_batch}")
        if pre_mo:
            so_data['mo_number'] = pre_mo
            print(f"   MO number from PDF: {pre_mo}")
        
        # Use pre-extracted addresses if available AND they look like valid addresses
        if pre_sold_to and len(pre_sold_to) > 10:
            if not so_data['sold_to']['address'] or len(pre_sold_to) > len(so_data['sold_to']['address']):
                print(f"📋 Using PRE-EXTRACTED Sold To address (more complete)")
                so_data['sold_to']['address'] = pre_sold_to.replace('\n', ', ')
                # Also set company name from first line if not already set
                if not so_data['sold_to']['company_name']:
                    first_line = pre_sold_to.split('\n')[0].strip()
                    so_data['sold_to']['company_name'] = first_line
                    so_data['customer_name'] = first_line
        
        if pre_ship_to and len(pre_ship_to) > 10:
            # CRITICAL: Use pre-extracted Ship To if it's DIFFERENT from Sold To
            if pre_ship_to.upper() != pre_sold_to.upper():
                print(f"📋 Using PRE-EXTRACTED Ship To address (different from Sold To)")
                so_data['ship_to']['address'] = pre_ship_to.replace('\n', ', ')
                # Also set company name from first line if not already set
                if not so_data['ship_to']['company_name']:
                    first_line = pre_ship_to.split('\n')[0].strip()
                    so_data['ship_to']['company_name'] = first_line
            elif not so_data['ship_to']['address']:
                so_data['ship_to']['address'] = pre_ship_to.replace('\n', ', ')
        
        # TRY GPT FIRST for smart address parsing (handles any format)
        # Fall back to regex only if GPT fails
        billing_addr = None
        shipping_addr = None
        
        # Try GPT for billing address
        if so_data['sold_to']['address']:
            billing_addr = parse_address_with_gpt(so_data['sold_to']['address'], "billing")
        
        # Try GPT for shipping address  
        if so_data['ship_to']['address']:
            shipping_addr = parse_address_with_gpt(so_data['ship_to']['address'], "shipping")
        
        # Fall back to regex parser if GPT failed
        if not billing_addr:
            print("📝 Using regex fallback for billing address")
            billing_addr = parse_address(so_data['sold_to']['address'], sold_to_contact)
        if not shipping_addr:
            print("📝 Using regex fallback for shipping address")
            shipping_addr = parse_address(so_data['ship_to']['address'], ship_to_contact)
        
        # Clean the raw addresses before storing
        cleaned_billing_address = clean_address_string(so_data['sold_to']['address'])
        cleaned_shipping_address = clean_address_string(so_data['ship_to']['address'])
        
        # Use detected country if available, otherwise use parsed country
        final_country = so_data.get('detected_country') or billing_addr['country']
        
        so_data['billing_address'] = {
            'company': so_data['sold_to']['company_name'],
            'contact_person': so_data['sold_to']['contact_person'],
            'contact': so_data['sold_to']['contact_person'],  # Alias for frontend
            'full_address': so_data['sold_to']['address'],  # Full original address
            'address': cleaned_billing_address,  # CLEANED address
            'street': billing_addr['street'] or cleaned_billing_address,  # Use cleaned if parsing fails
            'street_address': billing_addr['street'] or cleaned_billing_address,  # Alias
            'city': billing_addr['city'],
            'province': billing_addr['province'],
            'postal': billing_addr['postal_code'],  # Frontend uses 'postal'
            'postal_code': billing_addr['postal_code'],
            'country': final_country,
            'phone': so_data['sold_to']['phone'],
            'email': so_data['sold_to']['email']
        }
        
        # Handle PICKUP orders - use billing address but add PICK UP instruction
        if so_data.get('is_pickup_order'):
            # For pickup orders, use billing address as base but mark as pickup
            so_data['shipping_address'] = {
                'company': so_data['ship_to']['company_name'],
                'contact_person': so_data['ship_to']['contact_person'] or billing_addr.get('contact_person', ''),
                'contact': so_data['ship_to']['contact_person'] or billing_addr.get('contact', ''),
                'full_address': cleaned_billing_address,  # Use billing address
                'address': cleaned_billing_address,  # Use billing address
                'street': billing_addr['street'] or cleaned_billing_address,
                'street_address': billing_addr['street'] or cleaned_billing_address,
                'city': billing_addr['city'],
                'province': billing_addr['province'],
                'postal': billing_addr['postal_code'],
                'postal_code': billing_addr['postal_code'],
                'country': final_country,
                'phone': so_data['ship_to']['phone'] or so_data['sold_to']['phone'],
                'email': so_data['ship_to']['email'] or so_data['sold_to']['email'],
                'is_pickup': True,
                'pickup_instruction': 'PICK UP'  # This will be shown as a note
            }
            print(f"📦 Ship To marked as PICKUP order - will show instruction note")
        else:
            so_data['shipping_address'] = {
                'company': so_data['ship_to']['company_name'],
                'contact_person': so_data['ship_to']['contact_person'],
                'contact': so_data['ship_to']['contact_person'],  # Alias for frontend
                'full_address': so_data['ship_to']['address'],  # Full original address
                'address': cleaned_shipping_address,  # CLEANED address
                'street': shipping_addr['street'] or cleaned_shipping_address,  # Use cleaned if parsing fails
                'street_address': shipping_addr['street'] or cleaned_shipping_address,  # Alias
                'city': shipping_addr['city'],
                'province': shipping_addr['province'],
                'postal': shipping_addr['postal_code'],  # Frontend uses 'postal'
                'postal_code': shipping_addr['postal_code'],
                'country': so_data.get('detected_country') or shipping_addr['country'],
                'phone': so_data['ship_to']['phone'],
                'email': so_data['ship_to']['email'],
                'is_pickup': False
            }
        
        # Set status
        filename = os.path.basename(pdf_path)
        if 'R1' in filename or 'R2' in filename or 'R3' in filename:
            so_data['status'] = 'Revised'
        elif 'cancelled' in filename.lower():
            so_data['status'] = 'Cancelled'
        elif 'completed' in filename.lower():
            so_data['status'] = 'Completed'
        else:
            so_data['status'] = 'Unknown'
        
        # FALLBACK: Calculate subtotal from items if not extracted from PDF
        if so_data['subtotal'] == 0 and so_data['items']:
            calculated_subtotal = sum(item.get('amount', 0) or item.get('total_price', 0) or 0 for item in so_data['items'])
            if calculated_subtotal > 0:
                so_data['subtotal'] = calculated_subtotal
                print(f"  📊 Calculated subtotal from items: ${calculated_subtotal:.2f}")
        
        # FALLBACK: Calculate total_amount if not extracted from PDF
        if so_data['total_amount'] == 0:
            # Total = subtotal + tax
            if so_data['subtotal'] > 0:
                so_data['total_amount'] = so_data['subtotal'] + so_data['tax']
                print(f"  📊 Calculated total from subtotal + tax: ${so_data['total_amount']:.2f}")
        
        print(f"  SO Number: {so_data['so_number']}")
        print(f"  Customer: {so_data['customer_name']}")
        print(f"  Sold To: {so_data['sold_to']['company_name']}")
        print(f"  Ship To: {so_data['ship_to']['company_name']} - {so_data['ship_to']['contact_person']}")
        print(f"  Items: {len(so_data['items'])}")
        print(f"  Subtotal: ${so_data['subtotal']:.2f}")
        print(f"  Tax: ${so_data['tax']:.2f}")
        print(f"  Total: ${so_data['total_amount']:.2f}")
        
        # DEBUG: Print final addresses for verification
        print(f"\n📍 FINAL PARSED ADDRESSES:")
        print(f"  BILLING:")
        print(f"    Company: {so_data['billing_address'].get('company', 'N/A')}")
        print(f"    Full Address: {so_data['billing_address'].get('full_address', 'N/A')[:80]}...")
        print(f"    Street: {so_data['billing_address'].get('street', 'N/A')}")
        print(f"    City: {so_data['billing_address'].get('city', 'N/A')}, Province: {so_data['billing_address'].get('province', 'N/A')}")
        print(f"    Postal: {so_data['billing_address'].get('postal_code', 'N/A')}, Country: {so_data['billing_address'].get('country', 'N/A')}")
        print(f"  SHIPPING:")
        print(f"    Company: {so_data['shipping_address'].get('company', 'N/A')}")
        print(f"    Full Address: {so_data['shipping_address'].get('full_address', 'N/A')[:80]}...")
        print(f"    Street: {so_data['shipping_address'].get('street', 'N/A')}")
        print(f"    City: {so_data['shipping_address'].get('city', 'N/A')}, Province: {so_data['shipping_address'].get('province', 'N/A')}")
        print(f"    Postal: {so_data['shipping_address'].get('postal_code', 'N/A')}, Country: {so_data['shipping_address'].get('country', 'N/A')}")
        
        # SAFETY CHECK: Warn if billing and shipping addresses are identical (might be parsing error)
        billing_full = so_data['billing_address'].get('full_address', '')
        shipping_full = so_data['shipping_address'].get('full_address', '')
        if billing_full and shipping_full and billing_full.strip() == shipping_full.strip():
            # Check if table extraction showed different addresses
            if table_sold_to and table_ship_to and table_sold_to != table_ship_to:
                print(f"⚠️ WARNING: Billing and shipping addresses are IDENTICAL but table extraction showed DIFFERENT values!")
                print(f"   Table Sold To: {table_sold_to[:60]}...")
                print(f"   Table Ship To: {table_ship_to[:60]}...")
        
        return so_data
        
    except Exception as e:
        print(f"Error extracting SO data from {pdf_path}: {str(e)}")
        return None

# Add missing function alias for logistics module compatibility
def parse_sales_order_pdf(pdf_path):
    """Alias for extract_so_data_from_pdf - maintains compatibility with logistics module"""
    return extract_so_data_from_pdf(pdf_path)

# OLD PARSER - DEPRECATED (2025-01-03) - Use extract_so_data_from_pdf instead
def extract_so_data_from_pdf_old(pdf_path):
    """OLD SO PARSER - DEPRECATED 2025-01-03 - Use extract_so_data_from_pdf instead"""
    return extract_so_data_from_pdf(pdf_path)

def load_real_so_data():
    """Load real Sales Order data from ALL folders and subfolders recursively - OPTIMIZED"""
    so_data = []
    base_directory = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
    
    print(f"SEARCH: OPTIMIZED SO SCAN: Starting efficient recursive scan from {base_directory}")
    
    # Get all PDF files recursively
    pdf_files = []
    for root, dirs, files in os.walk(base_directory):
        for file in files:
            if file.lower().endswith('.pdf') and 'salesorder' in file.lower():
                pdf_files.append(os.path.join(root, file))
    
    print(f"SEARCH: Found {len(pdf_files)} SO PDF files")
    
    # Process files in batches for better performance
    batch_size = 50
    for i in range(0, len(pdf_files), batch_size):
        batch = pdf_files[i:i + batch_size]
        print(f"SEARCH: Processing batch {i//batch_size + 1}/{(len(pdf_files) + batch_size - 1)//batch_size}")
        
        for pdf_path in batch:
            try:
                so_data_item = extract_so_data_from_pdf(pdf_path)
                if so_data_item:
                    so_data.append(so_data_item)
            except Exception as e:
                print(f"ERROR: Failed to process {pdf_path}: {str(e)}")
                continue
    
    print(f"SUCCESS: Loaded {len(so_data)} Sales Orders")
    return so_data

def load_real_so_data():
    """Load real Sales Order data from ALL folders and subfolders recursively - OPTIMIZED"""
    so_data = []
    base_directory = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
    
    print(f"SEARCH: OPTIMIZED SO SCAN: Starting efficient recursive scan from {base_directory}")
    
    if not os.path.exists(base_directory):
        print(f"ERROR: Base SO directory not found: {base_directory}")
        return so_data
    
    # Get all PDF files recursively
    pdf_files = []
    for root, dirs, files in os.walk(base_directory):
        for file in files:
            if file.lower().endswith('.pdf') and 'salesorder' in file.lower():
                pdf_files.append(os.path.join(root, file))
    
    print(f"SEARCH: Found {len(pdf_files)} SO PDF files")
    
    # Process files in batches for better performance
    batch_size = 50
    for i in range(0, len(pdf_files), batch_size):
        batch = pdf_files[i:i + batch_size]
        print(f"SEARCH: Processing batch {i//batch_size + 1}/{(len(pdf_files) + batch_size - 1)//batch_size}")
        
        for pdf_path in batch:
            try:
                so_data_item = extract_so_data_from_pdf(pdf_path)
                if so_data_item:
                    so_data.append(so_data_item)
            except Exception as e:
                print(f"ERROR: Failed to process {pdf_path}: {str(e)}")
                continue
    
    print(f"SUCCESS: Loaded {len(so_data)} Sales Orders")
    return so_data

def load_real_so_data():
    """Load real Sales Order data from ALL folders and subfolders recursively - OPTIMIZED"""
    so_data = []
    base_directory = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
    
    print(f"SEARCH: OPTIMIZED SO SCAN: Starting efficient recursive scan from {base_directory}")
    
    if not os.path.exists(base_directory):
        print(f"ERROR: Base SO directory not found: {base_directory}")
        return so_data
    
    # Get all PDF files recursively
    pdf_files = []
    for root, dirs, files in os.walk(base_directory):
        for file in files:
            if file.lower().endswith('.pdf') and 'salesorder' in file.lower():
                pdf_files.append(os.path.join(root, file))
    
    print(f"SEARCH: Found {len(pdf_files)} SO PDF files")
    
    # Process files in batches for better performance
    batch_size = 50
    for i in range(0, len(pdf_files), batch_size):
        batch = pdf_files[i:i + batch_size]
        print(f"SEARCH: Processing batch {i//batch_size + 1}/{(len(pdf_files) + batch_size - 1)//batch_size}")
        
        for pdf_path in batch:
            try:
                so_data_item = extract_so_data_from_pdf(pdf_path)
                if so_data_item:
                    so_data.append(so_data_item)
            except Exception as e:
                print(f"ERROR: Failed to process {pdf_path}: {str(e)}")
                continue
    
    print(f"SUCCESS: Loaded {len(so_data)} Sales Orders")
    return so_data

def load_real_so_data():
    """Load real Sales Order data from ALL folders and subfolders recursively - OPTIMIZED"""
    so_data = []
    base_directory = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
    
    print(f"SEARCH: OPTIMIZED SO SCAN: Starting efficient recursive scan from {base_directory}")
    
    if not os.path.exists(base_directory):
        print(f"ERROR: Base SO directory not found: {base_directory}")
        return so_data
    
    # Get all PDF files recursively
    pdf_files = []
    for root, dirs, files in os.walk(base_directory):
        for file in files:
            if file.lower().endswith('.pdf') and 'salesorder' in file.lower():
                pdf_files.append(os.path.join(root, file))
    
    print(f"SEARCH: Found {len(pdf_files)} SO PDF files")
    
    # Process files in batches for better performance
    batch_size = 50
    for i in range(0, len(pdf_files), batch_size):
        batch = pdf_files[i:i + batch_size]
        print(f"SEARCH: Processing batch {i//batch_size + 1}/{(len(pdf_files) + batch_size - 1)//batch_size}")
        
        for pdf_path in batch:
            try:
                so_data_item = extract_so_data_from_pdf(pdf_path)
                if so_data_item:
                    so_data.append(so_data_item)
            except Exception as e:
                print(f"ERROR: Failed to process {pdf_path}: {str(e)}")
                continue
    
    print(f"SUCCESS: Loaded {len(so_data)} Sales Orders")
    return so_data

# G: Drive base paths - EXACT paths where data is located
GDRIVE_BASE = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\API Extractions"
SALES_ORDERS_BASE = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
MPS_BASE = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\API Extractions"
MPS_EXCEL_PATH = os.path.join(MPS_BASE, "MPS.xlsx")  # Fallback Excel file

# Full Company Data (MISys "Export All Company Data") - for import / direct fetch
# Local path (when backend runs where G: is mounted):
# Local path when Google Drive for Desktop is mounted as G: (dev/office only). On Render/cloud this is not used — we use Google Drive API and FULL_COMPANY_DATA_DRIVE_PATH instead.
GDRIVE_FULL_COMPANY_DATA = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\Full Company Data as of 02_10_2026"
# Drive-relative path for Google Drive API (Render/cloud); override via env if needed
FULL_COMPANY_DATA_DRIVE_PATH = os.getenv(
    "FULL_COMPANY_DATA_DRIVE_PATH",
    "MiSys/Misys Extracted Data/Full Company Data as of 02_10_2026"
)

def get_latest_folder():
    """Get the latest folder from G: Drive OR Google Drive API - OPTIMIZED for speed"""
    try:
        # CRITICAL: On cloud environments (Render/Cloud Run), NEVER try local G: Drive paths
        if IS_CLOUD_ENVIRONMENT:
            print("☁️ Cloud environment detected - using Google Drive API only")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
                else:
                    return None, "Could not find latest folder via Google Drive API"
            else:
                return None, "Google Drive API service not available on cloud environment"
        
        # LOCAL DEVELOPMENT: Try Google Drive API first if enabled, then fallback to local G: Drive
        if USE_GOOGLE_DRIVE_API:
            print("SEARCH: Using Google Drive API to find latest folder...")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
            # If API fails locally, fall through to try local G: Drive
        
        # LOCAL ONLY: Try local G: Drive path
        print(f"SEARCH: Checking local G: Drive path: {GDRIVE_BASE}")
        
        if not os.path.exists(GDRIVE_BASE):
            print(f"ERROR: G: Drive path not accessible: {GDRIVE_BASE}")
            # Try to initialize Google Drive API as fallback
            print("RETRY: Attempting to use Google Drive API as fallback...")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API fallback): {latest_name}")
                    return latest_name, None
            
            return None, f"G: Drive path not accessible and API fallback failed: {GDRIVE_BASE}"
        
        # FAST: Get folders and sort by name (assuming date format YYYY-MM-DD)
        folders = [f for f in os.listdir(GDRIVE_BASE) if os.path.isdir(os.path.join(GDRIVE_BASE, f))]
        print(f"📁 Found folders: {folders}")
        
        if not folders:
            print("ERROR: No folders found in G: Drive path")
            return None, "No folders found in G: Drive path"
        
        # OPTIMIZED: Sort by folder name (assuming YYYY-MM-DD format) instead of checking mtime
        folders.sort(reverse=True)  # Most recent first
        latest_folder = folders[0]
        
        print(f"SUCCESS: Latest folder found: {latest_folder}")
        return latest_folder, None
            
    except Exception as e:
        print(f"ERROR: Error in get_latest_folder: {e}")
        return None, str(e)

def load_json_file(file_path):
    """Load JSON file from G: Drive with proper error handling"""
    try:
        if os.path.exists(file_path):
            print(f"📄 Loading file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                print(f"SUCCESS: Successfully loaded {len(data)} records from {os.path.basename(file_path)}")
                return data
        else:
            print(f"⚠️ File not found: {file_path}")
        return []
    except Exception as e:
        print(f"ERROR: Error loading {file_path}: {e}")
        return []

# Global cache variables
_data_cache = None
_cache_timestamp = None
_cache_duration = 3600  # 1 hour cache (was 5 minutes - too short, causing frequent reloads)

app = Flask(__name__)

# MPS Sheet Configuration
MPS_SHEET_ID = '1zAOY7ngP2mLVi-W_FL9tsPiKDPqbU6WEUmrrTDeKygw'
MPS_CSV_URL = f'https://docs.google.com/spreadsheets/d/{MPS_SHEET_ID}/export?format=csv'

# CORS Configuration - Allow all origins for Cloud Run
# Cloud Run uses different URLs for each deployment
CORS(app, resources={
    r"/api/*": {
        "origins": "*",
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization", "Accept", "Origin", "X-Requested-With"],
        "expose_headers": ["Content-Type", "Content-Length"],
        "supports_credentials": False,
        "max_age": 3600
    }
})

# Enable GZIP compression - reduces large payloads (e.g. 60–70MB Full Company Data) to ~10MB over the wire
Compress(app)
print("GZIP compression enabled")

# Register logistics automation blueprint
if LOGISTICS_AVAILABLE:
    app.register_blueprint(logistics_bp)
    print("Logistics automation module loaded")
else:
    print("Logistics automation module not available")

# Register purchase requisition service blueprint
if PR_SERVICE_AVAILABLE:
    app.register_blueprint(pr_service)
    print("Purchase Requisition service loaded")
else:
    print("Purchase Requisition service not available")

# Register BOL HTML blueprint - DISABLED (duplicate endpoint in logistics_automation.py)
# if BOL_HTML_AVAILABLE:
#     app.register_blueprint(bol_html_bp)
#     print("BOL HTML module loaded")
# else:
#     print("BOL HTML module not available")

# Initialize OpenAI client
try:
    from openai import OpenAI
    # Get API key from environment variable - NEVER hardcode API keys!
    openai_api_key = os.getenv('OPENAI_API_KEY')
    
    if not openai_api_key or openai_api_key == "your_openai_api_key_here":
        print("ERROR: ERROR: OPENAI_API_KEY environment variable not set or using placeholder")
        print("💡 Set it with: set OPENAI_API_KEY=sk-proj-your_actual_key_here")
        print("💡 Get your API key from: https://platform.openai.com/api-keys")
        client = None
        openai_available = False
    else:
        try:
            client = OpenAI(api_key=openai_api_key)
            # Test the API key with a simple request
            openai_available = True
            print("OpenAI client initialized successfully")
        except Exception as e:
            print(f"ERROR: OpenAI client initialization failed: {e}")
            client = None
            openai_available = False
except ImportError:
    print("ERROR: OpenAI library not found. Install with: pip install openai")
    client = None
    openai_available = False
except Exception as e:
    print(f"ERROR: OpenAI initialization failed: {e}")
    client = None
    openai_available = False

def parse_address_with_gpt(raw_address_string, address_type="billing"):
    """
    Use GPT to intelligently parse ANY address format into structured components.
    This handles international addresses, various formats, and edge cases that regex can't.
    """
    global client, openai_available
    
    if not openai_available or not client or not raw_address_string:
        return None  # Fall back to regex parser
    
    try:
        prompt = f"""Parse this {address_type} address into structured components. Handle ANY format.

RAW ADDRESS:
{raw_address_string}

Return a JSON object with these fields (use empty string if not found, NEVER use "N/A"):
{{
    "street": "street address with building number",
    "city": "city name only",
    "province": "2-letter province/state code (e.g., QC for Quebec, ON for Ontario, TX for Texas)",
    "postal_code": "postal/zip code",
    "country": "Canada or USA or other country name"
}}

RULES:
1. "Quebec" → province: "QC", "Ontario" → province: "ON", "Texas" → province: "TX", etc.
2. Canadian postal codes look like "A1A 1A1", US zip codes are "12345" or "12345-6789"
3. If province is written as full name (e.g., "Quebec", "British Columbia"), convert to 2-letter code
4. City is usually before the province/state
5. For addresses like "Rouyn-Noranda, Quebec J9X 5B5":
   - city: "Rouyn-Noranda"
   - province: "QC"
   - postal_code: "J9X 5B5"

Return ONLY the JSON, no explanations."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert address parser. Parse addresses with 100% accuracy. Always convert province/state names to 2-letter codes. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=300,
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content.strip()
        parsed = json.loads(result)
        
        print(f"🤖 GPT Address Parse ({address_type}): {raw_address_string[:50]}... → city='{parsed.get('city')}', province='{parsed.get('province')}', postal='{parsed.get('postal_code')}'")
        
        return {
            'street': parsed.get('street', ''),
            'city': parsed.get('city', ''),
            'province': parsed.get('province', ''),
            'postal_code': parsed.get('postal_code', ''),
            'country': parsed.get('country', 'Canada')
        }
        
    except Exception as e:
        print(f"⚠️ GPT address parsing failed, using fallback: {e}")
        return None  # Fall back to regex parser

# Google Drive API - LAZY INITIALIZATION (only when needed, not on startup)
# This prevents startup failures if Google Drive API has issues
# CRITICAL: On Cloud Run/Render, ALWAYS use Google Drive API (G: Drive is NEVER accessible)
IS_CLOUD_RUN = os.getenv('K_SERVICE') is not None
IS_RENDER = os.getenv('RENDER') is not None or os.getenv('RENDER_SERVICE_ID') is not None
IS_CLOUD_ENVIRONMENT = IS_CLOUD_RUN or IS_RENDER
SKIP_GOOGLE_DRIVE = os.getenv('SKIP_GOOGLE_DRIVE', '').lower() in ('1', 'true', 'yes')
USE_GOOGLE_DRIVE_API = (IS_CLOUD_ENVIRONMENT or os.getenv('USE_GOOGLE_DRIVE_API', 'false').lower() == 'true') and not SKIP_GOOGLE_DRIVE
google_drive_service = None

if IS_CLOUD_ENVIRONMENT:
    if SKIP_GOOGLE_DRIVE:
        print("Backend: Cloud — SKIP_GOOGLE_DRIVE=1, not using Google Drive (return empty data)")
    elif IS_RENDER:
        print("Backend: Render/cloud — data from Google Drive API only (no local drive paths)")
    else:
        print("Backend: Cloud Run — data from Google Drive API only (no local drive paths)")
else:
    print("Backend: Local/dev — data from G: path when available, else Google Drive API")

def get_google_drive_service():
    """Lazy initialization of Google Drive service - only when actually needed"""
    global google_drive_service
    
    if not USE_GOOGLE_DRIVE_API:
        return None
    
    if google_drive_service is not None:
        return google_drive_service
    
    # Initialize only when first needed
    try:
        print("Initializing Google Drive API service (lazy load)...")
        from google_drive_service import GoogleDriveService
        google_drive_service = GoogleDriveService()
        if google_drive_service.authenticate():
            print("Google Drive API service initialized successfully")
            return google_drive_service
        else:
            print("WARNING: Google Drive API authentication failed - will fall back to G: Drive if available")
            google_drive_service = None
            return None
    except Exception as e:
        print(f"WARNING: Failed to initialize Google Drive service: {e}")
        import traceback
        traceback.print_exc()
        google_drive_service = None
        return None

def get_latest_folder():
    """Get the latest folder from G: Drive OR Google Drive API - OPTIMIZED for speed"""
    try:
        # CRITICAL: On cloud environments (Render/Cloud Run), NEVER try local G: Drive paths
        if IS_CLOUD_ENVIRONMENT:
            print("☁️ Cloud environment detected - using Google Drive API only")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
                else:
                    return None, "Could not find latest folder via Google Drive API"
            else:
                return None, "Google Drive API service not available on cloud environment"
        
        # LOCAL DEVELOPMENT: Try Google Drive API first if enabled, then fallback to local G: Drive
        if USE_GOOGLE_DRIVE_API:
            print("SEARCH: Using Google Drive API to find latest folder...")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
            # If API fails locally, fall through to try local G: Drive
        
        # LOCAL ONLY: Try local G: Drive path
        print(f"SEARCH: Checking local G: Drive path: {GDRIVE_BASE}")
        
        if not os.path.exists(GDRIVE_BASE):
            print(f"ERROR: G: Drive path not accessible: {GDRIVE_BASE}")
            # Try to initialize Google Drive API as fallback (local dev only)
            if not IS_CLOUD_ENVIRONMENT:
                print("RETRY: Attempting to use Google Drive API as fallback...")
                service = get_google_drive_service()
                if service:
                    latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                    if latest_name:
                        print(f"SUCCESS: Latest folder found (via API fallback): {latest_name}")
                        return latest_name, None
            
            return None, f"G: Drive path not accessible: {GDRIVE_BASE}"
        
        # FAST: Get folders and sort by name (assuming date format YYYY-MM-DD)
        folders = [f for f in os.listdir(GDRIVE_BASE) if os.path.isdir(os.path.join(GDRIVE_BASE, f))]
        print(f"📁 Found folders: {folders}")
        
        if not folders:
            print("ERROR: No folders found in G: Drive path")
            return None, "No folders found in G: Drive path"
        
        # OPTIMIZED: Sort by folder name (assuming YYYY-MM-DD format) instead of checking mtime
        folders.sort(reverse=True)  # Most recent first
        latest_folder = folders[0]
        
        print(f"SUCCESS: Latest folder found: {latest_folder}")
        return latest_folder, None
            
    except Exception as e:
        print(f"ERROR: Error in get_latest_folder: {e}")
        return None, str(e)

if USE_GOOGLE_DRIVE_API:
    print("Google Drive API enabled - will initialize when needed")
else:
    print("Google Drive API not enabled (USE_GOOGLE_DRIVE_API=false or not set)")

# G: Drive base paths - EXACT paths where data is located
GDRIVE_BASE = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\API Extractions"
SALES_ORDERS_BASE = r"G:\Shared drives\Sales_CSR\Customer Orders\Sales Orders"
MPS_BASE = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\API Extractions"
MPS_EXCEL_PATH = os.path.join(MPS_BASE, "MPS.xlsx")  # Fallback Excel file

def get_latest_folder():
    """Get the latest folder from G: Drive OR Google Drive API - OPTIMIZED for speed"""
    try:
        # CRITICAL: On cloud environments (Render/Cloud Run), NEVER try local G: Drive paths
        if IS_CLOUD_ENVIRONMENT:
            print("☁️ Cloud environment detected - using Google Drive API only")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
                else:
                    return None, "Could not find latest folder via Google Drive API"
            else:
                return None, "Google Drive API service not available on cloud environment"
        
        # LOCAL DEVELOPMENT: Try Google Drive API first if enabled, then fallback to local G: Drive
        if USE_GOOGLE_DRIVE_API:
            print("SEARCH: Using Google Drive API to find latest folder...")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API): {latest_name}")
                    return latest_name, None
            # If API fails locally, fall through to try local G: Drive
        
        # LOCAL ONLY: Try local G: Drive path
        print(f"SEARCH: Checking local G: Drive path: {GDRIVE_BASE}")
        
        if not os.path.exists(GDRIVE_BASE):
            print(f"ERROR: G: Drive path not accessible: {GDRIVE_BASE}")
            # Try to initialize Google Drive API as fallback
            print("RETRY: Attempting to use Google Drive API as fallback...")
            service = get_google_drive_service()
            if service:
                latest_id, latest_name, _ = service.find_latest_api_extractions_folder()
                if latest_name:
                    print(f"SUCCESS: Latest folder found (via API fallback): {latest_name}")
                    return latest_name, None
            
            return None, f"G: Drive path not accessible and API fallback failed: {GDRIVE_BASE}"
        
        # FAST: Get folders and sort by name (assuming date format YYYY-MM-DD)
        folders = [f for f in os.listdir(GDRIVE_BASE) if os.path.isdir(os.path.join(GDRIVE_BASE, f))]
        print(f"📁 Found folders: {folders}")
        
        if not folders:
            print("ERROR: No folders found in G: Drive path")
            return None, "No folders found in G: Drive path"
        
        # OPTIMIZED: Sort by folder name (assuming YYYY-MM-DD format) instead of checking mtime
        folders.sort(reverse=True)  # Most recent first
        latest_folder = folders[0]
        
        print(f"SUCCESS: Latest folder found: {latest_folder}")
        return latest_folder, None
            
    except Exception as e:
        print(f"ERROR: Error in get_latest_folder: {e}")
        return None, str(e)

def load_json_file(file_path):
    """Load JSON file from G: Drive with proper error handling - optimized for speed"""
    try:
        if os.path.exists(file_path):
            # Reduced logging for performance - only log errors
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data
        else:
            # File not found - return empty array silently (expected for optional files)
            return []
    except Exception as e:
        print(f"ERROR: Error loading {file_path}: {e}")
        return []



# Global cache for data - store PRE-SERIALIZED JSON to avoid 35s serialization time
_data_cache = None  # Stores raw data dict (for internal use)
_response_cache = None  # Stores pre-serialized JSON bytes
_cache_timestamp = None
_cache_duration = 3600  # 1 hour cache (was 5 minutes - too short)
# Portal store: persisted MOs, inventory adjustments, transfers, item overrides, BOM edits (see portal_store.py)
try:
    import portal_store
except ImportError:
    portal_store = None

def _merge_portal_store(data):
    """Apply portal store (created MOs, adjustments, transfers, overrides, BOM edits) to data."""
    if not data or portal_store is None:
        return data
    return portal_store.apply_to_data(data)

def get_empty_app_data_structure():
    """Return the exact app data shape the frontend expects (all keys, empty values). Framework for Full Company Data - app is ready to receive once conversion is implemented."""
    return {
        'CustomAlert5.json': [], 'Items.json': [], 'MIITEM.json': [], 'MIILOC.json': [],
        'BillsOfMaterial.json': [], 'BillOfMaterialDetails.json': [], 'MIBOMH.json': [], 'MIBOMD.json': [],
        'ManufacturingOrderHeaders.json': [], 'ManufacturingOrderDetails.json': [], 'ManufacturingOrderRoutings.json': [],
        'MIMOH.json': [], 'MIMOMD.json': [], 'MIMORD.json': [], 'Jobs.json': [], 'JobDetails.json': [],
        'MIJOBH.json': [], 'MIJOBD.json': [], 'MIPOH.json': [], 'MIPOD.json': [], 'MIPOHX.json': [],
        'MIPOC.json': [], 'MIPOCV.json': [], 'MIPODC.json': [], 'MIWOH.json': [], 'MIWOD.json': [], 'MIBORD.json': [],
        'PurchaseOrderDetails.json': [], 'PurchaseOrderExtensions.json': [], 'PurchaseOrders.json': [],
        'WorkOrderHeaders.json': [], 'WorkOrderDetails.json': [], 'WorkOrders.json': [], 'ParsedSalesOrders.json': [],
        'SalesOrderHeaders.json': [], 'SalesOrderDetails.json': [],
        'PurchaseOrderAdditionalCosts.json': [], 'PurchaseOrderAdditionalCostsTaxes.json': [], 'PurchaseOrderDetailAdditionalCosts.json': [],
        'SalesOrders.json': [], 'SalesOrdersByStatus': {}, 'TotalOrders': 0, 'StatusFolders': [], 'ScanMethod': '',
        'MPS.json': {'mps_orders': [], 'summary': {'total_orders': 0}},
        'LotSerialHistory.json': [], 'LotSerialDetail.json': [],
        'MIILOCQT.json': [], 'MIBINQ.json': [], 'MISLBINQ.json': [], 'MISLHIST.json': [],
        'MILOGH.json': [], 'MIBINH.json': [], 'MIICST.json': [], 'MIITEMX.json': [], 'MIITEMA.json': [],
        'MIQMFG.json': [], 'MISUPL.json': [],
    }

@app.route('/api/data', methods=['GET'])
def get_all_data():
    """Single source: MISys Full Company Data export (your CSV/Excel export folder). When that folder is available, we use it only. Use ?source=default to force legacy API Extractions instead."""
    global _data_cache, _response_cache, _cache_timestamp
    import time  # Import at function level for use throughout function
    from flask import request

    try:
        print("/api/data endpoint called")
        data_source_param = request.args.get('source')

        # Explicit Full Company Data requested
        if data_source_param == 'full_company_data':
            full_data = None
            err_msg = None
            try:
                from full_company_data_converter import load_from_folder, load_from_drive_api
            except ImportError:
                try:
                    from .full_company_data_converter import load_from_folder, load_from_drive_api
                except ImportError:
                    load_from_folder = load_from_drive_api = None
            if load_from_folder and not IS_CLOUD_ENVIRONMENT and os.path.exists(GDRIVE_FULL_COMPANY_DATA):
                full_data, err_msg = load_from_folder(GDRIVE_FULL_COMPANY_DATA)
            if full_data is None and load_from_drive_api:
                gdrive_service = get_google_drive_service()
                drive_id = gdrive_service.find_shared_drive("IT_Automation") if gdrive_service else None
                if drive_id and gdrive_service and getattr(gdrive_service, "authenticated", False):
                    full_data, err_msg = load_from_drive_api(gdrive_service, drive_id, FULL_COMPANY_DATA_DRIVE_PATH)
            if full_data is not None:
                full_data = _merge_portal_store(full_data)
                file_count = sum(1 for v in full_data.values() if isinstance(v, list) and len(v) > 0)
                return jsonify({
                    "data": full_data,
                    "folderInfo": {
                        "folderName": "Full Company Data as of 02_10_2026",
                        "syncDate": datetime.now().isoformat(),
                        "lastModified": datetime.now().isoformat(),
                        "folder": FULL_COMPANY_DATA_DRIVE_PATH,
                        "created": datetime.now().isoformat(),
                        "size": "N/A",
                        "fileCount": file_count,
                    },
                    "LoadTimestamp": datetime.now().isoformat(),
                    "source": "full_company_data",
                    "fullCompanyDataReady": True,
                })
            empty_data = get_empty_app_data_structure()
            return jsonify({
                "data": empty_data,
                "folderInfo": {
                    "folderName": "Full Company Data (framework)",
                    "syncDate": datetime.now().isoformat(),
                    "lastModified": datetime.now().isoformat(),
                    "folder": FULL_COMPANY_DATA_DRIVE_PATH,
                    "created": datetime.now().isoformat(),
                    "size": "0",
                    "fileCount": 0,
                },
                "LoadTimestamp": datetime.now().isoformat(),
                "source": "full_company_data (conversion failed or folder not available)",
                "fullCompanyDataReady": False,
                "message": err_msg or "Full Company Data folder not found or converter error",
            })
        
        # Default load (no ?source or any other value): prefer Full Company Data when available so MISys export "just works"
        if data_source_param != 'default':
            # When SKIP_GOOGLE_DRIVE: skip all G Drive attempts, return empty immediately
            if SKIP_GOOGLE_DRIVE:
                print("📂 SKIP_GOOGLE_DRIVE=1: skipping Google Drive and G: path, returning empty data")
                empty_data = get_empty_app_data_structure()
                return jsonify({
                    "data": empty_data,
                    "folderInfo": {
                        "folderName": "SKIP_GOOGLE_DRIVE",
                        "syncDate": datetime.now().isoformat(),
                        "lastModified": datetime.now().isoformat(),
                        "folder": "Skipped",
                        "created": datetime.now().isoformat(),
                        "size": "0",
                        "fileCount": 0,
                    },
                    "LoadTimestamp": datetime.now().isoformat(),
                    "source": "empty (SKIP_GOOGLE_DRIVE=1)",
                    "fullCompanyDataReady": False,
                })
            full_data = None
            err_msg = None
            if IS_CLOUD_ENVIRONMENT:
                print("📂 Full Company Data: loading via Google Drive API (cloud — no local path)")
            else:
                print(f"📂 Full Company Data: trying local path first (exists={os.path.exists(GDRIVE_FULL_COMPANY_DATA)})")
            try:
                from full_company_data_converter import load_from_folder, load_from_drive_api
            except ImportError:
                try:
                    from .full_company_data_converter import load_from_folder, load_from_drive_api
                except ImportError:
                    load_from_folder = load_from_drive_api = None
            if load_from_folder and not IS_CLOUD_ENVIRONMENT and os.path.exists(GDRIVE_FULL_COMPANY_DATA):
                full_data, err_msg = load_from_folder(GDRIVE_FULL_COMPANY_DATA)
                print(f"📂 Full Company Data load_from_folder: full_data={'present' if full_data else None}, err_msg={err_msg!r}")
                if full_data:
                    cnt = sum(1 for v in full_data.values() if isinstance(v, list) and len(v) > 0)
                    print(f"📂 Full Company Data: {cnt} non-empty lists")
            else:
                if not load_from_folder:
                    print("📂 Full Company Data: skipped (converter import failed)")
                elif IS_CLOUD_ENVIRONMENT:
                    print("📂 Full Company Data: using Drive API (cloud deployment)")
                elif not os.path.exists(GDRIVE_FULL_COMPANY_DATA):
                    print(f"📂 Full Company Data: local folder not found, will try Drive API")
            if full_data is None and load_from_drive_api:
                gdrive_service = get_google_drive_service()
                drive_id = gdrive_service.find_shared_drive("IT_Automation") if gdrive_service else None
                if drive_id and gdrive_service and getattr(gdrive_service, "authenticated", False):
                    full_data, err_msg = load_from_drive_api(gdrive_service, drive_id, FULL_COMPANY_DATA_DRIVE_PATH)
                    print(f"📂 Full Company Data load_from_drive_api: full_data={'present' if full_data else None}, err_msg={err_msg!r}")
                else:
                    print("📂 Full Company Data: Drive API not available or not authenticated")
            if full_data is not None:
                has_any = any(isinstance(v, list) and len(v) > 0 for v in full_data.values())
                if has_any:
                    full_data = _merge_portal_store(full_data)
                    file_count = sum(1 for v in full_data.values() if isinstance(v, list) and len(v) > 0)
                    print(f"✅ Using Full Company Data (MISys export) – {file_count} files with data")
                    return jsonify({
                        "data": full_data,
                        "folderInfo": {
                            "folderName": "Full Company Data as of 02_10_2026",
                            "syncDate": datetime.now().isoformat(),
                            "lastModified": datetime.now().isoformat(),
                            "folder": FULL_COMPANY_DATA_DRIVE_PATH,
                            "created": datetime.now().isoformat(),
                            "size": "N/A",
                            "fileCount": file_count,
                        },
                        "LoadTimestamp": datetime.now().isoformat(),
                        "cached": False,
                        "source": "full_company_data",
                        "fullCompanyDataReady": True,
                    })
                else:
                    print("📂 Full Company Data: converter returned data but all lists empty (check file names: need MIITEM.csv, Item.csv, MIBOMD.csv, MIPOH.csv, etc.)")
            else:
                print("📂 Full Company Data: not loaded. Falling back to API Extractions.")
        
        # Check if we have valid cached response (pre-serialized)
        if _response_cache and _cache_timestamp:
            cache_age = time.time() - _cache_timestamp
            print(f"SEARCH: Cache check: age={cache_age:.1f}s, duration={_cache_duration}s, valid={cache_age < _cache_duration}")
            if cache_age < _cache_duration:
                print("SUCCESS: Returning PRE-SERIALIZED cached response (instant!)")
                # Return pre-serialized response directly - no jsonify() overhead
                from flask import Response
                return Response(_response_cache, mimetype='application/json')
        
        print("RETRY: Cache expired or missing, loading fresh data...")
        
        # Check if G: Drive is accessible
        if not os.path.exists(GDRIVE_BASE):
            print(f"ERROR: G: Drive not accessible at: {GDRIVE_BASE}")
            
            # Try to use Google Drive API as fallback (lazy initialization)
            gdrive_service = get_google_drive_service()
            if gdrive_service and gdrive_service.authenticated:
                print("🔄 G: Drive not accessible, falling back to Google Drive API...")
                try:
                    gdrive_data, gdrive_folder_info = gdrive_service.get_all_data()
                    if gdrive_data and gdrive_folder_info:
                        print(f"✅ Successfully loaded data from Google Drive API")
                        gdrive_data = _merge_portal_store(gdrive_data)
                        # Cache the data AND pre-serialize the response
                        import json as json_module
                        _data_cache = gdrive_data
                        _cache_timestamp = time.time()
                        
                        # Pre-serialize the response to avoid 35s jsonify overhead on cache hits
                        response_dict = {
                            "data": gdrive_data,
                            "folderInfo": gdrive_folder_info,
                            "LoadTimestamp": datetime.now().isoformat(),
                            "cached": False,
                            "source": "Google Drive API"
                        }
                        _response_cache = json_module.dumps(response_dict)
                        print(f"💾 Data + Response cached for {_cache_duration} seconds (pre-serialized: {len(_response_cache)/1024/1024:.1f}MB)")
                        
                        from flask import Response
                        return Response(_response_cache, mimetype='application/json')
                    else:
                        print("⚠️ Google Drive API returned empty data")
                except Exception as e:
                    print(f"❌ Error loading data from Google Drive API: {e}")
                    import traceback
                    traceback.print_exc()
            
            # If Google Drive API also failed or not available, return empty data structure
            print("⚠️ Both G: Drive and Google Drive API unavailable - returning empty data")
            empty_data = get_empty_app_data_structure()
            empty_data['ScanMethod'] = 'No G: Drive Access'
            return jsonify({
                "data": empty_data,
                "folderInfo": {
                    "folderName": "No G: Drive Access",
                    "syncDate": datetime.now().isoformat(),
                    "lastModified": datetime.now().isoformat(),
                    "folder": "Not Connected",
                    "created": datetime.now().isoformat(),
                    "size": "0",
                    "fileCount": 0
                },
                "LoadTimestamp": datetime.now().isoformat(),
                "warning": "G: Drive not accessible - returning empty data",
                "source": "None (G: Drive not accessible)"
            })
        
        latest_folder, error = get_latest_folder()
        if error:
            print(f"ERROR: Error getting latest folder: {error}")
            return jsonify({"error": error}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        print(f"📂 Loading data from folder: {folder_path}")
        
        # FAST LOADING - Only load essential files first for speed
        raw_data = {}
        
        # Essential files for proper data loading - 10 critical files
        essential_files = [
            "CustomAlert5.json",  # PRIMARY: Complete item data
            "MIILOC.json",        # Inventory location data
            "SalesOrderHeaders.json",  # Sales orders
            "SalesOrderDetails.json",  # Sales order details
            "ManufacturingOrderHeaders.json",  # Manufacturing orders
            "ManufacturingOrderDetails.json",  # Manufacturing order details
            "BillsOfMaterial.json",  # BOM data
            "BillOfMaterialDetails.json",  # BOM details
            "PurchaseOrders.json",  # Purchase orders
            "PurchaseOrderDetails.json"  # Purchase order details
        ]
        
        print(f"⚡ LOADING: Loading G: Drive data...")
        
        # Load only essential files first (optimized - minimal logging for speed)
        for file_name in essential_files:
            file_path = os.path.join(folder_path, file_name)
            file_data = load_json_file(file_path)
            raw_data[file_name] = file_data
            if file_data:
                print(f"✅ {file_name}: {len(file_data) if isinstance(file_data, list) else 1} records")
        
        # Optional files: load from folder when present (so item modal, BOM Where Used, Work Orders, PO costs get real data for testing)
        optional_files = [
            'Items.json', 'MIITEM.json', 'MIBOMH.json', 'MIBOMD.json',
            'ManufacturingOrderRoutings.json', 'MIMOH.json', 'MIMOMD.json', 'MIMORD.json',
            'Jobs.json', 'JobDetails.json', 'MIJOBH.json', 'MIJOBD.json',
            'MIPOH.json', 'MIPOD.json', 'MIPOHX.json', 'MIPOC.json', 'MIPOCV.json',
            'MIPODC.json', 'MIWOH.json', 'MIWOD.json', 'MIBORD.json',
            'PurchaseOrderExtensions.json', 'WorkOrders.json', 'WorkOrderDetails.json',
            'PurchaseOrderAdditionalCosts.json', 'PurchaseOrderAdditionalCostsTaxes.json',
            'PurchaseOrderDetailAdditionalCosts.json'
        ]
        for file_name in optional_files:
            file_path = os.path.join(folder_path, file_name)
            file_data = load_json_file(file_path)
            raw_data[file_name] = file_data if isinstance(file_data, list) else ([] if file_data is None else [])
            if file_data and isinstance(file_data, list) and len(file_data) > 0:
                print(f"✅ {file_name}: {len(file_data)} records")
        
        loaded_count = len([k for k, v in raw_data.items() if isinstance(v, list) and len(v) > 0])
        print(f"⚡ Loaded {loaded_count} files with data")
        
        # Use data AS-IS - no conversion needed!
        
        # Get folder info
        folder_info = {
            "folderName": latest_folder,
            "syncDate": datetime.fromtimestamp(os.path.getmtime(folder_path)).isoformat(),
            "lastModified": datetime.fromtimestamp(os.path.getmtime(folder_path)).isoformat(),
            "folder": folder_path,
            "created": datetime.fromtimestamp(os.path.getctime(folder_path)).isoformat(),
            "size": "G: Drive",
            "fileCount": len([f for f in raw_data.keys() if f.endswith('.json')])
        }
        
        # Load Sales Orders data - ULTRA-FAST OPTIMIZED
        print("Loading Sales Orders - Performance Optimized...")
        try:
            from so_performance_optimizer import get_optimized_so_data, get_so_performance_stats
            try:
                from so_background_refresh import start_so_background_refresh, get_so_refresh_status
                # Start background refresh service if not already running
                try:
                    start_so_background_refresh()
                except Exception as bg_error:
                    print(f"⚠️ Background refresh service not available: {bg_error}")
            except ImportError as import_error:
                print(f"⚠️ so_background_refresh module not available: {import_error}")
            except Exception as bg_error:
                print(f"⚠️ Background refresh service error: {bg_error}")
            
            
            sales_orders_data = get_optimized_so_data()
            if sales_orders_data:
                raw_data.update(sales_orders_data)
                load_time = sales_orders_data.get('LoadTime', 0)
                total_orders = sales_orders_data.get('TotalOrders', 0)
                load_method = sales_orders_data.get('LoadMethod', 'Unknown')
                print(f"⚡ SO LOAD: {total_orders} orders in {load_time:.3f}s ({load_method})")
                
                # Add performance stats to response
                perf_stats = get_so_performance_stats()
                raw_data['SOPerformanceStats'] = perf_stats
                
                # Add background refresh status
                try:
                    refresh_status = get_so_refresh_status()
                    raw_data['SOBackgroundRefresh'] = refresh_status
                except:
                    pass
        except Exception as e:
            print(f"⚠️ SO Optimizer not available, falling back to standard loader: {e}")
            # Fallback to original method
            sales_orders_data = load_sales_orders()
            if sales_orders_data:
                raw_data.update(sales_orders_data)
                print(f"SUCCESS: Added {sales_orders_data.get('TotalOrders', 0)} sales orders to data")
        
        # Load cached parsed SO data for instant lookups
        print("RETRY: Loading cached parsed SO data...")
        cached_so_data = load_cached_so_data()
        if cached_so_data:
            raw_data.update(cached_so_data)
            print(f"SUCCESS: Added {len(cached_so_data.get('ParsedSalesOrders.json', []))} parsed SOs to data")
        
        # Enterprise SO Service integration (lazy import to avoid circular dependency)
        try:
            # Use importlib to avoid circular import issues
            import importlib
            enterprise_so_module = importlib.import_module('enterprise_so_service')
            get_so_service_health = getattr(enterprise_so_module, 'get_so_service_health', None)
            if get_so_service_health:
                so_health = get_so_service_health()
                raw_data['SOServiceHealth'] = so_health
                print(f"SUCCESS: SO Service Health: {so_health['status']} - {so_health['total_sos']} SOs cached")
            else:
                print("⚠️ Enterprise SO Service function not found")
        except ImportError as import_error:
            print(f"⚠️ Enterprise SO Service module not available: {import_error}")
        except Exception as e:
            print(f"⚠️ Enterprise SO Service error: {e}")
        
        # Load MPS (Master Production Schedule) data
        print("RETRY: Loading MPS data...")
        mps_data = load_mps_data()
        if mps_data and 'error' not in mps_data:
            raw_data['MPS.json'] = mps_data
            print(f"SUCCESS: Added MPS data with {len(mps_data.get('mps_orders', []))} production orders")
        else:
            print(f"MPS data not available: {mps_data.get('error', 'Unknown error')}")
            raw_data['MPS.json'] = {"mps_orders": [], "summary": {"total_orders": 0}}
        
        print(f"SUCCESS: Successfully loaded data from {latest_folder}")
        raw_data = _merge_portal_store(raw_data)
        
        # SAFE DATA SUMMARY - Only process list data types
        safe_summary = []
        for k, v in raw_data.items():
            if isinstance(v, list):
                safe_summary.append(f'{k}: {len(v)} records')
            elif isinstance(v, (str, int, float, bool)):
                safe_summary.append(f'{k}: {type(v).__name__} value')
            else:
                safe_summary.append(f'{k}: {type(v).__name__}')
        
        print(f"📊 Data summary: {safe_summary}")
        
        # Cache the data AND pre-serialize the response for future requests
        _data_cache = raw_data
        _cache_timestamp = time.time()
        
        # Detect if running on Cloud Run (no G: Drive access) vs local (has G: Drive)
        is_cloud_run = os.getenv('K_SERVICE') is not None
        data_source = "Google Drive API" if is_cloud_run else "Local G: Drive"
        
        # Pre-serialize the response to avoid 35s jsonify overhead on cache hits
        import json as json_module
        response_dict = {
            "data": raw_data,
            "folderInfo": folder_info,
            "LoadTimestamp": datetime.now().isoformat(),
            "cached": False,
            "source": data_source
        }
        _response_cache = json_module.dumps(response_dict)
        print(f"💾 Data + Response cached for {_cache_duration} seconds (pre-serialized: {len(_response_cache)/1024/1024:.1f}MB)")
        
        from flask import Response
        return Response(_response_cache, mimetype='application/json')
        
    except Exception as e:
        import traceback
        print(f"ERROR: Error in get_all_data: {e}")
        traceback.print_exc()
        # Return 200 with empty structure so the app can load (e.g. on Render when G: Drive/API unavailable)
        empty_data = get_empty_app_data_structure()
        empty_data['ScanMethod'] = 'Error'
        return jsonify({
            "data": empty_data,
            "folderInfo": {
                "folderName": "Data load failed",
                "syncDate": datetime.now().isoformat(),
                "lastModified": datetime.now().isoformat(),
                "folder": "Not Connected",
                "created": datetime.now().isoformat(),
                "size": "0",
                "fileCount": 0,
            },
            "LoadTimestamp": datetime.now().isoformat(),
            "source": "error",
            "message": str(e),
        })

@app.route('/api/data-source', methods=['GET'])
def get_data_source_status():
    """Framework: report which data source is in use and whether Full Company Data is available/ready."""
    try:
        # Check if Full Company Data folder is reachable (list count)
        full_company_available = False
        try:
            if not IS_CLOUD_ENVIRONMENT and os.path.exists(GDRIVE_FULL_COMPANY_DATA):
                full_company_available = len([f for f in os.listdir(GDRIVE_FULL_COMPANY_DATA) if os.path.isfile(os.path.join(GDRIVE_FULL_COMPANY_DATA, f))]) > 0
            else:
                service = get_google_drive_service()
                if service and getattr(service, 'authenticated', False):
                    drive_id = service.find_shared_drive("IT_Automation")
                    folder_id = service.find_folder_by_path(drive_id, FULL_COMPANY_DATA_DRIVE_PATH) if drive_id else None
                    if folder_id:
                        files = service.list_all_files_in_folder(folder_id, drive_id)
                        full_company_available = len(files) > 0
        except Exception:
            full_company_available = False
        return jsonify({
            "currentSource": "default",
            "defaultLabel": "MISys Full Company Data (export folder)",
            "fullCompanyDataPath": FULL_COMPANY_DATA_DRIVE_PATH,
            "fullCompanyDataAvailable": full_company_available,
            "fullCompanyDataReady": full_company_available,
            "message": "Data comes from your MISys Full Company Data export. Place your export CSVs in the configured folder; no separate API extractions are required.",
        })
    except Exception as e:
        return jsonify({"error": str(e), "currentSource": "default", "fullCompanyDataReady": False}), 500


def _get_company_data_for_export():
    """Return raw company data dict for export. Uses _data_cache if set; else tries Google Drive API. Returns (data_dict, None) or (None, error_message)."""
    global _data_cache
    if _data_cache is not None and isinstance(_data_cache, dict):
        return _data_cache, None
    try:
        gdrive_service = get_google_drive_service()
        if gdrive_service and getattr(gdrive_service, 'authenticated', False):
            gdrive_data, _ = gdrive_service.get_all_data()
            if gdrive_data:
                return _merge_portal_store(gdrive_data), None
    except Exception as e:
        print(f"Export: Google Drive API fallback failed: {e}")
    return None, "No data available. Open the app or call GET /api/data first to load data, then try export again."


def _export_company_data_xlsx(data):
    """Build one Excel workbook, one sheet per list entity. Returns bytes."""
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    for key in sorted(data.keys()):
        val = data[key]
        if not isinstance(val, list) or not val:
            continue
        sheet_name = key.replace(".json", "")[:31]
        sheet_name = re.sub(r'[\:\*\?\/\\\[\]]', '_', sheet_name)
        ws = wb.create_sheet(title=sheet_name)
        if val and isinstance(val[0], dict):
            headers = list(val[0].keys())
            for col, h in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=str(h))
            for row_idx, row in enumerate(val, 2):
                for col_idx, h in enumerate(headers, 1):
                    v = row.get(h)
                    if v is not None:
                        ws.cell(row=row_idx, column=col_idx, value=v)
        else:
            for row_idx, row in enumerate(val, 1):
                ws.cell(row=row_idx, column=1, value=str(row))
    if len(wb.worksheets) == 0:
        ws = wb.create_sheet(title="Info")
        ws.cell(row=1, column=1, value="No list data to export.")
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _export_company_data_csv_single(data):
    """Export first substantial list (e.g. CustomAlert5) as single CSV. Returns bytes."""
    for key in ["CustomAlert5.json", "ManufacturingOrderHeaders.json", "MIILOC.json", "Items.json"]:
        val = data.get(key)
        if isinstance(val, list) and val and isinstance(val[0], dict):
            buf = io.StringIO()
            writer = csv.writer(buf)
            headers = list(val[0].keys())
            writer.writerow(headers)
            for row in val:
                writer.writerow([row.get(h) for h in headers])
            return buf.getvalue().encode("utf-8-sig")
    return b""


def _export_company_data_csv_multiple(data):
    """One CSV per list entity. Returns dict of filename -> bytes."""
    out = {}
    for key in sorted(data.keys()):
        val = data[key]
        if not isinstance(val, list) or not val:
            continue
        name = key.replace(".json", "") + ".csv"
        if isinstance(val[0], dict):
            buf = io.StringIO()
            writer = csv.writer(buf)
            writer.writerow(list(val[0].keys()))
            for row in val:
                writer.writerow([row.get(h) for h in val[0].keys()])
            out[name] = buf.getvalue().encode("utf-8-sig")
        else:
            out[name] = "\n".join(str(r) for r in val).encode("utf-8-sig")
    return out


def _export_company_data_xml_single(data):
    """One XML with root and one element per entity (list as child elements). Returns bytes."""
    root = ET.Element("CompanyData")
    for key in sorted(data.keys()):
        val = data[key]
        if not isinstance(val, list):
            continue
        entity = ET.SubElement(root, "entity", name=key.replace(".json", ""))
        for row in val:
            if isinstance(row, dict):
                row_el = ET.SubElement(entity, "row")
                for k, v in row.items():
                    if v is not None:
                        child = ET.SubElement(row_el, "cell", key=str(k))
                        child.text = str(v)
            else:
                ET.SubElement(entity, "row").text = str(row)
    return ET.tostring(root, encoding="unicode", default_namespace="").encode("utf-8")


def _export_company_data_xml_multiple(data):
    """One XML file per list entity. Returns dict of filename -> bytes."""
    out = {}
    for key in sorted(data.keys()):
        val = data[key]
        if not isinstance(val, list) or not val:
            continue
        name = key.replace(".json", "") + ".xml"
        root = ET.Element("data", name=key.replace(".json", ""))
        for row in val:
            if isinstance(row, dict):
                row_el = ET.SubElement(root, "row")
                for k, v in row.items():
                    if v is not None:
                        child = ET.SubElement(row_el, "cell", key=str(k))
                        child.text = str(v)
            else:
                ET.SubElement(root, "row").text = str(row)
        out[name] = ET.tostring(root, encoding="unicode", default_namespace="").encode("utf-8")
    return out


@app.route('/api/export/company-data', methods=['GET'])
def export_company_data():
    """Export all company data as xlsx, csv, or xml. User-friendly and automation: GET with format and optional multiple.
    Query: format=xlsx|csv|xml, multiple=true|false (for csv/xml), filename=... (optional, for Content-Disposition).
    """
    try:
        data, err = _get_company_data_for_export()
        if err:
            return jsonify({"error": err}), 503
        fmt = (request.args.get("format") or "xlsx").strip().lower()
        multiple = request.args.get("multiple", "false").strip().lower() in ("true", "1", "yes")
        filename_param = (request.args.get("filename") or "").strip()
        base_name = filename_param or f"company_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        if fmt == "xlsx":
            raw = _export_company_data_xlsx(data)
            fname = f"{base_name}.xlsx"
            return send_file(
                io.BytesIO(raw),
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                as_attachment=True,
                download_name=fname,
            )
        if fmt == "csv":
            if multiple:
                files = _export_company_data_csv_multiple(data)
                if not files:
                    return jsonify({"error": "No list data to export"}), 400
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for name, content in files.items():
                        zf.writestr(name, content)
                buf.seek(0)
                fname = f"{base_name}.zip"
                return send_file(buf, mimetype="application/zip", as_attachment=True, download_name=fname)
            raw = _export_company_data_csv_single(data)
            if not raw:
                return jsonify({"error": "No list data to export"}), 400
            fname = f"{base_name}.csv"
            return send_file(io.BytesIO(raw), mimetype="text/csv; charset=utf-8", as_attachment=True, download_name=fname)
        if fmt == "xml":
            if multiple:
                files = _export_company_data_xml_multiple(data)
                if not files:
                    return jsonify({"error": "No list data to export"}), 400
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for name, content in files.items():
                        zf.writestr(name, content)
                buf.seek(0)
                fname = f"{base_name}.zip"
                return send_file(buf, mimetype="application/zip", as_attachment=True, download_name=fname)
            raw = _export_company_data_xml_single(data)
            fname = f"{base_name}.xml"
            return send_file(io.BytesIO(raw), mimetype="application/xml; charset=utf-8", as_attachment=True, download_name=fname)
        return jsonify({"error": "format must be xlsx, csv, or xml"}), 400
    except Exception as e:
        print(f"ERROR export_company_data: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/manufacturing-orders', methods=['POST'])
def create_manufacturing_order():
    """Create a new manufacturing order. Persisted via portal_store (survives restart)."""
    global _cache_timestamp, _response_cache
    try:
        body = request.get_json() or {}
        build_item_no = (body.get('build_item_no') or body.get('Build Item No.') or '').strip()
        quantity = body.get('quantity') or body.get('Ordered') or body.get('ordered')
        due_date = (body.get('due_date') or body.get('Due Date') or body.get('Sales Order Ship Date') or '').strip()
        batch_number = (body.get('batch_number') or body.get('Batch No.') or body.get('Batch Number') or '').strip()
        sales_order_no = (body.get('sales_order_no') or body.get('SO No.') or body.get('Sales Order No.') or '').strip()
        description = (body.get('description') or '').strip()

        if not build_item_no:
            return jsonify({"error": "build_item_no is required"}), 400
        try:
            qty = float(quantity) if quantity is not None else 0
        except (TypeError, ValueError):
            qty = 0
        if qty <= 0:
            return jsonify({"error": "quantity must be greater than 0"}), 400

        # Generate MO number (portal-created: prefix so we can tell from MISys)
        import uuid
        mo_no = f"MO-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        now_iso = datetime.now().isoformat()
        order_date = datetime.now().strftime('%Y-%m-%d') if due_date else ''
        due_date_val = due_date if due_date else order_date

        mo_record = {
            'Mfg. Order No.': mo_no,
            'Build Item No.': build_item_no,
            'Ordered': qty,
            'Completed': 0,
            'Status': 0,  # Pending
            'Order Date': order_date,
            'Release Date': '',
            'Completion Date': '',
            'Sales Order Ship Date': due_date_val or '',
            'Customer': description or 'Portal',
            'Description': description or build_item_no,
            'Location No.': '',
            'Projected Material Cost': 0,
            'Cumulative Cost': 0,
            'Batch No.': batch_number,
            'Sales Order No.': sales_order_no or '',
            '_created_at': now_iso,
            '_source': 'portal',
        }
        if portal_store:
            portal_store.add_created_mo(mo_record)
        _cache_timestamp = None
        _response_cache = None
        print(f"Created MO: {mo_no} Item={build_item_no} Qty={qty} Batch={batch_number or '(none)'} SO={sales_order_no or '(none)'}")
        return jsonify(mo_record), 201
    except Exception as e:
        print(f"ERROR create_manufacturing_order: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/manufacturing-orders/<mo_no>/release', methods=['POST'])
def release_manufacturing_order(mo_no):
    """D5: Release MO (status to released/in progress). Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        mo_no = (mo_no or '').strip()
        if not mo_no:
            return jsonify({"error": "mo_no required"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        now = datetime.now().strftime('%Y-%m-%d')
        portal_store.add_mo_update(mo_no, status=1, release_date=now)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "mo_no": mo_no, "status": 1, "release_date": now}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _get_bom_components_for_item(data, parent_item_no):
    """Return list of (component_item_no, required_qty) for the given parent (build) item."""
    out = []
    for key in ("BillOfMaterialDetails.json", "MIBOMD.json"):
        details = data.get(key) or []
        for row in details:
            if not isinstance(row, dict):
                continue
            parent = (row.get("Parent Item No.") or row.get("bomItem") or "").strip()
            if parent != parent_item_no:
                continue
            comp = (row.get("Component Item No.") or row.get("partId") or "").strip()
            qty = float(row.get("Required Quantity") or row.get("qty") or 0)
            if comp:
                out.append((comp, qty))
    return out


@app.route('/api/manufacturing-orders/<mo_no>/complete', methods=['POST'])
def complete_manufacturing_order(mo_no):
    """D6/D7: Complete MO (report completed qty, set status, backflush components). Persisted in portal_store."""
    global _cache_timestamp, _response_cache, _data_cache
    try:
        mo_no = (mo_no or '').strip()
        if not mo_no:
            return jsonify({"error": "mo_no required"}), 400
        body = request.get_json() or {}
        completed_qty = body.get("completed_qty") or body.get("Completed") or body.get("completed")
        try:
            completed_qty = float(completed_qty) if completed_qty is not None else 0
        except (TypeError, ValueError):
            completed_qty = 0
        if completed_qty <= 0:
            return jsonify({"error": "completed_qty must be positive"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        data = _data_cache
        if not data or not isinstance(data, dict):
            return jsonify({"error": "Load app data first (open the app or call GET /api/data), then complete MO"}), 503
        all_mos = data.get("ManufacturingOrderHeaders.json") or []
        build_item_no = None
        for mo in all_mos:
            if not isinstance(mo, dict):
                continue
            if (mo.get("Mfg. Order No.") or mo.get("mohId")) == mo_no:
                build_item_no = (mo.get("Build Item No.") or mo.get("buildItem") or "").strip()
                break
        if not build_item_no:
            return jsonify({"error": f"MO {mo_no} not found"}), 404
        components = _get_bom_components_for_item(data, build_item_no)
        for comp_item, req_qty in components:
            portal_store.add_inventory_adjustment(
                comp_item, "", - (req_qty * completed_qty),
                reason=f"Backflush MO {mo_no}",
            )
        now = datetime.now().strftime('%Y-%m-%d')
        portal_store.add_mo_update(mo_no, status=2, completed=completed_qty, completion_date=now)
        lot = (body.get("lot") or body.get("Lot No.") or body.get("batch_number") or "").strip()
        if lot and build_item_no:
            portal_store.add_mo_completion_lot(mo_no, build_item_no, completed_qty, lot)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "mo_no": mo_no, "completed_qty": completed_qty, "backflushed": len(components), "lot": lot or None}), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/work-orders/<wo_no>/release', methods=['POST'])
def release_work_order(wo_no):
    """E2: Release WO (status to released). Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        wo_no = (wo_no or '').strip()
        if not wo_no:
            return jsonify({"error": "wo_no required"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        now = datetime.now().strftime('%Y-%m-%d')
        portal_store.add_wo_update(wo_no, status=1, release_date=now)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "wo_no": wo_no, "status": 1, "release_date": now}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/work-orders/<wo_no>/complete', methods=['POST'])
def complete_work_order(wo_no):
    """E3: Report WO completion (qty, scrap). Persisted in portal_store."""
    global _cache_timestamp, _response_cache, _data_cache
    try:
        wo_no = (wo_no or '').strip()
        if not wo_no:
            return jsonify({"error": "wo_no required"}), 400
        body = request.get_json() or {}
        completed_qty = body.get("completed_qty") or body.get("Completed") or body.get("completed")
        try:
            completed_qty = float(completed_qty) if completed_qty is not None else 0
        except (TypeError, ValueError):
            completed_qty = 0
        if completed_qty <= 0:
            return jsonify({"error": "completed_qty must be positive"}), 400
        scrap = body.get("scrap") or body.get("Scrap")
        try:
            scrap = float(scrap) if scrap is not None else 0
        except (TypeError, ValueError):
            scrap = 0
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        data = _data_cache
        current_completed = 0
        if data and isinstance(data, dict):
            for wo in (data.get("WorkOrders.json") or []):
                if not isinstance(wo, dict):
                    continue
                if (wo.get("Work Order No.") or wo.get("Job No.") or "").strip() == wo_no:
                    current_completed = float(wo.get("Completed") or 0)
                    break
        new_total = current_completed + completed_qty
        now = datetime.now().strftime('%Y-%m-%d')
        portal_store.add_wo_update(wo_no, status=2, completed=new_total, completion_date=now, scrap=scrap if scrap else None)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "wo_no": wo_no, "completed_qty": completed_qty, "total_completed": new_total, "scrap": scrap}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/inventory/adjustment', methods=['POST'])
def inventory_adjustment():
    """B5: Add or remove qty for an item (optionally by location). Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        body = request.get_json() or {}
        item_no = (body.get('item_no') or body.get('Item No.') or '').strip()
        location = (body.get('location') or body.get('Location No.') or '').strip()
        delta = body.get('delta') or body.get('qty')
        reason = (body.get('reason') or '').strip()
        if not item_no:
            return jsonify({"error": "item_no is required"}), 400
        try:
            delta = float(delta)
        except (TypeError, ValueError):
            return jsonify({"error": "delta must be a number"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        portal_store.add_inventory_adjustment(item_no, location, delta, reason)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "item_no": item_no, "delta": delta}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/inventory/transfer', methods=['POST'])
def inventory_transfer():
    """B6: Move qty of an item from one location to another. Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        body = request.get_json() or {}
        from_loc = (body.get('from_loc') or body.get('From Location') or '').strip()
        to_loc = (body.get('to_loc') or body.get('To Location') or '').strip()
        item_no = (body.get('item_no') or body.get('Item No.') or '').strip()
        qty = body.get('qty') or body.get('quantity')
        if not item_no or not from_loc or not to_loc:
            return jsonify({"error": "item_no, from_loc, to_loc are required"}), 400
        try:
            qty = float(qty)
        except (TypeError, ValueError):
            return jsonify({"error": "qty must be a number"}), 400
        if qty <= 0:
            return jsonify({"error": "qty must be positive"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        portal_store.add_location_transfer(from_loc, to_loc, item_no, qty)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "item_no": item_no, "from_loc": from_loc, "to_loc": to_loc, "qty": qty}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/items/<item_no>/reorder', methods=['PATCH', 'PUT'])
def item_reorder_update(item_no):
    """B4: Update min/max/reorder level/reorder qty for an item. Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        item_no = (item_no or '').strip()
        if not item_no:
            return jsonify({"error": "item_no required"}), 400
        body = request.get_json() or {}
        minimum = body.get('Minimum') if 'Minimum' in body else body.get('minimum')
        maximum = body.get('Maximum') if 'Maximum' in body else body.get('maximum')
        reorder_level = body.get('Reorder Level') if 'Reorder Level' in body else body.get('reorder_level')
        reorder_quantity = body.get('Reorder Quantity') if 'Reorder Quantity' in body else body.get('reorder_quantity')
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        portal_store.set_item_override(item_no, minimum=minimum, maximum=maximum, reorder_level=reorder_level, reorder_quantity=reorder_quantity)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "item_no": item_no}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/bom', methods=['POST'])
def bom_create():
    """C5: Create BOM header and/or lines. Persisted in portal_store."""
    global _cache_timestamp, _response_cache
    try:
        body = request.get_json() or {}
        parent = (body.get('parent_item_no') or body.get('Parent Item No.') or body.get('bomItem') or '').strip()
        revision = (body.get('revision') or body.get('Revision No.') or body.get('bomRev') or '1').strip()
        components = body.get('components') or body.get('lines') or []
        if not parent:
            return jsonify({"error": "parent_item_no (or Parent Item No.) is required"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        header = {
            'Parent Item No.': parent,
            'Revision No.': revision,
            'BOM Item': parent,
            'BOM Rev': revision,
            'Build Quantity': float(body.get('Build Quantity') or body.get('mult') or 1),
            '_source': 'portal',
        }
        portal_store.add_bom_header(header)
        for line in components:
            comp = (line.get('component_item_no') or line.get('Component Item No.') or line.get('partId') or '').strip()
            qty = float(line.get('Required Quantity') or line.get('qty') or 0)
            if not comp:
                continue
            detail = {
                'Parent Item No.': parent,
                'Revision No.': revision,
                'Component Item No.': comp,
                'Required Quantity': qty,
                'BOM Item': parent,
                'BOM Rev': revision,
                'partId': comp,
                'qty': qty,
                '_source': 'portal',
            }
            portal_store.add_bom_detail(detail)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "parent_item_no": parent, "revision": revision, "lines": len(components)}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/purchase-orders', methods=['POST'])
def create_purchase_order():
    """F2: Create PO – full header (supplier, terms, ship, cost, etc.) + lines. Persisted via portal_store."""
    global _cache_timestamp, _response_cache
    try:
        body = request.get_json() or {}
        supplier_no = (body.get('supplier_no') or body.get('Supplier No.') or '').strip()
        supplier_name = (body.get('supplier_name') or body.get('Name') or body.get('supplier_name') or '').strip()
        lines = body.get('lines') or body.get('Lines') or []
        if not supplier_no:
            return jsonify({"error": "supplier_no (or Supplier No.) is required"}), 400
        if not lines or not isinstance(lines, list):
            return jsonify({"error": "lines array is required (e.g. [{item_no, qty, unit_cost?}]"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        import uuid
        po_no = f"PO-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        now_date = datetime.now().strftime('%Y-%m-%d')
        order_date = (body.get('order_date') or body.get('Order Date') or now_date)
        if isinstance(order_date, str):
            order_date = order_date.strip() or now_date
        else:
            order_date = now_date
        requested_date = (body.get('requested_date') or body.get('Requested Date') or '').strip() or None
        terms = (body.get('terms') or body.get('Terms') or '').strip() or None
        ship_via = (body.get('ship_via') or body.get('Ship Via') or '').strip() or None
        fob = (body.get('fob') or body.get('FOB') or '').strip() or None
        contact = (body.get('contact') or body.get('Contact') or '').strip() or None
        buyer = (body.get('buyer') or body.get('Buyer') or '').strip() or None
        try:
            freight = float(body.get('freight') or body.get('Freight') or 0)
        except (TypeError, ValueError):
            freight = 0
        currency = (body.get('currency') or body.get('Source Currency') or body.get('Home Currency') or 'USD').strip()
        description = (body.get('description') or body.get('Description') or '').strip() or None
        lines_subtotal = 0.0
        valid_details = []
        for idx, line in enumerate(lines):
            item_no = (line.get('item_no') or line.get('Item No.') or '').strip()
            qty = float(line.get('qty') or line.get('Ordered') or line.get('Quantity') or 0)
            unit_cost = float(line.get('unit_cost') or line.get('Unit Cost') or line.get('unit_price') or line.get('Unit Price') or 0)
            if not item_no or qty <= 0:
                continue
            line_desc = (line.get('description') or line.get('Description') or '').strip() or None
            required_date = (line.get('required_date') or line.get('Required Date') or line.get('Due Date') or '').strip() or None
            ext = qty * unit_cost
            lines_subtotal += ext
            detail = {
                'PO No.': po_no,
                'Item No.': item_no,
                'Ordered': qty,
                'Received': 0,
                'Unit Cost': unit_cost,
                'Unit Price': unit_cost,
                'Extended Price': round(ext, 2),
                'Line No.': len(valid_details) + 1,
                '_source': 'portal',
            }
            if line_desc:
                detail['Description'] = line_desc
            if required_date:
                detail['Required Date'] = required_date
            valid_details.append(detail)
            portal_store.add_created_po_detail(detail)
        total_amount = round(lines_subtotal + freight, 2)
        header = {
            'PO No.': po_no,
            'Supplier No.': supplier_no,
            'Name': supplier_name or supplier_no,
            'Order Date': order_date,
            'Status': 0,
            'Total Amount': total_amount,
            '_source': 'portal',
        }
        if requested_date:
            header['Requested Date'] = requested_date
        if terms:
            header['Terms'] = terms
        if ship_via:
            header['Ship Via'] = ship_via
        if fob:
            header['FOB'] = fob
        if contact:
            header['Contact'] = contact
        if buyer:
            header['Buyer'] = buyer
        if freight and freight != 0:
            header['Freight'] = round(freight, 2)
        if currency:
            header['Source Currency'] = currency
            header['Home Currency'] = currency
        if description:
            header['Description'] = description
        portal_store.add_created_po(header)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "po_no": po_no, "lines": len(valid_details), "total_amount": total_amount}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/purchase-orders/<po_no>/receive', methods=['POST'])
def receive_against_po(po_no):
    """F4: Receive against PO – qty, location, optional lot. Updates PO line Received and inventory."""
    global _cache_timestamp, _response_cache
    try:
        po_no = (po_no or '').strip()
        if not po_no:
            return jsonify({"error": "po_no required"}), 400
        body = request.get_json() or {}
        item_no = (body.get('item_no') or body.get('Item No.') or '').strip()
        qty = body.get('qty') or body.get('quantity')
        location = (body.get('location') or body.get('Location No.') or '').strip()
        lot = (body.get('lot') or body.get('Lot No.') or '').strip()
        if not item_no:
            return jsonify({"error": "item_no (or Item No.) is required"}), 400
        try:
            qty = float(qty)
        except (TypeError, ValueError):
            return jsonify({"error": "qty must be a number"}), 400
        if qty <= 0:
            return jsonify({"error": "qty must be positive"}), 400
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        portal_store.add_po_receive(po_no, item_no, qty, location, lot)
        portal_store.add_inventory_adjustment(item_no, location, qty, f"PO receive {po_no}")
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "po_no": po_no, "item_no": item_no, "qty": qty}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/inventory/by-lot', methods=['GET'])
def inventory_by_lot():
    """H4: Inventory by lot – aggregates portal PO receives and MO completion lots."""
    try:
        if portal_store is None:
            return jsonify({"by_lot": []}), 200
        from collections import defaultdict
        balances = defaultdict(float)
        for rec in portal_store.get_po_receives() or []:
            key = (rec.get("item_no") or "", rec.get("lot") or "")
            balances[key] += float(rec.get("qty", 0))
        for rec in portal_store.get_mo_completion_lots() or []:
            key = (rec.get("item_no") or "", rec.get("lot") or "")
            balances[key] += float(rec.get("qty", 0))
        by_lot = [{"item_no": k[0], "lot": k[1], "qty": v} for k, v in balances.items() if v != 0]
        return jsonify({"by_lot": by_lot}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/lot-history', methods=['GET'])
def lot_history():
    """H5: Lot/serial history – PO receives and MO completions with lot."""
    try:
        if portal_store is None:
            return jsonify({"history": []}), 200
        history = []
        for rec in portal_store.get_po_receives() or []:
            history.append({
                "type": "receive",
                "po_no": rec.get("po_no"),
                "item_no": rec.get("item_no"),
                "qty": float(rec.get("qty", 0)),
                "location": rec.get("location"),
                "lot": rec.get("lot") or "",
                "at": rec.get("at"),
            })
        for rec in portal_store.get_mo_completion_lots() or []:
            history.append({
                "type": "production",
                "mo_no": rec.get("mo_no"),
                "item_no": rec.get("item_no"),
                "qty": float(rec.get("qty", 0)),
                "lot": rec.get("lot") or "",
                "at": rec.get("at"),
            })
        history.sort(key=lambda x: x.get("at") or "", reverse=True)
        return jsonify({"history": history}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _compute_shortage(data):
    """J1/J2: Time-phased shortage – reorder_level - (stock + open_po). Returns list of { item_no, shortage_qty, on_hand, open_po, reorder_level }."""
    if not data or not isinstance(data, dict):
        return []
    items = data.get("CustomAlert5.json") or data.get("Items.json") or []
    details = data.get("PurchaseOrderDetails.json") or []
    open_po_by_item = {}
    for line in details:
        if not isinstance(line, dict):
            continue
        item_no = line.get("Item No.") or line.get("Item No") or ""
        ordered = float(line.get("Ordered") or line.get("Ordered Qty") or 0)
        received = float(line.get("Received") or line.get("Received Qty") or 0)
        if item_no and ordered > received:
            open_po_by_item[item_no] = open_po_by_item.get(item_no, 0) + (ordered - received)
    shortage_list = []
    for item in items:
        if not isinstance(item, dict):
            continue
        item_no = item.get("Item No.") or item.get("item_no") or ""
        if not item_no:
            continue
        reorder_level = float(item.get("Reorder Level") or item.get("reorder_level") or 0)
        if reorder_level <= 0:
            continue
        on_hand = float(item.get("Stock") or item.get("stock") or 0)
        open_po = open_po_by_item.get(item_no, 0)
        shortage = max(0, reorder_level - (on_hand + open_po))
        if shortage > 0:
            shortage_list.append({
                "item_no": item_no,
                "shortage_qty": round(shortage, 4),
                "on_hand": on_hand,
                "open_po": open_po,
                "reorder_level": reorder_level,
                "reorder_quantity": float(item.get("Reorder Quantity") or item.get("reorder_quantity") or 0),
            })
    return shortage_list


@app.route('/api/shortage', methods=['GET'])
def get_shortage():
    """J2: Shortage report – items below reorder level (stock + open PO). Uses cached app data."""
    global _data_cache
    try:
        data = _data_cache
        if not data:
            return jsonify({"shortage": [], "message": "Load app data first (GET /api/data)"}), 200
        shortage = _compute_shortage(data)
        return jsonify({"shortage": shortage}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/mrp/auto-create-po', methods=['POST'])
def mrp_auto_create_po():
    """J3/N2: Auto-create PO from shortage. Body: optional item_nos[] and supplier_no (or use first shortage item's last supplier)."""
    global _cache_timestamp, _response_cache, _data_cache
    try:
        data = _data_cache
        if not data:
            return jsonify({"error": "Load app data first (GET /api/data)"}), 503
        body = request.get_json() or {}
        item_nos = body.get("item_nos") or body.get("items") or []
        supplier_no = (body.get("supplier_no") or body.get("Supplier No.") or "").strip()
        shortage = _compute_shortage(data)
        if not item_nos and shortage:
            item_nos = [s["item_no"] for s in shortage[:50]]
        if not item_nos:
            return jsonify({"ok": True, "message": "No shortage items to order", "po_no": None}), 200
        if not supplier_no:
            for line in (data.get("PurchaseOrderDetails.json") or []):
                if not isinstance(line, dict):
                    continue
                if (line.get("Item No.") or line.get("Item No")) in item_nos:
                    po_no = line.get("PO No.")
                    for po in (data.get("PurchaseOrders.json") or []):
                        if isinstance(po, dict) and (po.get("PO No.") or po.get("Purchase Order No")) == po_no:
                            supplier_no = (po.get("Supplier No.") or po.get("Name") or "").strip()
                            break
                    if supplier_no:
                        break
        if not supplier_no:
            supplier_no = "SUPPLIER"
        shortage_map = {s["item_no"]: s for s in shortage}
        lines = []
        for ino in item_nos:
            qty = 1
            if ino in shortage_map:
                qty = max(shortage_map[ino]["shortage_qty"], shortage_map[ino].get("reorder_quantity") or 1)
            lines.append({"item_no": ino, "qty": qty, "unit_cost": 0})
        if portal_store is None:
            return jsonify({"error": "Portal store not available"}), 503
        import uuid
        po_no = f"PO-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        now_date = datetime.now().strftime('%Y-%m-%d')
        header = {"PO No.": po_no, "Supplier No.": supplier_no, "Name": supplier_no, "Order Date": now_date, "Status": 0, "Total Amount": 0, "_source": "portal"}
        portal_store.add_created_po(header)
        for idx, line in enumerate(lines):
            detail = {"PO No.": po_no, "Item No.": line["item_no"], "Ordered": line["qty"], "Received": 0, "Unit Cost": line.get("unit_cost", 0), "Line No.": idx + 1, "_source": "portal"}
            portal_store.add_created_po_detail(detail)
        _cache_timestamp = None
        _response_cache = None
        return jsonify({"ok": True, "po_no": po_no, "lines": len(lines), "supplier_no": supplier_no}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint - FAST, doesn't check G: Drive (Cloud Run can't access it)
    CRITICAL: This endpoint must ALWAYS return 200 to allow Cloud Run to start serving traffic
    """
    try:
        # Check Google Drive API status (but don't fail if it errors)
        google_drive_api_enabled = False
        google_drive_authenticated = False
        
        try:
            if USE_GOOGLE_DRIVE_API:
                google_drive_api_enabled = True
                service = get_google_drive_service()
                if service is not None and hasattr(service, 'authenticated'):
                    google_drive_authenticated = service.authenticated
        except Exception as gdrive_err:
            # Don't fail health check if Google Drive check fails
            print(f"⚠️ Health check: Google Drive status check failed (non-critical): {gdrive_err}")
        
        # ALWAYS return 200 - server is ready to accept requests
        # Data loading can happen in background
        return jsonify({
            "status": "ready",  # Changed from "healthy" to "ready" for frontend compatibility
            "message": "Backend is ready and accepting requests",
            "timestamp": datetime.now().isoformat(),
            "google_drive_api_enabled": google_drive_api_enabled,
            "google_drive_authenticated": google_drive_authenticated,
            "is_cloud_run": IS_CLOUD_RUN
        }), 200
    except Exception as e:
        # Even if there's an error, try to return 200 so Cloud Run doesn't kill the container
        print(f"⚠️ Health check error (returning ready anyway): {e}")
        return jsonify({
            "status": "ready",
            "message": "Backend is ready (health check had minor error)",
            "timestamp": datetime.now().isoformat(),
            "error": str(e)
        }), 200

def scan_folder_recursively(folder_path, status, path_parts=[]):
    """Recursively scan any folder structure and find all PDF files"""
    orders = []
    
    try:
        if not os.path.exists(folder_path):
            return orders
            
        items = os.listdir(folder_path)
        so_files = [item for item in items if item.lower().endswith(('.pdf', '.doc', '.docx'))]
        subfolders = [item for item in items if os.path.isdir(os.path.join(folder_path, item)) and item.lower() not in ['desktop.ini', 'thumbs.db']]
        
        # Process ALL SO files in current folder (PDF, DOC, DOCX)
        for file in so_files:
            # Look for any SO pattern, not just "salesorder_"
            if any(pattern in file.lower() for pattern in ['salesorder', 'sales order', 'so_', 'so-', 'order']):
                try:
                    # Extract order number from various filename patterns
                    order_num = 'Unknown'
                    import re
                    
                    # Try different patterns to extract SO number
                    patterns = [
                        r'salesorder[_\s-]*(\d+)',
                        r'sales\s*order[_\s-]*(\d+)',
                        r'so[_\s-]*(\d+)',
                        r'order[_\s-]*(\d+)',
                        r'(\d+)'  # fallback - any number
                    ]
                    
                    for pattern in patterns:
                        match = re.search(pattern, file.lower())
                        if match:
                            order_num = match.group(1)
                            break
                    
                    file_path = os.path.join(folder_path, file)
                    file_stat = os.stat(file_path)
                    
                    # Build comprehensive path hierarchy
                    path_info = {
                        'Order No.': order_num,
                        'Customer': 'Customer Data',
                        'Order Date': datetime.fromtimestamp(file_stat.st_ctime).strftime('%Y-%m-%d'),
                        'Ship Date': datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d'),
                        'Status': status,
                        'File': file,
                        'File Path': file_path,
                        'File Type': file.split('.')[-1].upper(),
                        'Last Modified': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                        'Folder Path': '/'.join(path_parts) if path_parts else 'Root',
                        'Full Path': '/'.join([status] + path_parts) if path_parts else status
                    }
                    
                    # Add dynamic path levels
                    for i, part in enumerate(path_parts):
                        path_info[f'Level_{i+1}'] = part
                    
                    orders.append(path_info)
                    
                except Exception as e:
                    print(f"ERROR: Error processing file {file}: {e}")
        
        # Recursively process subfolders
        for subfolder in subfolders:
            subfolder_path = os.path.join(folder_path, subfolder)
            new_path_parts = path_parts + [subfolder]
            subfolder_orders = scan_folder_recursively(subfolder_path, status, new_path_parts)
            orders.extend(subfolder_orders)
            
            if subfolder_orders:
                path_str = '/'.join(new_path_parts)
                print(f"📄 Found {len(subfolder_orders)} orders in {status}/{path_str}")
        
        return orders
        
    except Exception as e:
        print(f"ERROR: Error scanning folder {folder_path}: {e}")
        return orders

def load_cached_so_data():
    """Load cached parsed SO data for instant lookups"""
    try:
        cache_dir = "cache"
        parsed_sos_file = os.path.join(cache_dir, "ParsedSalesOrders.json")
        item_index_file = os.path.join(cache_dir, "SOItemIndex.json")
        cache_status_file = os.path.join(cache_dir, "SOCacheStatus.json")
        
        cached_data = {}
        
        # Load parsed SOs
        if os.path.exists(parsed_sos_file):
            with open(parsed_sos_file, 'r') as f:
                parsed_sos = json.load(f)
                cached_data['ParsedSalesOrders.json'] = parsed_sos
                print(f"SUCCESS: Loaded {len(parsed_sos)} parsed SOs from cache")
        
        # Load item index
        if os.path.exists(item_index_file):
            with open(item_index_file, 'r') as f:
                item_index = json.load(f)
                cached_data['SOItemIndex.json'] = item_index
                print(f"SUCCESS: Loaded {len(item_index)} item index entries from cache")
        
        # Load cache status
        if os.path.exists(cache_status_file):
            with open(cache_status_file, 'r') as f:
                cache_status = json.load(f)
                cached_data['SOCacheStatus'] = cache_status
                print(f"SUCCESS: Cache last updated: {cache_status.get('last_updated', 'Unknown')}")
        
        return cached_data if cached_data else None
        
    except Exception as e:
        print(f"ERROR: Error loading cached SO data: {e}")
        return None

def load_sales_orders():
    """Smart Sales Orders loader - discovers ANY folder structure dynamically"""
    try:
        # Detect if we're on Cloud Run (can't access local G: Drive)
        is_cloud_run = os.getenv('K_SERVICE') is not None
        
        if is_cloud_run:
            # Cloud Run: Use Google Drive API directly (G: Drive doesn't exist in cloud)
            print("☁️ Cloud Run detected - using Google Drive API for Sales Orders")
            gdrive_service = get_google_drive_service()
            if gdrive_service and gdrive_service.authenticated:
                try:
                    # load_sales_orders_data() finds Sales_CSR drive automatically (ignores drive_id parameter)
                    # Only load active folders (In Production, New and Revised) to reduce response size
                    so_data = gdrive_service.load_sales_orders_data(None, filter_folders=['In Production', 'New and Revised'])
                    if so_data and so_data.get('TotalOrders', 0) > 0:
                        print(f"✅ Successfully loaded {so_data.get('TotalOrders', 0)} sales orders from Google Drive API")
                        return so_data
                    else:
                        print("⚠️ Google Drive API returned empty sales orders data")
                except Exception as e:
                    print(f"❌ Error loading sales orders from Google Drive API: {e}")
                    import traceback
                    traceback.print_exc()
            else:
                print("⚠️ Google Drive API not available on Cloud Run")
                return {
                    'SalesOrders.json': [],
                    'SalesOrdersByStatus': {},
                    'TotalOrders': 0,
                    'StatusFolders': [],
                    'ScanMethod': 'Cloud Run - Google Drive API unavailable',
                    'LoadTimestamp': datetime.now().isoformat()
                }
        
        # Local development: Try G: Drive first, fallback to Google Drive API
        print(f"💻 Local development - checking G: Drive: {SALES_ORDERS_BASE}")
        
        if not os.path.exists(SALES_ORDERS_BASE):
            print(f"⚠️ G: Drive not accessible: {SALES_ORDERS_BASE}")
            
            # Try Google Drive API as fallback (for local dev when G: Drive is disconnected)
            gdrive_service = get_google_drive_service()
            if gdrive_service and gdrive_service.authenticated:
                print("🔄 Falling back to Google Drive API for Sales Orders...")
                try:
                    so_data = gdrive_service.load_sales_orders_data(None, filter_folders=['In Production', 'New and Revised'])
                    if so_data and so_data.get('TotalOrders', 0) > 0:
                        print(f"✅ Successfully loaded {so_data.get('TotalOrders', 0)} sales orders from Google Drive API")
                        return so_data
                except Exception as e:
                    print(f"❌ Error loading sales orders from Google Drive API: {e}")
            
            # Both failed
            print("⚠️ Both G: Drive and Google Drive API unavailable - returning empty data")
            return {
                'SalesOrders.json': [],
                'SalesOrdersByStatus': {},
                'TotalOrders': 0,
                'StatusFolders': [],
                'ScanMethod': 'No G: Drive Access',
                'LoadTimestamp': datetime.now().isoformat()
            }
        
        # Discover all status folders dynamically
        base_items = os.listdir(SALES_ORDERS_BASE)
        status_folders = [item for item in base_items if os.path.isdir(os.path.join(SALES_ORDERS_BASE, item)) and item != 'desktop.ini']
        
        print(f"SEARCH: Discovered status folders: {status_folders}")
        
        sales_data = {}
        all_orders = []
        
        for status in status_folders:
            folder_path = os.path.join(SALES_ORDERS_BASE, status)
            print(f"RETRY: Scanning {status}...")
            
            # Recursively scan this status folder
            orders = scan_folder_recursively(folder_path, status)
            sales_data[status] = orders
            all_orders.extend(orders)
            
            print(f"SUCCESS: Total loaded {len(orders)} orders from {status}")
        
        # Smart sorting by order number
        def smart_sort(order):
            try:
                order_num = order['Order No.']
                import re
                numbers = re.findall(r'\d+', order_num)
                return int(numbers[0]) if numbers else 0
            except:
                return 0
        
        all_orders.sort(key=smart_sort, reverse=True)
        
        print(f"🎉 Smart scan complete! Found {len(all_orders)} total sales orders across {len(status_folders)} status folders")
        
        return {
            'SalesOrders.json': all_orders,
            'SalesOrdersByStatus': sales_data,
            'LoadTimestamp': datetime.now().isoformat(),
            'TotalOrders': len(all_orders),
            'StatusFolders': status_folders,
            'ScanMethod': 'Smart Recursive Discovery'
        }
        
    except Exception as e:
        print(f"ERROR: Error in smart sales orders loading: {e}")
        return {}

def load_mps_from_excel():
    """Fallback: Load MPS data from Excel file"""
    try:
        MPS_BASE = r"G:\Shared drives\IT_Automation\MiSys\Misys Extracted Data\API Extractions"
        MPS_EXCEL_PATH = os.path.join(MPS_BASE, "MPS.xlsx")
        
        if not os.path.exists(MPS_EXCEL_PATH):
            print(f"ERROR: MPS Excel file not found at: {MPS_EXCEL_PATH}")
            return {"error": "MPS file not found"}
        
        print(f"📊 Loading MPS data from Excel: {MPS_EXCEL_PATH}")
        
        # Load Excel file
        import pandas as pd
        df = pd.read_excel(MPS_EXCEL_PATH)
        
        print(f"📊 Loaded {len(df)} rows from Excel file")
        
        # Convert to MPS format
        mps_orders = []
        for _, row in df.iterrows():
            mps_order = {
                'order_id': str(row.get('Order ID', '')),
                'item_code': str(row.get('Item Code', '')),
                'description': str(row.get('Description', '')),
                'quantity': float(row.get('Quantity', 1)),
                'start_date': str(row.get('Start Date', '')),
                'due_date': str(row.get('Due Date', '')),
                'status': str(row.get('Status', 'scheduled')),
                'priority': str(row.get('Priority', 'medium')),
                'customer': str(row.get('Customer', '')),
                'work_center': str(row.get('Work Center', 'Production')),
                'estimated_hours': float(row.get('Estimated Hours', 8)),
                'actual_hours': float(row.get('Actual Hours', 0)),
                'materials': [],
                'revenue': float(row.get('Revenue', 0))
            }
            mps_orders.append(mps_order)
        
        print(f"📊 Processed {len(mps_orders)} MPS orders from Excel")
        
        return {
            'mps_orders': mps_orders,
            'source': 'excel',
            'total_orders': len(mps_orders)
        }
        
    except Exception as e:
        print(f"ERROR: Failed to load MPS from Excel: {e}")
        return {"error": f"Excel loading failed: {e}"}

def load_mps_data():
    """Load MPS (Master Production Schedule) data from Google Sheets using alternative method"""
    # Import requests at the function level
    import requests
    
    try:
        # Use Google Sheets API to read the MPS data directly
        sheet_id = MPS_SHEET_ID
        
        print(f"Loading MPS data from Google Sheets: {sheet_id}")
        
        # Use the working URL format immediately (format 4 that works)
        csv_url = MPS_CSV_URL
        
        print(f"Using working Google Sheets URL: {csv_url}")
        
        try:
            response = requests.get(csv_url, timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            })
            
            if response.status_code == 200:
                csv_data = response.text
                print(f"Success with Google Sheets URL")
            else:
                print(f"ERROR: Google Sheets failed: Status {response.status_code}")
                return load_mps_from_excel()
                
        except Exception as e:
            print(f"ERROR: Google Sheets error: {e}")
            return load_mps_from_excel()
        
        
        # Parse CSV data
        import csv
        import io
        
        reader = csv.reader(io.StringIO(csv_data))
        rows = list(reader)
        
        print(f"Retrieved {len(rows)} rows from Google Sheets")
        
        if len(rows) < 5:
            print("ERROR: Not enough data in Google Sheets")
            return load_mps_from_excel()
        
        # Find header row (should be row 4, index 3)
        headers = None
        data_start_row = 4  # Row 5 in 1-based indexing
        
        if len(rows) > 3:
            headers = rows[3]  # Row 4 (0-based index 3)
            print(f"Headers: {headers[:10]}...")  # Show first 10 headers
        
        if not headers:
            print("ERROR: Could not find headers in Google Sheets")
            return load_mps_from_excel()
        
        # Process production orders
        mps_orders = []
        
        for i, row in enumerate(rows[data_start_row:], start=data_start_row + 1):
            if len(row) < 2:  # Skip empty rows
                continue
            
            # Skip rows where first column is empty or just whitespace
            if not row[0] or not row[0].strip():
                continue
            
            try:
                # Map row data to our structure - CORRECT column mapping:
                # 0=Line#, 1=SO, 2=MO, 3=WIP, 4=WorkCenter, 5=Status, 6=Product, 
                # 7=CustomerCode, 8=Packaging, 9=Required, 10=Ready, 11=Planned%, 
                # 12=Actual%, 13=Promised, 14=StartDate, 15=EndDate, 16=Duration, 
                # 17=DTC, 18=ActionItems
                order_data = {
                    'order_number': row[0].strip() if len(row) > 0 else '',
                    'so_number': row[1].strip() if len(row) > 1 else '',
                    'mo_number': row[2].strip() if len(row) > 2 else '',
                    'wip': row[3].strip() if len(row) > 3 else '',
                    'work_center': row[4].strip() if len(row) > 4 else '',
                    'status': row[5].strip() if len(row) > 5 else '',
                    'product': row[6].strip() if len(row) > 6 else '',
                    'customer_code': row[7].strip() if len(row) > 7 else '',
                    'packaging': row[8].strip() if len(row) > 8 else '',
                    'required': row[9].strip() if len(row) > 9 else '',
                    'ready': row[10].strip() if len(row) > 10 else '',
                    'planned': row[11].strip() if len(row) > 11 else '',
                    'actual': row[12].strip() if len(row) > 12 else '',
                    'promised': row[13].strip() if len(row) > 13 else '',
                    'start_date': row[14].strip() if len(row) > 14 else '',
                    'end_date': row[15].strip() if len(row) > 15 else '',
                    'duration': row[16].strip() if len(row) > 16 else '',
                    'dtc': row[17].strip() if len(row) > 17 else '',
                    'action_items': row[18].strip() if len(row) > 18 else ''
                }
                
                # Only add if we have essential data
                if order_data['order_number'] and order_data['so_number']:
                    mps_orders.append(order_data)
                    
            except Exception as e:
                print(f"⚠️ Error processing row {i}: {e}")
                continue
        
        print(f"SUCCESS: Processed {len(mps_orders)} MPS orders from Google Sheets")
        
        if mps_orders:
            print(f"Sample order: {mps_orders[0]}")
        
        return {
            "mps_orders": mps_orders,
            "summary": {
                "total_orders": len(mps_orders),
                "source": "Google Sheets API",
                "sheet_id": sheet_id,
                "last_updated": datetime.now().isoformat()
            }
        }
        
    except requests.RequestException as e:
        print(f"ERROR: Error accessing Google Sheets: {e}")
        return {"error": f"Could not access Google Sheets: {str(e)}"}
    except Exception as e:
        print(f"ERROR: Error parsing Google Sheets data: {e}")
        return {"error": f"Error parsing data: {str(e)}"}

@app.route('/api/mps', methods=['GET'])
def get_mps_data():
    """Get MPS schedule from Google Sheets - returns RAW CSV for frontend to parse.
    This matches the reference implementation exactly."""
    import requests as req
    try:
        # Fetch CSV directly from Google Sheets and return as-is
        # This lets the frontend parse it with PapaParse (same as reference app)
        response = req.get(MPS_CSV_URL, timeout=15, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        if response.ok:
            return response.text, 200, {'Content-Type': 'text/csv'}
        return jsonify({'error': 'Failed to fetch MPS data from Google Sheets'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/warmup', methods=['GET'])
def warmup():
    """Lightweight warmup endpoint"""
    return jsonify({
        'status': 'warm',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/data/clear-cache', methods=['POST'])
def clear_data_cache():
    """Clear the data cache to force fresh reload"""
    global _data_cache, _cache_timestamp
    _data_cache = None
    _cache_timestamp = None
    print("Data cache cleared - next request will load fresh data")
    return jsonify({"message": "Cache cleared successfully"})

@app.route('/api/data/lazy-load', methods=['POST'])
def lazy_load_additional_data():
    """Lazy load additional data files when needed"""
    try:
        data = request.get_json()
        requested_files = data.get('files', [])
        
        if not requested_files:
            return jsonify({"error": "No files specified"}), 400
        
        # Check if G: Drive is accessible
        if not os.path.exists(GDRIVE_BASE):
            return jsonify({"error": "G: Drive not accessible"}), 500
        
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": error}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        loaded_files = {}
        
        for file_name in requested_files:
            file_path = os.path.join(folder_path, file_name)
            if os.path.exists(file_path):
                file_data = load_json_file(file_path)
                loaded_files[file_name] = file_data
                print(f"SUCCESS: Lazy loaded {file_name} with {len(file_data) if isinstance(file_data, list) else 1} records")
            else:
                loaded_files[file_name] = []
                print(f"⚠️ {file_name} not found")
        
        return jsonify({
            "data": loaded_files,
            "LoadTimestamp": datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"ERROR: Error in lazy_load_additional_data: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/sales-orders', methods=['GET'])
def get_sales_orders():
    """Get all Sales Orders from G: Drive - organized by status"""
    try:
        sales_data = load_sales_orders()
        return jsonify(sales_data)
    except Exception as e:
        print(f"ERROR: Error in get_sales_orders: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/enterprise/so-service/health', methods=['GET'])
def get_enterprise_so_health():
    """Get Enterprise SO Service health status"""
    try:
        from enterprise_so_service import get_so_service_health
        health = get_so_service_health()
        return jsonify(health)
    except Exception as e:
        return jsonify({"error": str(e), "status": "unavailable"}), 500

@app.route('/api/so-performance', methods=['GET'])
def get_so_performance():
    """Get SO loading performance statistics"""
    try:
        from so_performance_optimizer import get_so_performance_stats
        stats = get_so_performance_stats()
        return jsonify({
            "status": "success",
            "performance": stats,
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e), "status": "unavailable"}), 500

@app.route('/api/so-performance/refresh', methods=['POST'])
def refresh_so_performance_cache():
    """Force refresh SO performance cache"""
    try:
        from so_performance_optimizer import force_so_refresh
        refreshed_data = force_so_refresh()
        return jsonify({
            "status": "success",
            "message": "SO cache refreshed successfully",
            "data": {
                "total_orders": refreshed_data.get('TotalOrders', 0),
                "load_time": refreshed_data.get('LoadTime', 0),
                "load_method": refreshed_data.get('LoadMethod', 'Unknown')
            },
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e), "status": "failed"}), 500

@app.route('/api/so-background-refresh', methods=['GET'])
def get_so_background_refresh_status():
    """Get SO background refresh service status"""
    try:
        from so_background_refresh import get_so_refresh_status
        status = get_so_refresh_status()
        return jsonify({
            "status": "success",
            "background_refresh": status,
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e), "status": "unavailable"}), 500

@app.route('/api/so-background-refresh/force', methods=['POST'])
def force_so_background_refresh():
    """Force immediate SO background refresh"""
    try:
        from so_background_refresh import force_so_refresh_now
        force_so_refresh_now()
        return jsonify({
            "status": "success",
            "message": "Background refresh triggered successfully",
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e), "status": "failed"}), 500

@app.route('/api/full-company-data/list', methods=['GET'])
def full_company_data_list():
    """List files in the Full Company Data folder (core MISys export). Uses existing data path; no conflict with current app data."""
    try:
        # Local path: list from G: drive folder
        if not IS_CLOUD_ENVIRONMENT and os.path.exists(GDRIVE_FULL_COMPANY_DATA):
            names = [f for f in os.listdir(GDRIVE_FULL_COMPANY_DATA) if os.path.isfile(os.path.join(GDRIVE_FULL_COMPANY_DATA, f))]
            files = [{"id": "", "name": n, "mimeType": "local"} for n in sorted(names)]
            return jsonify({
                "path": GDRIVE_FULL_COMPANY_DATA,
                "source": "local",
                "files": files,
                "count": len(files)
            })
        # Cloud or API: use Google Drive API
        service = get_google_drive_service()
        if not service or not getattr(service, 'authenticated', False):
            return jsonify({"error": "Google Drive not available", "path": FULL_COMPANY_DATA_DRIVE_PATH}), 503
        drive_id = service.find_shared_drive("IT_Automation")
        if not drive_id:
            return jsonify({"error": "Shared drive IT_Automation not found", "path": FULL_COMPANY_DATA_DRIVE_PATH}), 404
        folder_id = service.find_folder_by_path(drive_id, FULL_COMPANY_DATA_DRIVE_PATH)
        if not folder_id:
            return jsonify({"error": "Full Company Data folder not found", "path": FULL_COMPANY_DATA_DRIVE_PATH}), 404
        files = service.list_all_files_in_folder(folder_id, drive_id)
        return jsonify({
            "path": FULL_COMPANY_DATA_DRIVE_PATH,
            "source": "google_drive_api",
            "files": files,
            "count": len(files)
        })
    except Exception as e:
        print(f"ERROR full_company_data_list: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/full-company-data/preview', methods=['GET'])
def full_company_data_preview():
    """Preview first rows of a file in Full Company Data (for mapping). Query param: file=Item.csv"""
    import pandas as pd
    from flask import request
    file_name = request.args.get("file", "").strip()
    if not file_name:
        return jsonify({"error": "Missing query param: file=YourFile.csv"}), 400
    max_rows = min(int(request.args.get("rows", 20)), 100)
    try:
        # Local path
        if not IS_CLOUD_ENVIRONMENT and os.path.exists(GDRIVE_FULL_COMPANY_DATA):
            file_path = os.path.join(GDRIVE_FULL_COMPANY_DATA, file_name)
            if not os.path.isfile(file_path):
                return jsonify({"error": f"File not found: {file_name}"}), 404
            if file_name.lower().endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='skip', nrows=max_rows)
            elif file_name.lower().endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path, nrows=max_rows)
            else:
                return jsonify({"error": "Only .csv and .xlsx supported for preview"}), 400
            columns = list(df.columns)
            rows = df.fillna("").to_dict(orient="records")
            return jsonify({"file": file_name, "columns": columns, "rows": rows, "rowCount": len(rows)})
        # Google Drive API
        service = get_google_drive_service()
        if not service or not getattr(service, 'authenticated', False):
            return jsonify({"error": "Google Drive not available"}), 503
        drive_id = service.find_shared_drive("IT_Automation")
        folder_id = service.find_folder_by_path(drive_id, FULL_COMPANY_DATA_DRIVE_PATH)
        if not folder_id:
            return jsonify({"error": "Full Company Data folder not found"}), 404
        query = f"name='{file_name}' and '{folder_id}' in parents and trashed=false"
        list_params = {"q": query, "supportsAllDrives": True, "includeItemsFromAllDrives": True, "fields": "files(id, name)", "pageSize": 1}
        if drive_id:
            list_params["corpora"] = "drive"
            list_params["driveId"] = drive_id
        results = service.service.files().list(**list_params).execute()
        files = results.get("files", [])
        if not files:
            return jsonify({"error": f"File not found: {file_name}"}), 404
        content = service.download_file(files[0]["id"], file_name)
        if content is None:
            return jsonify({"error": "Failed to download file"}), 500
        if isinstance(content, (list, dict)):
            return jsonify({"error": "File is JSON, use list endpoint; preview is for CSV/Excel"}), 400
        raw = content.decode("utf-8", errors="replace")
        if file_name.lower().endswith(".csv"):
            from io import StringIO
            df = pd.read_csv(StringIO(raw), on_bad_lines="skip", nrows=max_rows)
        elif file_name.lower().endswith((".xlsx", ".xls")):
            from io import BytesIO
            df = pd.read_excel(BytesIO(content), nrows=max_rows)
        else:
            return jsonify({"error": "Only .csv and .xlsx supported for preview"}), 400
        columns = list(df.columns)
        rows = df.fillna("").to_dict(orient="records")
        return jsonify({"file": file_name, "columns": columns, "rows": rows, "rowCount": len(rows)})
    except Exception as e:
        print(f"ERROR full_company_data_preview: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/enterprise/so-service/refresh', methods=['POST'])
def refresh_so_cache():
    """Trigger incremental SO cache refresh"""
    try:
        from enterprise_so_service import EnterpriseSO, get_so_service_health
        service = EnterpriseSO()
        service.incremental_parse()
        
        health = get_so_service_health()
        return jsonify({
            "success": True,
            "message": "SO cache refreshed",
            "health": health
        })
    except Exception as e:
        return jsonify({"error": str(e), "success": False}), 500

@app.route('/api/sales-order-pdf/<path:file_path>', methods=['GET'])
def get_sales_order_pdf(file_path):
    """Serve a Sales Order PDF file by full path"""
    try:
        print(f"SEARCH: Serving PDF: {file_path}")
        
        # Decode URL encoding
        import urllib.parse
        decoded_path = urllib.parse.unquote(file_path)
        
        if not os.path.exists(decoded_path):
            print(f"ERROR: PDF file not found: {decoded_path}")
            return jsonify({"error": f"PDF file not found: {decoded_path}"}), 404
        
        # Serve the PDF file directly
        from flask import send_file
        filename = os.path.basename(decoded_path)
        
        return send_file(
            decoded_path,
            as_attachment=False,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"ERROR: Error serving PDF {file_path}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/gdrive/preview/<file_id>', methods=['GET'])
def preview_gdrive_file(file_id):
    """Preview a file from Google Drive by file ID - streams the file content"""
    try:
        print(f"📄 Previewing Google Drive file: {file_id}")
        
        gdrive_service = get_google_drive_service()
        if not gdrive_service or not gdrive_service.authenticated:
            return jsonify({"error": "Google Drive not authenticated"}), 503
        
        # Get file metadata first
        file_metadata = gdrive_service.service.files().get(
            fileId=file_id,
            fields='name, mimeType, size',
            supportsAllDrives=True
        ).execute()
        
        filename = file_metadata.get('name', 'file')
        mime_type = file_metadata.get('mimeType', 'application/octet-stream')
        
        print(f"📄 File: {filename}, Type: {mime_type}")
        
        # Download file content
        from io import BytesIO
        from googleapiclient.http import MediaIoBaseDownload
        
        request = gdrive_service.service.files().get_media(fileId=file_id)
        file_buffer = BytesIO()
        downloader = MediaIoBaseDownload(file_buffer, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_buffer.seek(0)
        
        from flask import send_file
        return send_file(
            file_buffer,
            as_attachment=False,
            download_name=filename,
            mimetype=mime_type
        )
        
    except Exception as e:
        print(f"❌ Error previewing Google Drive file {file_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/gdrive/download/<file_id>', methods=['GET'])
def download_gdrive_file(file_id):
    """Download a file from Google Drive by file ID"""
    try:
        print(f"📥 Downloading Google Drive file: {file_id}")
        
        gdrive_service = get_google_drive_service()
        if not gdrive_service or not gdrive_service.authenticated:
            return jsonify({"error": "Google Drive not authenticated"}), 503
        
        # Get file metadata first
        file_metadata = gdrive_service.service.files().get(
            fileId=file_id,
            fields='name, mimeType, size',
            supportsAllDrives=True
        ).execute()
        
        filename = file_metadata.get('name', 'file')
        mime_type = file_metadata.get('mimeType', 'application/octet-stream')
        
        print(f"📥 Downloading: {filename}, Type: {mime_type}")
        
        # Download file content
        from io import BytesIO
        from googleapiclient.http import MediaIoBaseDownload
        
        request = gdrive_service.service.files().get_media(fileId=file_id)
        file_buffer = BytesIO()
        downloader = MediaIoBaseDownload(file_buffer, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_buffer.seek(0)
        
        from flask import send_file
        return send_file(
            file_buffer,
            as_attachment=True,  # Force download
            download_name=filename,
            mimetype=mime_type
        )
        
    except Exception as e:
        print(f"❌ Error downloading Google Drive file {file_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/sales-orders/folder/<path:folder_path>', methods=['GET'])
def get_sales_order_folder(folder_path):
    """Get Sales Order folder contents dynamically - REAL TIME SYNC
    
    Works on both:
    - Local development: Uses G: Drive directly
    - Cloud (Render/Cloud Run): Uses Google Drive API
    """
    try:
        # URL decode the path to handle spaces and special characters
        from urllib.parse import unquote
        folder_path = unquote(folder_path)
        print(f"SEARCH: Loading Sales Order folder: {folder_path}")
        
        # Detect if we're on cloud (Render or Cloud Run)
        is_cloud = os.getenv('RENDER') is not None or os.getenv('K_SERVICE') is not None
        
        # Check if G: Drive is accessible (local development)
        full_path = os.path.join(SALES_ORDERS_BASE, folder_path)
        use_gdrive_api = is_cloud or not os.path.exists(SALES_ORDERS_BASE)
        
        if use_gdrive_api:
            # Cloud deployment or G: Drive not accessible - use Google Drive API
            print(f"☁️ Using Google Drive API for folder: {folder_path}")
            gdrive_service = get_google_drive_service()
            
            if not gdrive_service or not gdrive_service.authenticated:
                print("❌ Google Drive API not available")
                return jsonify({
                    "error": "Google Drive API not available",
                    "folders": [],
                    "files": [],
                    "total_folders": 0,
                    "total_files": 0
                }), 503
            
            # Use the new browse method
            result = gdrive_service.browse_sales_orders_folder(folder_path)
            
            if 'error' in result and result.get('folders') == [] and result.get('files') == []:
                print(f"❌ Folder not found via Google Drive API: {folder_path}")
                return jsonify(result), 404
            
            print(f"✅ Loaded {result.get('total_folders', 0)} folders and {result.get('total_files', 0)} files via Google Drive API")
            return jsonify(result)
        
        # Local development with G: Drive access
        print(f"💻 Using local G: Drive: {full_path}")
        
        if not os.path.exists(full_path):
            print(f"ERROR: Folder not found at path: {full_path}")
            return jsonify({"error": f"Folder not found: {folder_path}"}), 404
        
        # Get folder contents
        items = os.listdir(full_path)
        folders = []
        files = []
        
        for item in items:
            item_path = os.path.join(full_path, item)
            if os.path.isdir(item_path) and item != 'desktop.ini':
                # Count files in subfolder
                try:
                    subfolder_items = os.listdir(item_path)
                    file_count = len([f for f in subfolder_items if os.path.isfile(os.path.join(item_path, f))])
                    folder_count = len([f for f in subfolder_items if os.path.isdir(os.path.join(item_path, f))])
                    
                    folders.append({
                        'name': item,
                        'type': 'folder',
                        'file_count': file_count,
                        'folder_count': folder_count,
                        'path': os.path.join(folder_path, item).replace('\\', '/')
                    })
                except:
                    folders.append({
                        'name': item,
                        'type': 'folder',
                        'file_count': 0,
                        'folder_count': 0,
                        'path': os.path.join(folder_path, item).replace('\\', '/')
                    })
                    
            elif os.path.isfile(item_path) and not item.startswith('.'):
                # Get file info
                try:
                    stat = os.stat(item_path)
                    size = stat.st_size
                    modified = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')
                    
                    files.append({
                        'name': item,
                        'type': 'file',
                        'size': size,
                        'modified': modified,
                        'path': item_path,
                        'is_pdf': item.lower().endswith('.pdf'),
                        'is_excel': item.lower().endswith(('.xlsx', '.xls'))
                    })
                except:
                    files.append({
                        'name': item,
                        'type': 'file',
                        'size': 0,
                        'modified': 'Unknown',
                        'path': item_path,
                        'is_pdf': item.lower().endswith('.pdf'),
                        'is_excel': item.lower().endswith(('.xlsx', '.xls'))
                    })
        
        # Sort folders and files
        folders.sort(key=lambda x: x['name'])
        files.sort(key=lambda x: x['name'])
        
        print(f"SUCCESS: Loaded {len(folders)} folders and {len(files)} files from {folder_path}")
        
        return jsonify({
            'path': folder_path,
            'full_path': full_path,
            'folders': folders,
            'files': files,
            'total_folders': len(folders),
            'total_files': len(files),
            'source': 'G: Drive'
        })
        
    except Exception as e:
        print(f"ERROR: Error loading folder {folder_path}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/sales-orders/find/<so_number>', methods=['GET'])
def find_sales_order_file(so_number):
    """Find a specific Sales Order file by SO number"""
    try:
        print(f"SEARCH: Searching for SO file: {so_number}")
        
        # Use the actual Sales Orders path from G: Drive
        base_path = SALES_ORDERS_BASE
        
        if not os.path.exists(base_path):
            return jsonify({"found": False, "error": "Sales Orders folder not found"}), 404
        
        # Search through all subfolders for the SO file - PRIORITIZE LATEST VERSION
        matching_files = []
        
        for root, dirs, files in os.walk(base_path):
            for file in files:
                if file.lower().endswith('.pdf'):
                    # Check if SO number is in filename
                    if so_number in file or f"SO_{so_number}" in file or f"SO{so_number}" in file:
                        file_path = os.path.join(root, file)
                        matching_files.append(file_path)
                        print(f"📄 Found SO PDF: {file_path}")
        
        # Sort files to prioritize latest revision (R999 > R3 > R2 > R1 > original)
        if matching_files:
            def get_version_priority(filepath):
                filename = os.path.basename(filepath).lower()
                # Extract revision number from filename - handles any revision number
                import re
                rev_match = re.search(r'_r(\d+)', filename)
                if rev_match:
                    rev_num = int(rev_match.group(1))
                    return 1000 + rev_num  # R999=1999, R3=1003, R2=1002, R1=1001
                else:
                    return 1  # Original version (lowest priority)
            
            matching_files.sort(key=get_version_priority, reverse=True)
            so_file_path = matching_files[0]
            selected_file = os.path.basename(so_file_path)
            print(f"SUCCESS: Selected latest version: {selected_file}")
            
            # Verify it's actually the latest
            if len(matching_files) > 1:
                print(f"📋 Available versions: {[os.path.basename(f) for f in matching_files[:3]]}")
            
            return jsonify({
                "found": True,
                "filePath": so_file_path,
                "fileName": selected_file,
                "folder": os.path.basename(os.path.dirname(so_file_path))
            })
        
        print(f"ERROR: SO file not found: {so_number}")
        return jsonify({"found": False, "error": f"SO {so_number} file not found"}), 404
        
    except Exception as e:
        print(f"ERROR: Error searching for SO file {so_number}: {e}")
        return jsonify({"found": False, "error": str(e)}), 500


@app.route('/api/check-changes', methods=['GET'])
def check_changes():
    """Check for real-time changes in G: Drive folders"""
    try:
        import time
        current_time = time.time()
        changes_detected = []
        
        # Check MISys data folder
        if os.path.exists(GDRIVE_BASE):
            for filename in os.listdir(GDRIVE_BASE):
                if filename.endswith('.json'):
                    file_path = os.path.join(GDRIVE_BASE, filename)
                    file_stat = os.stat(file_path)
                    
                    # Check if file was modified since last check
                    if not hasattr(check_changes, 'last_check_time'):
                        check_changes.last_check_time = {}
                    
                    if filename not in check_changes.last_check_time:
                        check_changes.last_check_time[filename] = file_stat.st_mtime
                    elif check_changes.last_check_time[filename] < file_stat.st_mtime:
                        changes_detected.append({
                            'type': 'file_modified',
                            'file': filename,
                            'path': GDRIVE_BASE
                        })
                        check_changes.last_check_time[filename] = file_stat.st_mtime
        
        # Check Sales Orders folder
        if os.path.exists(SALES_ORDERS_BASE):
            for status_folder in os.listdir(SALES_ORDERS_BASE):
                status_path = os.path.join(SALES_ORDERS_BASE, status_folder)
                if os.path.isdir(status_path):
                    # Check for new files or modifications
                    for root, dirs, files in os.walk(status_path):
                        for file in files:
                            if file.endswith('.pdf'):
                                file_path = os.path.join(root, file)
                                file_stat = os.stat(file_path)
                                
                                file_key = f"{status_folder}_{file}"
                                if not hasattr(check_changes, 'sales_last_check_time'):
                                    check_changes.sales_last_check_time = {}
                                
                                if file_key not in check_changes.sales_last_check_time:
                                    check_changes.sales_last_check_time[file_key] = file_stat.st_mtime
                                elif check_changes.sales_last_check_time[file_key] < file_stat.st_mtime:
                                    changes_detected.append({
                                        'type': 'sales_order_modified',
                                        'file': file,
                                        'file_path': file_path,
                                        'status': status_folder
                                    })
                                    check_changes.sales_last_check_time[file_key] = file_stat.st_mtime
        
        return jsonify({
            'hasChanges': len(changes_detected) > 0,
            'changes': changes_detected,
            'changeType': 'File modifications detected' if changes_detected else 'No changes',
            'folderPath': GDRIVE_BASE if changes_detected else '',
            'timestamp': current_time
        })
        
    except Exception as e:
        print(f"ERROR: Error checking for changes: {e}")
        return jsonify({"error": str(e)}), 500

def analyze_inventory_data(data, query):
    """Analyze inventory data to answer user queries"""
    try:
        # Extract relevant data structures - ONLY using CustomAlert5 for item data
        customalert5_items = data.get('CustomAlert5.json', [])
        bom_headers = data.get('BillsOfMaterial.json', [])
        bom_details = data.get('BillOfMaterialDetails.json', [])
        mo_headers = data.get('ManufacturingOrderHeaders.json', [])
        mo_details = data.get('ManufacturingOrderDetails.json', [])
        po_headers = data.get('PurchaseOrders.json', [])
        po_details = data.get('PurchaseOrderDetails.json', [])
        
        # Create a comprehensive data summary for ChatGPT
        data_summary = {
            "total_items": len(customalert5_items),
            "items_with_stock": len([item for item in customalert5_items if safe_float(item.get("Stock", 0)) > 0]),
            "total_boms": len(bom_headers),
            "active_manufacturing_orders": len([mo for mo in mo_headers if mo.get("Status", 0) in [1, 2]]),
            "active_purchase_orders": len([po for po in po_headers if po.get("Status", 0) == 1]),
        }
        
        # Sample some data for context
        sample_items = customalert5_items[:3] if customalert5_items else []
        sample_boms = bom_headers[:3] if bom_headers else []
        sample_mos = mo_headers[:3] if mo_headers else []
        
        return {
            "data_summary": data_summary,
            "sample_items": sample_items,
            "sample_boms": sample_boms,
            "sample_mos": sample_mos,
            "available_data_files": list(data.keys())
        }
    except Exception as e:
        print(f"ERROR: Error analyzing inventory data: {e}")
        return {"error": str(e)}

def find_item_usage_and_availability(data, item_description_or_no, quantity_needed=None):
    """Find detailed information about item usage and availability"""
    try:
        # Use ONLY CustomAlert5.json for item data - it has everything
        customalert5_items = data.get('CustomAlert5.json', [])
        bom_details = data.get('BillOfMaterialDetails.json', [])
        mo_headers = data.get('ManufacturingOrderHeaders.json', [])
        mo_details = data.get('ManufacturingOrderDetails.json', [])
        
        # Find the item - CustomAlert5 has all item info AND stock info
        target_item = None
        for item in customalert5_items:
            if (item_description_or_no.lower() in (item.get("Description", "").lower()) or 
                item_description_or_no.upper() == item.get("Item No.", "").upper()):
                target_item = item
                break
        
        if not target_item:
            return {"error": f"Item '{item_description_or_no}' not found"}
        
        item_no = target_item.get("Item No.", "")
        
        # Stock information is ALREADY in CustomAlert5 - no need for separate lookup
        stock_info = target_item  # CustomAlert5 has all stock fields: Stock, WIP, Reserve, On Order
        
        # Find where this item is used in BOMs
        used_in_boms = []
        for bom_detail in bom_details:
            if bom_detail.get("Component Item No.") == item_no:
                used_in_boms.append({
                    "parent_item": bom_detail.get("Parent Item No."),
                    "quantity_per": bom_detail.get("Quantity Per", 0),
                    "scrap_factor": bom_detail.get("Scrap Factor", 0)
                })
        
        # Find active manufacturing orders using this item
        active_usage = []
        for mo in mo_headers:
            if mo.get("Status", 0) in [1, 2]:  # Released or WIP
                # Check if this MO uses our item
                mo_no = mo.get("Mfg. Order No.", "")
                for mo_detail in mo_details:
                    if (mo_detail.get("Mfg. Order No.") == mo_no and 
                        mo_detail.get("Item No.") == item_no):
                        active_usage.append({
                            "mo_number": mo_no,
                            "build_item": mo.get("Build Item No."),
                            "quantity_needed": mo_detail.get("Quantity", 0),
                            "status": mo.get("Status"),
                            "due_date": mo.get("Due Date")
                        })
        
        # Calculate current stock and availability
        current_stock = 0
        if stock_info:
            current_stock = safe_float(stock_info.get("Quantity on Hand", 0))
        elif target_item:
            current_stock = safe_float(target_item.get("Quantity on Hand", 0))
        
        # Calculate allocated quantity
        allocated_qty = sum([usage.get("quantity_needed", 0) for usage in active_usage])
        available_qty = current_stock - allocated_qty
        
        result = {
            "item": target_item,
            "current_stock": current_stock,
            "allocated_quantity": allocated_qty,
            "available_quantity": available_qty,
            "used_in_boms": used_in_boms,
            "active_usage": active_usage,
            "stock_info": stock_info
        }
        
        # If quantity needed is specified, check if we have enough
        if quantity_needed:
            result["quantity_requested"] = quantity_needed
            result["sufficient_stock"] = available_qty >= quantity_needed
            result["shortfall"] = max(0, quantity_needed - available_qty)
        
        return result
        
    except Exception as e:
        print(f"ERROR: Error finding item usage: {e}")
        return {"error": str(e)}

def get_product_analytics_data(data, query):
    """
    RAG Analytics Function - Extract and calculate product analytics from all data sources
    Returns structured analytics data for sales, costs, profits, restocking, etc.
    """
    try:
        analytics = {
            "sales_data": [],
            "cost_data": [],
            "profit_analysis": [],
            "restocking_needs": [],
            "product_performance": [],
            "summary": {}
        }
        
        # Extract data sources
        items = data.get('CustomAlert5.json', [])
        so_headers = data.get('SalesOrderHeaders.json', [])
        so_details = data.get('SalesOrderDetails.json', [])
        real_sos = data.get('RealSalesOrders', [])
        po_headers = data.get('PurchaseOrders.json', [])
        po_details = data.get('PurchaseOrderDetails.json', [])
        mo_headers = data.get('ManufacturingOrderHeaders.json', [])
        mo_details = data.get('ManufacturingOrderDetails.json', [])
        
        # SALES ANALYTICS - Extract from Sales Orders
        sales_by_product = {}
        sales_by_customer = {}
        total_revenue = 0
        
        # Process structured SO data
        for so in so_headers:
            so_number = so.get('Order Number', so.get('SO Number', ''))
            customer = so.get('Customer Name', so.get('Customer', 'Unknown'))
            order_date = so.get('Order Date', so.get('Date', ''))
            total_amount = safe_float(so.get('Total Amount', so.get('Amount', 0)))
            
            if total_amount > 0:
                total_revenue += total_amount
                sales_by_customer[customer] = sales_by_customer.get(customer, 0) + total_amount
        
        # Process SO details for product-level sales
        for detail in so_details:
            item_no = detail.get('Item No.', detail.get('Item Number', ''))
            description = detail.get('Description', '')
            quantity = safe_float(detail.get('Quantity', detail.get('Qty', 0)))
            unit_price = safe_float(detail.get('Unit Price', detail.get('Price', 0)))
            line_total = quantity * unit_price
            
            if item_no and quantity > 0:
                if item_no not in sales_by_product:
                    sales_by_product[item_no] = {
                        'item_no': item_no,
                        'description': description,
                        'total_quantity_sold': 0,
                        'total_revenue': 0,
                        'avg_price': 0,
                        'order_count': 0
                    }
                sales_by_product[item_no]['total_quantity_sold'] += quantity
                sales_by_product[item_no]['total_revenue'] += line_total
                sales_by_product[item_no]['order_count'] += 1
        
        # Process real SO data from PDFs
        for so in real_sos:
            customer = so.get('customer_name', 'Unknown')
            total_amount = safe_float(so.get('total_amount', 0))
            if total_amount > 0:
                total_revenue += total_amount
                sales_by_customer[customer] = sales_by_customer.get(customer, 0) + total_amount
            
            for item in so.get('items', []):
                item_desc = item.get('description', '')
                quantity = safe_float(item.get('quantity', 0))
                price = safe_float(item.get('price', item.get('unit_price', 0)))
                line_total = quantity * price
                
                # Try to match with item number from items data
                item_no = None
                for itm in items:
                    if item_desc.lower() in itm.get('Description', '').lower():
                        item_no = itm.get('Item No.', '')
                        break
                
                if item_no and quantity > 0:
                    if item_no not in sales_by_product:
                        sales_by_product[item_no] = {
                            'item_no': item_no,
                            'description': item_desc,
                            'total_quantity_sold': 0,
                            'total_revenue': 0,
                            'avg_price': 0,
                            'order_count': 0
                        }
                    sales_by_product[item_no]['total_quantity_sold'] += quantity
                    sales_by_product[item_no]['total_revenue'] += line_total
                    sales_by_product[item_no]['order_count'] += 1
        
        # Calculate average prices
        for item_no, data in sales_by_product.items():
            if data['total_quantity_sold'] > 0:
                data['avg_price'] = data['total_revenue'] / data['total_quantity_sold']
        
        # COST ANALYTICS - Extract from Purchase Orders and Items
        cost_by_product = {}
        
        # Get costs from items (CustomAlert5 has Recent Cost, Standard Cost, Unit Cost)
        for item in items:
            item_no = item.get('Item No.', '')
            recent_cost = safe_float(item.get('Recent Cost', 0))
            standard_cost = safe_float(item.get('Standard Cost', 0))
            unit_cost = safe_float(item.get('Unit Cost', 0))
            
            # Priority: Recent Cost > Standard Cost > Unit Cost
            cost = recent_cost if recent_cost > 0 else (standard_cost if standard_cost > 0 else unit_cost)
            
            if item_no and cost > 0:
                cost_by_product[item_no] = {
                    'item_no': item_no,
                    'description': item.get('Description', ''),
                    'recent_cost': recent_cost,
                    'standard_cost': standard_cost,
                    'unit_cost': unit_cost,
                    'current_cost': cost
                }
        
        # Get costs from Purchase Orders (more recent/accurate)
        for po_detail in po_details:
            item_no = po_detail.get('Item No.', '')
            po_cost = safe_float(po_detail.get('Cost', po_detail.get('Unit Cost', 0)))
            
            if item_no and po_cost > 0:
                if item_no not in cost_by_product:
                    cost_by_product[item_no] = {
                        'item_no': item_no,
                        'description': '',
                        'recent_cost': 0,
                        'standard_cost': 0,
                        'unit_cost': 0,
                        'current_cost': po_cost
                    }
                # Update with PO cost if it's more recent
                if po_cost > cost_by_product[item_no].get('current_cost', 0):
                    cost_by_product[item_no]['current_cost'] = po_cost
                    cost_by_product[item_no]['recent_cost'] = po_cost
        
        # PROFIT ANALYSIS - Combine sales and costs
        profit_analysis = []
        for item_no, sales_info in sales_by_product.items():
            cost_info = cost_by_product.get(item_no, {})
            cost = cost_info.get('current_cost', 0)
            revenue = sales_info.get('total_revenue', 0)
            quantity = sales_info.get('total_quantity_sold', 0)
            
            if cost > 0 and quantity > 0:
                total_cost = cost * quantity
                profit = revenue - total_cost
                margin = (profit / revenue * 100) if revenue > 0 else 0
                
                profit_analysis.append({
                    'item_no': item_no,
                    'description': sales_info.get('description', cost_info.get('description', '')),
                    'quantity_sold': quantity,
                    'revenue': revenue,
                    'cost': cost,
                    'total_cost': total_cost,
                    'profit': profit,
                    'margin_percent': margin,
                    'avg_selling_price': sales_info.get('avg_price', 0)
                })
        
        # RESTOCKING ANALYSIS - Items below reorder level
        restocking_needs = []
        for item in items:
            item_no = item.get('Item No.', '')
            stock = safe_float(item.get('Stock', item.get('Quantity on Hand', 0)))
            reorder_level = safe_float(item.get('Reorder Level', item.get('Minimum', 0)))
            reorder_qty = safe_float(item.get('Reorder Quantity', item.get('Maximum', 0)))
            
            if reorder_level > 0 and stock < reorder_level:
                restocking_needs.append({
                    'item_no': item_no,
                    'description': item.get('Description', ''),
                    'current_stock': stock,
                    'reorder_level': reorder_level,
                    'reorder_quantity': reorder_qty if reorder_qty > 0 else reorder_level * 2,
                    'shortfall': reorder_level - stock
                })
        
        # PRODUCT PERFORMANCE - Top sellers, top revenue, etc.
        top_sellers = sorted(sales_by_product.values(), 
                           key=lambda x: x.get('total_quantity_sold', 0), 
                           reverse=True)[:10]
        top_revenue = sorted(sales_by_product.values(), 
                           key=lambda x: x.get('total_revenue', 0), 
                           reverse=True)[:10]
        top_profits = sorted(profit_analysis, 
                           key=lambda x: x.get('profit', 0), 
                           reverse=True)[:10]
        top_customers = sorted(sales_by_customer.items(), 
                             key=lambda x: x[1], 
                             reverse=True)[:10]
        
        # Build analytics summary
        analytics['sales_data'] = list(sales_by_product.values())
        analytics['cost_data'] = list(cost_by_product.values())
        analytics['profit_analysis'] = profit_analysis
        analytics['restocking_needs'] = restocking_needs
        analytics['product_performance'] = {
            'top_sellers': top_sellers,
            'top_revenue': top_revenue,
            'top_profits': top_profits,
            'top_customers': [{'customer': k, 'total_revenue': v} for k, v in top_customers]
        }
        analytics['summary'] = {
            'total_revenue': total_revenue,
            'total_products_sold': len(sales_by_product),
            'total_customers': len(sales_by_customer),
            'items_need_restocking': len(restocking_needs),
            'products_with_profit_data': len(profit_analysis),
            'total_sales_orders': len(so_headers) + len(real_sos)
        }
        
        return analytics
        
    except Exception as e:
        print(f"ERROR: Error in product analytics: {e}")
        import traceback
        traceback.print_exc()
        return {
            "sales_data": [],
            "cost_data": [],
            "profit_analysis": [],
            "restocking_needs": [],
            "product_performance": {},
            "summary": {},
            "error": str(e)
        }

@app.route('/api/chat', methods=['POST'])
def chat_query():
    """Handle ChatGPT queries about inventory data with smart SO search"""
    try:
        data = request.get_json()
        user_query = data.get('query', '')
        date_context = data.get('dateContext', {})
        
        # 🧠 SMART SALES ORDER DETECTION AND SEARCH
        import re
        
        # Only do smart search if OpenAI is available
        if not openai_available:
            print("⚠️ OpenAI not available, skipping smart SO search")
        
        # Check if this is a Sales Order query
        so_patterns = [
            r'sales?\s*order\s*#?\s*\d+',
            r'so\s*#?\s*\d+',
            r'order\s*#?\s*\d+',
            r'salesorder[_\s]*\d+'
        ]
        
        is_so_query = any(re.search(pattern, user_query.lower()) for pattern in so_patterns)
        
        if is_so_query and openai_available and client:
            print(f"🧠 Detected Sales Order query - using smart search")
            from smart_so_search import get_smart_search
            smart_search = get_smart_search(client)
            smart_result = smart_search.smart_so_search(user_query)
            
            if smart_result['success']:
                # Found the SO! Use this data for the AI response
                print(f"SUCCESS: Smart search found SO {smart_result['so_number']}")
                
                # Create a focused response using the found data
                so_data = smart_result['extracted_data'][0]  # Use first match
                
                focused_response = f"""**SALES ORDER #{smart_result['so_number']} FOUND!**

**📋 ORDER DETAILS:**
- **SO Number:** {so_data.get('so_number', 'N/A')}
- **Customer:** {so_data.get('customer_name', 'N/A')}
- **Order Date:** {so_data.get('order_date', 'N/A')}
- **Due Date:** {so_data.get('due_date', 'N/A')}
- **Total Amount:** ${so_data.get('total_amount', 0):,.2f}
- **Status:** {so_data.get('folder_status', so_data.get('status', 'N/A'))}
- **Items:** {len(so_data.get('items', []))} items

**📦 ITEMS:**"""
                
                for i, item in enumerate(so_data.get('items', [])[:5], 1):
                    focused_response += f"\n{i}. {item.get('description', 'N/A')} - Qty: {item.get('quantity', 'N/A')}"
                
                if len(so_data.get('items', [])) > 5:
                    focused_response += f"\n... and {len(so_data.get('items', [])) - 5} more items"
                
                focused_response += f"\n\n**📁 File Location:** {so_data.get('folder_status', 'N/A')} folder"
                focused_response += f"\n**📄 Source:** {os.path.basename(so_data.get('file_path', 'N/A'))}"
                
                return jsonify({
                    "query": user_query,
                    "response": focused_response,
                    "smart_search_used": True,
                    "so_number": smart_result['so_number'],
                    "data_context": {
                        "folder": "Smart SO Search",
                        "total_items": len(so_data.get('items', [])),
                        "so_total": so_data.get('total_amount', 0)
                    }
                })
            else:
                # Smart search failed, continue with normal processing but mention the search attempt
                print(f"ERROR: Smart search failed: {smart_result.get('error', 'Unknown error')}")
                user_query += f" [Smart search attempted but no files found for the requested SO number]"
        
        if not user_query:
            return jsonify({"error": "No query provided"}), 400
        
        print(f"🤖 Processing ChatGPT query: {user_query}")
        
        # Check if OpenAI is available
        if not openai_available:
            return jsonify({
                "query": user_query,
                "response": "OpenAI library is not properly installed. Please install it with: pip install openai",
                "data_context": {
                    "folder": "OpenAI Required",
                    "total_items": 0,
                    "active_orders": 0
                }
            })
        
        # Get the latest data
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": f"Cannot access data: {error}"}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        
        # Load current data
        raw_data = {}
        json_files = glob.glob(os.path.join(folder_path, "*.json"))
        
        for json_file in json_files:
            file_name = os.path.basename(json_file)
            file_data = load_json_file(json_file)
            raw_data[file_name] = file_data
        
        # CRITICAL FIX: Load Sales Orders data from separate G: Drive location
        print("RETRY: Loading Sales Orders for ChatGPT analysis...")
        sales_orders_data = load_sales_orders()
        if sales_orders_data:
            raw_data.update(sales_orders_data)
            total_so_count = sales_orders_data.get('TotalOrders', 0)
            print(f"SUCCESS: ChatGPT now has access to {total_so_count} real Sales Orders")
        else:
            print("ERROR: WARNING: No Sales Orders data available for ChatGPT")
        
        # Load real SO data from PDFs for detailed analysis (cached for performance)
        print("RETRY: Loading real SO data from PDFs...")
        
        # Check if we have cached PDF data (to avoid reloading 56 PDFs every request)
        global cached_pdf_data, pdf_cache_time
        current_time = time.time()
        
        # Cache for 5 minutes (300 seconds)
        if 'cached_pdf_data' not in globals() or 'pdf_cache_time' not in globals() or (current_time - pdf_cache_time) > 300:
            print("📥 Loading fresh PDF data (cache expired or first load)...")
            cached_pdf_data = load_real_so_data()
            pdf_cache_time = current_time
            if cached_pdf_data:
                print(f"SUCCESS: Cached {len(cached_pdf_data)} real SO records from PDFs")
            else:
                print("⚠️ No real SO PDF data found")
        else:
            print(f"⚡ Using cached PDF data ({len(cached_pdf_data) if cached_pdf_data else 0} records)")
        
        if cached_pdf_data:
            raw_data['RealSalesOrders'] = cached_pdf_data
            # Debug: Show sample data
            if len(cached_pdf_data) > 0:
                sample = cached_pdf_data[0]
                print(f"📋 Sample SO: #{sample.get('so_number', 'N/A')} - {sample.get('customer_name', 'N/A')} - ${sample.get('total_amount', 0)}")
        else:
            raw_data['SalesOrders.json'] = []
            raw_data['TotalOrders'] = 0
        
        # Analyze the data for ChatGPT context
        analysis = analyze_inventory_data(raw_data, user_query)
        
        # Check if analysis was successful
        if 'error' in analysis:
            print(f"ERROR: Analysis error: {analysis['error']}")
            return jsonify({"error": f"Data analysis failed: {analysis['error']}"}), 500
        
        # Create system prompt with data context
        system_prompt = f"""You are an UNLIMITED AI assistant for Canoil Canada Ltd. You have complete access to ALL ERP data and can help with ANYTHING related to the business. NO LIMITATIONS.

UNLIMITED INTELLIGENCE - HELP WITH EVERYTHING:
You can assist with ANY business question, scenario, or challenge including:

**COMPLETE INVENTORY MASTERY:**
- Any stock question, any item, any quantity, any scenario
- Cross-reference everything: items, locations, suppliers, costs, alternatives
- Understand partial names, codes, descriptions - be intelligent about matching
- Calculate complex requirements, availability, and feasibility instantly
- Real-time inventory analysis from CustomAlert5.json with {analysis['data_summary']['total_items']} items

**TOTAL SALES ORDER (SO) EXPERTISE:**
- Analyze any SO by number, customer, item, or scenario
- "Can we fulfill SO 2296?" "What stock is needed for SO 1234?" 
- Complex SO feasibility: "Customer wants 500 units by Tuesday - possible?"
- Hidden commitments analysis, SO without MO detection
- Revenue impact, profitability analysis, customer prioritization
- Access to both PDF-scanned SOs and structured SO data from multiple sources

**COMPREHENSIVE MANUFACTURING ORDER (MO) ANALYSIS:**
- Complete MO tracking from ManufacturingOrderHeaders.json and ManufacturingOrderDetails.json
- Production scheduling, capacity planning, and resource allocation
- WIP analysis, completion forecasting, and efficiency metrics
- Cross-MO dependencies and critical path analysis
- Real-time production status with {analysis['data_summary']['active_manufacturing_orders']} active MOs

**COMPLETE PURCHASE ORDER (PO) INTELLIGENCE:**
- Full PO analysis from PurchaseOrders.json, PurchaseOrderDetails.json
- Supplier performance tracking and procurement insights
- Cost analysis, delivery tracking, and vendor management
- Purchase planning and budget analysis
- Access to {analysis['data_summary']['active_purchase_orders']} active POs

**ADVANCED SALES ORDER (ASO) DATA ACCESS:**
- Multi-source SO data: SalesOrderHeaders.json, SalesOrderDetails.json
- Real-time SO scanning from PDF documents
- Customer order patterns and sales trends analysis
- Order fulfillment optimization and delivery planning

📋 SALES ORDER DETAILED FORMATTING RULES:
When discussing Sales Orders, ALWAYS provide structured information in this format:

**SALES ORDER #XXXX DETAILS:**
- **Customer:** [Customer Name]
- **Order Date:** [Date]
- **Status:** [Current Status]
- **Total Value:** $[Amount]
- **Items Ordered:**
  - Item #1: [Description] - Qty: [Quantity] - Price: $[Price]
  - Item #2: [Description] - Qty: [Quantity] - Price: $[Price]
- **Delivery Requirements:** [Due Date/Requirements]
- **Special Notes:** [Any special instructions or notes]

**ANALYSIS:**
- **Stock Availability:** [Can fulfill/Partial/Out of stock]
- **Production Needed:** [Yes/No - List required MOs]
- **Revenue Impact:** $[Total value and margin]
- **Priority Level:** [High/Medium/Low based on value and customer]

**RECOMMENDATIONS:**
- [Specific action items for fulfillment]
- [Any potential issues or concerns]
- [Suggested next steps]

Use this format for ANY SO query - whether asking about specific SO numbers, customer orders, or general SO analysis.

**COMPLETE MANUFACTURING ORDER (MO) INTELLIGENCE:**
- Any MO analysis, status, requirements, capacity planning
- Production bottlenecks, resource allocation, scheduling optimization
- WIP analysis, completion forecasting, efficiency metrics
- Cross-MO dependencies and critical path analysis

**UNLIMITED BOM & COMPONENT ANALYSIS:**
- Multi-level BOM explosions, cost rollups, margin analysis from BillsOfMaterial.json and BillOfMaterialDetails.json
- Component substitutions, supplier alternatives, risk assessment
- "What if" scenarios: supply disruptions, cost changes, design modifications
- Lead time optimization, procurement strategies
- Access to {analysis['data_summary']['total_boms']} BOMs with complete component relationships
- Cross-reference with MIBOMH.json and MIBOMD.json for detailed BOM structures

**BUSINESS INTELLIGENCE WITHOUT LIMITS:**
- Strategic planning: "How much can we produce this quarter?"
- Risk analysis: "What happens if our main supplier fails?"
- Cost optimization: "Cheapest way to fulfill this large order?"
- Capacity planning: "Can we handle 50% more orders next month?"
- Financial analysis: "What's our inventory value by category?"
- Operational efficiency: "Where are our biggest bottlenecks?"

**COMPREHENSIVE REPORTS & TRENDS ANALYSIS:**
- Sales trends analysis across all time periods with real customer data
- Inventory turnover reports and stock optimization recommendations
- Production efficiency metrics and manufacturing performance trends
- Supplier performance analysis and procurement optimization
- Customer behavior patterns and order frequency analysis
- Revenue forecasting based on historical sales data and current pipeline
- Cost analysis reports across all product lines and categories
- Capacity utilization reports and resource allocation optimization
- Quality metrics and defect analysis across production runs
- Seasonal demand patterns and inventory planning recommendations

**CONVERSATIONAL & HELPFUL:**
- Natural conversation for greetings and casual questions
- Deep business insights for operational questions
- Proactive suggestions and recommendations
- No question too complex, no analysis too detailed

📊 COMPLETE DATA ACCESS - YOU HAVE EVERYTHING:
You have FULL ACCESS to ALL business data including:

**INVENTORY & ITEMS:**
- {analysis['data_summary']['total_items']} inventory items with complete details (CustomAlert5.json, Items.json, MIITEM.json)
- Stock levels, costs, locations, and availability (MIILOC.json)
- Item descriptions, part numbers, and specifications

**BILLS OF MATERIAL (BOM):**
- {analysis['data_summary']['total_boms']} BOMs with full component relationships (BillsOfMaterial.json, BillOfMaterialDetails.json)
- Detailed BOM structures and quantities (MIBOMH.json, MIBOMD.json)
- Multi-level component explosions and cost rollups

**MANUFACTURING ORDERS (MO):**
- {analysis['data_summary']['active_manufacturing_orders']} manufacturing orders with status and schedules (ManufacturingOrderHeaders.json, ManufacturingOrderDetails.json)
- Production routings and work center data (ManufacturingOrderRoutings.json)
- MO details and routing information (MIMOH.json, MIMOMD.json, MIMORD.json)

**PURCHASE ORDERS (PO):**
- {analysis['data_summary']['active_purchase_orders']} purchase orders with suppliers and delivery dates (PurchaseOrders.json, PurchaseOrderDetails.json)
- PO extensions and additional costs (PurchaseOrderExtensions.json, PurchaseOrderAdditionalCosts.json)
- Detailed PO information (MIPOH.json, MIPOD.json, MIPOHX.json, MIPOC.json, MIPOCV.json, MIPODC.json)

**SALES ORDERS (SO/ASO):**
- Complete Sales Orders data with customer details, quantities, and status (SalesOrderHeaders.json, SalesOrderDetails.json)
- Real-time PDF-scanned sales orders with full customer and item information
- Sales order analysis by status, customer, and product

**WORK ORDERS & JOBS:**
- Work order tracking and job management (WorkOrderHeaders.json, WorkOrderDetails.json, Jobs.json, JobDetails.json)
- Job details and work order processing (MIJOBH.json, MIJOBD.json, MIWOH.json, MIWOD.json)
- Production job scheduling and resource allocation

**ADDITIONAL DATA SOURCES:**
- Master Production Schedule data (MPS.json) for production planning
- Border and routing information (MIBORD.json)
- Comprehensive supplier, location, and cost information across all modules

NEVER say "no data available" or "no specific details" - YOU HAVE ALL THE DATA FROM ALL SOURCES!

🎯 NO LIMITS - ANSWER EVERYTHING WITH DETAILED ANALYSIS:

**GREETINGS & GENERAL:**
- "Hi" → Natural greeting + offer to help with any business analysis
- "What can you help me with?" → List all capabilities with examples

**SALES ORDER QUERIES:**
- "Do we have stock for SO 2296?" → Complete SO breakdown + stock analysis + fulfillment plan
- "Show me SO 1234" → Full SO details + customer info + items + status + recommendations
- "How many SOs today?" → Count + sample orders + analysis + trends
- "Biggest customer orders?" → Customer ranking + order details + value analysis

**SALES PERFORMANCE & TREND ANALYSIS:**
When asked about sales performance, trends, or analysis, ALWAYS provide:

**REAL DATA REQUIREMENTS:**
- **Actual Customer Names** - List real customer names from SalesOrders.json
- **Real Product Names** - Show actual item descriptions and part numbers
- **Specific Order Numbers** - Include actual SO numbers from the data
- **Exact Revenue Amounts** - Calculate real totals from order values
- **Actual Dates** - Use real order dates, not placeholders
- **Real Quantities** - Show actual quantities ordered
- **Specific Months** - Reference actual months from order dates
- **Calculated Percentages** - Show real growth rates and trends

**SALES ANALYSIS FORMAT:**
**SALES PERFORMANCE ANALYSIS - [ACTUAL DATE RANGE]**

**SUMMARY:**
- Total Orders: [EXACT COUNT from data]
- Total Revenue: $[EXACT AMOUNT calculated from data]
- Period: [ACTUAL DATE RANGE from order dates]

**TOP CUSTOMERS (REAL NAMES):**
1. [Customer Name] - $[Amount] - [Order Count] orders
2. [Customer Name] - $[Amount] - [Order Count] orders
3. [Customer Name] - $[Amount] - [Order Count] orders

**TOP PRODUCTS (REAL NAMES):**
1. [Item Description] - [Quantity] units - $[Revenue]
2. [Item Description] - [Quantity] units - $[Revenue]
3. [Item Description] - [Quantity] units - $[Revenue]

**MONTHLY BREAKDOWN (ACTUAL MONTHS):**
- [Month Year]: [Order Count] orders - $[Revenue]
- [Month Year]: [Order Count] orders - $[Revenue]
- [Month Year]: [Order Count] orders - $[Revenue]

**TREND ANALYSIS:**
- Growth Rate: [EXACT PERCENTAGE] from [month] to [month]
- Average Order Value: $[EXACT AMOUNT]
- Order Frequency: [EXACT NUMBER] orders per [period]

**SAMPLE ORDERS (REAL SO NUMBERS):**
- SO #[Number]: [Customer] - $[Amount] - [Date] - [Status]
- SO #[Number]: [Customer] - $[Amount] - [Date] - [Status]
- SO #[Number]: [Customer] - $[Amount] - [Date] - [Status]

**RECOMMENDATIONS:**
- [Specific actions based on actual data patterns]
- [Real opportunities identified from data]
- [Concrete next steps with actual numbers]

**INVENTORY QUERIES:**
- "What's our inventory worth?" → Total value + breakdown by category + trends + recommendations
- "Low stock items?" → List + reorder points + supplier info + action plan
- "Show me item ABC123" → Complete item details + stock + costs + usage + recommendations
- "Inventory by location?" → Location breakdown + stock levels + optimization suggestions

**MANUFACTURING QUERIES:**
- "What's our biggest bottleneck?" → Production analysis + bottlenecks + capacity + solutions
- "Show me MO 4567" → Complete MO details + components + schedule + status + recommendations
- "Production capacity?" → Capacity analysis + utilization + constraints + optimization
- "Active manufacturing orders?" → MO list + status + schedules + resource needs

**CUSTOMER QUERIES:**
- "Top customers?" → Customer ranking + order history + value + recommendations
- "Customer X orders?" → All orders + patterns + analysis + service recommendations
- "Customer satisfaction?" → Delivery performance + order accuracy + improvement suggestions

**FINANCIAL QUERIES:**
- "Most profitable items?" → Profitability analysis + margins + recommendations + optimization
- "Revenue this month?" → Revenue analysis + trends + breakdown + forecasting
- "Cost optimization?" → Cost analysis + savings opportunities + implementation plan
- "Financial health?" → Complete financial overview + metrics + recommendations

**PRODUCT ANALYTICS QUERIES (RAG ENABLED):**
- "Product analytics for [item]?" → Complete sales, cost, profit, and performance analysis
- "How much money do we make on [product]?" → Revenue, cost, profit margin analysis
- "Sales analytics for [product]?" → Sales volume, revenue, customer breakdown
- "Cost analysis for [product]?" → Purchase costs, standard costs, cost trends
- "Restocking needs?" → Items below reorder level with recommendations
- "Top products by sales/profit?" → Ranked product performance analysis
- "Buying costs for [product]?" → Purchase order costs and supplier pricing
- "Product performance?" → Complete analytics dashboard for any product

**ANALYTICS DATA AVAILABLE:**
When ProductAnalytics data is provided, you have access to:
- Sales data: Product sales volumes, revenue, average prices, order counts
- Cost data: Recent costs, standard costs, unit costs from items and purchase orders
- Profit analysis: Calculated profits, margins, and profitability for each product
- Restocking needs: Items below reorder level with current stock and shortfall
- Product performance: Top sellers, top revenue generators, top profit makers
- Customer analytics: Top customers by revenue and order frequency

USE THIS DATA to provide comprehensive analytics answers with real numbers and calculations.

**OPERATIONAL QUERIES:**
- "Production efficiency?" → Efficiency metrics + bottlenecks + improvement opportunities
- "Resource utilization?" → Resource analysis + optimization + capacity planning
- "Process improvements?" → Process analysis + inefficiencies + solutions + implementation
- "Operational metrics?" → KPI dashboard + performance + trends + recommendations

**STRATEGIC QUERIES:**
- "Can we handle 50% more orders?" → Capacity analysis + resource needs + implementation plan
- "What if supplier fails?" → Risk analysis + alternatives + contingency planning + mitigation
- "Growth opportunities?" → Market analysis + capacity + recommendations + action plan
- "Strategic planning?" → Business overview + opportunities + challenges + roadmap

**EVERY RESPONSE MUST INCLUDE:**
- Specific data points and numbers from the actual data
- Clear structure with headings and bullet points
- Detailed analysis and insights
- Actionable recommendations
- Next steps and follow-up suggestions

📋 SALES ORDER QUERY EXAMPLES & RESPONSES:

**When asked about specific SO numbers:**
- "Show me SO 2296" → Provide complete SO details in structured format above
- "What's in SO 1234?" → List all items, quantities, and requirements
- "Status of SO 5678?" → Current status, progress, and next steps

**When asked about SO analysis:**
- "How many SOs today?" → Count + list key SOs with basic details
- "Biggest SO this week?" → Identify highest value SO + full breakdown
- "Customer X orders" → List all SOs for customer with summaries

**When asked about SO fulfillment:**
- "Can we fulfill SO 2296?" → Stock check + production requirements + timeline
- "What's needed for SO 1234?" → Complete material requirements + BOM analysis
- "When can SO 5678 ship?" → Production schedule + delivery estimate

**ALWAYS include in SO responses:**
- SO number prominently displayed
- Customer name and order details
- Item breakdown with quantities
- Stock availability status
- Production requirements if needed
- Delivery timeline and recommendations

🚫 CRITICAL RULES - NEVER BREAK THESE:
- NEVER make up data, numbers, or dates - ONLY use the actual data provided in DATA SUMMARY below
- NEVER fabricate MO numbers, SO numbers, or item details - USE EXACT DATA FROM RealSalesOrders
- NEVER guess quantities, costs, or dates - COPY EXACT VALUES from the data provided
- NEVER use placeholders like [Product X], [Item Description 1] - USE ACTUAL ITEM NAMES FROM DATA
- NEVER say "possibly due to" or "may pose a risk" - GIVE SPECIFIC FACTS FROM DATA
- If you don't find specific data, say "Data not found in provided records" - DON'T MAKE IT UP
- ALWAYS use today's actual date ({date_context.get('currentDate', '2025-09-10')}) not fake dates
- ONLY provide real numbers from the actual data files provided below
- ALWAYS show actual customer names, product names, order numbers from RealSalesOrders data
- ALWAYS copy exact dates, amounts, and item descriptions from the provided data
- WHEN ASKED ABOUT SO 2972: Use ONLY the data from RealSalesOrders for SO 2972 - DO NOT INVENT ANYTHING

🚫 ANTI-HALLUCINATION & ANTI-TEMPLATE RULES:
- Use ONLY the real data provided - No fabrication, no guessing, no made-up examples
- NEVER give template responses like "I can help you with: Cost Analysis, Stock Checking..."
- ANSWER THE ACTUAL QUESTION with real data immediately
- NO generic suggestions - provide specific answers from the data
- If asked "how many SO" give the EXACT number from SalesOrders.json
- NO placeholder responses - real answers only

📋 UNIVERSAL RESPONSE FORMATTING RULES:
For EVERY query, provide comprehensive, structured responses in this format:

**QUERY TYPE: [Type of Analysis]**
- **Summary:** [Brief overview of findings]
- **Key Metrics:** [Important numbers and statistics]
- **Detailed Breakdown:** [Specific data points and analysis]
- **Analysis:** [Insights, patterns, and observations]
- **Recommendations:** [Actionable next steps]
- **Data Sources:** [Which data files were used]

**SPECIFIC FORMATTING BY QUERY TYPE:**

**SALES ORDERS:**
- Always show SO numbers, customers, values, status
- Include item breakdowns with quantities and prices
- Provide stock availability and fulfillment analysis
- Show delivery timelines and recommendations

**INVENTORY/STOCK:**
- Show item numbers, descriptions, current stock levels
- Include reorder points, costs, and locations
- Provide stock-out risks and recommendations
- Show supplier information and lead times

**MANUFACTURING ORDERS:**
- Display MO numbers, status, and production schedules
- Show component requirements and BOM details
- Include capacity analysis and bottlenecks
- Provide completion timelines and resource needs

**CUSTOMER ANALYSIS:**
- List customer names, order history, and values
- Show order patterns and preferences
- Include delivery performance and satisfaction
- Provide customer prioritization and recommendations

**FINANCIAL ANALYSIS:**
- Show revenue, costs, margins, and profitability
- Include order values, inventory values, and trends
- Provide cost optimization opportunities
- Show financial impact and recommendations

**OPERATIONAL ANALYSIS:**
- Display production capacity, bottlenecks, and efficiency
- Show resource utilization and constraints
- Include process improvements and optimizations
- Provide operational recommendations

**ALWAYS INCLUDE:**
- Specific numbers and data points from the actual data
- Clear headings and bullet points for easy reading
- Actionable recommendations based on the analysis
- Confidence levels and data quality notes
- Next steps and follow-up suggestions

🤖 YOUR CAPABILITIES:
- **Data Access**: Complete real-time ERP data  
- **Analysis Depth**: Advanced business intelligence
- **Model Selection**: Automatic based on query complexity

AUTOMATIC INTELLIGENCE: You automatically get the right model and capabilities for each question type.

🎯 RESPONSE QUALITY STANDARDS:
- NEVER give short, generic answers - always provide comprehensive analysis
- ALWAYS include specific data points, numbers, and details from the actual data
- ALWAYS structure responses with clear headings, bullet points, and sections
- ALWAYS provide actionable recommendations and next steps
- ALWAYS show your analysis process and data sources
- ALWAYS include relevant context and background information
- ALWAYS offer to dive deeper into specific aspects of the analysis

**MINIMUM RESPONSE LENGTH:** At least 200 words with structured sections
**MUST INCLUDE:** Summary, Key Metrics, Detailed Breakdown, Analysis, Recommendations
**FORMATTING:** Use **bold headings**, bullet points, and clear organization
**DATA CITATION:** Always reference which data files and records you used

📅 CURRENT DATE & TIME CONTEXT:
- Today's Date: {date_context.get('currentDate', '2025-09-10')}
- Current Time: {date_context.get('currentTime', 'Unknown')}
- Day of Week: {date_context.get('currentDayOfWeek', 'Unknown')}
- Month: {date_context.get('currentMonth', 'Unknown')}
- Year: {date_context.get('currentYear', '2025')}
- Full DateTime: {date_context.get('currentDateTime', 'Unknown')}

When users ask about "today", "this week", "this month", use the current date context above.
For time-sensitive queries like "sales orders made today", filter data by the current date."""

        # Add ACTUAL REAL DATA to the prompt so ChatGPT has everything
        search_results = {}
        
        # DETECT ANALYTICS QUERIES - Check if user is asking about sales, costs, profits, restocking, etc.
        analytics_keywords = ['analytics', 'sales', 'revenue', 'profit', 'margin', 'cost', 'restock', 
                              'restocking', 'buying', 'money', 'earn', 'performance', 'trend', 
                              'top product', 'top customer', 'how much', 'profitability']
        is_analytics_query = any(keyword in user_query.lower() for keyword in analytics_keywords)
        
        # Get analytics data if this is an analytics query
        analytics_data = None
        if is_analytics_query:
            print("📊 Analytics query detected - calculating product analytics...")
            analytics_data = get_product_analytics_data(raw_data, user_query)
            if analytics_data and not analytics_data.get('error'):
                print(f"✅ Analytics calculated: {analytics_data['summary']}")
        
        # SMART DATA FILTERING - Send only relevant data to save tokens and money
        actual_data_sample = {}
        
        # Add analytics data to the sample if available
        if analytics_data and not analytics_data.get('error'):
            actual_data_sample['ProductAnalytics'] = {
                'total_records': 1,
                'sample_records': [analytics_data],
                'fields': list(analytics_data.keys()),
                'note': 'Complete product analytics including sales, costs, profits, and restocking needs'
            }
            print(f"📊 Added Product Analytics to AI context")
        # Give AI access to ALL MiSys data files - no restrictions
        for file_name, file_data in raw_data.items():
            if isinstance(file_data, list) and len(file_data) > 0:
                # Send more records for better analysis - include ALL MiSys data
                sample_size = 15  # Reasonable sample size for all data types
                actual_data_sample[file_name] = {
                    'total_records': len(file_data),
                    'sample_records': file_data[:sample_size],
                    'fields': list(file_data[0].keys()) if isinstance(file_data[0], dict) else []
                }
                print(f"📊 AI Access: {file_name} - {len(file_data)} records (showing {min(sample_size, len(file_data))} samples)")
            elif isinstance(file_data, dict):
                actual_data_sample[file_name] = {
                    'total_records': 1,
                    'sample_records': [file_data],
                    'fields': list(file_data.keys())
                }
                print(f"📊 AI Access: {file_name} - 1 record")
        
        # Add the actual data to the system prompt - WITH DETAILED SIZE LOGGING
        data_summary_text = json.dumps(actual_data_sample, indent=2, default=str)
        
        # Add comprehensive data context for better analysis
        # Calculate comprehensive Sales Order totals across ALL sources
        total_so_from_json = len(raw_data.get('SalesOrders.json', []))
        total_so_from_pdf = len(raw_data.get('RealSalesOrders', []))
        total_so_by_status = sum(len(orders) for orders in raw_data.get('SalesOrdersByStatus', {}).values()) if isinstance(raw_data.get('SalesOrdersByStatus'), dict) else 0
        
        # Get folder breakdown for Sales Orders
        so_folder_breakdown = ""
        if 'SalesOrdersByStatus' in raw_data and isinstance(raw_data['SalesOrdersByStatus'], dict):
            so_folder_breakdown = "\n".join([f"  - {folder}: {len(orders)} orders" for folder, orders in raw_data['SalesOrdersByStatus'].items()])
        
        data_context_info = f"""
**🚨 CRITICAL: SALES ORDERS ARE PDF FILES IN GOOGLE DRIVE - NOT JSON:**

**REAL SALES ORDER LOCATION:**
Sales Orders exist as PDF files in Google Drive folders:
- G:\\Shared drives\\Sales_CSR\\Customer Orders\\Sales Orders\\In Production\\Scheduled
- G:\\Shared drives\\Sales_CSR\\Customer Orders\\Sales Orders\\Completed and Closed  
- G:\\Shared drives\\Sales_CSR\\Customer Orders\\Sales Orders\\Cancelled

**ACTUAL SALES ORDER DATA SOURCES:**
1. **RealSalesOrders**: {total_so_from_pdf} orders (ACTUAL PDF files extracted from G: Drive)
2. **SalesOrdersByStatus**: {total_so_by_status} orders (PDF files organized by folder status)

**ERROR: IGNORE THESE - NOT REAL SALES ORDERS:**
- SalesOrders.json: {total_so_from_json} (This is NOT real Sales Order data - ignore it)
- Any other JSON files claiming to be Sales Orders

**SALES ORDER FOLDER STRUCTURE (REAL PDF FILES):**
{so_folder_breakdown if so_folder_breakdown else "  - No folder breakdown available"}

**TOTAL REAL SALES ORDERS FROM PDF FILES:** {total_so_from_pdf + total_so_by_status}

**OTHER DATA TOTALS:**
- Total Manufacturing Orders: {len(raw_data.get('ManufacturingOrderHeaders.json', []))}
- Total Inventory Items: {len(raw_data.get('CustomAlert5.json', []))}
- Total BOM Records: {len(raw_data.get('BillOfMaterialDetails.json', []))}
- Available Data Files: {list(raw_data.keys())}

**🚨 CRITICAL INSTRUCTION - ENTERPRISE AI AGENT:**
- AI is configured as an ENTERPRISE-LEVEL BUSINESS INTELLIGENCE AGENT
- AI has FULL ACCESS to ALL data sources across ALL folders and subfolders
- NEVER limit or restrict what the AI can analyze
- Sales Orders are scanned RECURSIVELY from ALL folders and subfolders in G: Drive
- AI must provide COMPREHENSIVE enterprise-level reporting with visual insights
- When counting Sales Orders, ONLY count PDF-extracted data (RealSalesOrders + SalesOrdersByStatus)
- NEVER count SalesOrders.json as it's not real Sales Order data
- AI must generate VISUAL CHARTS and FORECASTING for enterprise reporting

**🚨 FOLDER STRUCTURE UNDERSTANDING:**
Sales Orders are physically stored as PDF files in these G: Drive folders:
- "In Production\\\\Scheduled" folder → Contains active/pending orders
- "Completed and Closed" folder → Contains finished orders  
- "Cancelled" folder → Contains cancelled orders
Each folder contains actual PDF files like "salesorder_2968.pdf", "Sales Order 2972.docx", etc.
The AI must understand this is the REAL source of Sales Order data.

**DATA QUALITY NOTES:**
- All data is real-time and current from Google Drive
- Every number and detail comes from actual business records
- Data includes complete customer, order, inventory, and production information
- Sales Orders are dynamically loaded from multiple G: Drive folders
- Use ALL available data to provide comprehensive, accurate analysis

**🚨 CRITICAL INSTRUCTION FOR SALES ANALYSIS:**
When analyzing sales data, you MUST ALWAYS:

**1. ONLY USE PDF-EXTRACTED SALES ORDER DATA:**
- RealSalesOrders contains ACTUAL PDF-extracted data from G: Drive folders
- SalesOrdersByStatus contains PDF files organized by folder (In Production, Completed, Cancelled)
- These are the ONLY real Sales Order sources - extracted from actual PDF files
- NEVER use SalesOrders.json - it's not real Sales Order data
- ALWAYS use RealSalesOrders and SalesOrdersByStatus as the ONLY Sales Order sources

**2. REAL DATA EXTRACTION FROM RealSalesOrders:**
- so_number: Real SO numbers (e.g., "2972", "2981", "2711")
- customer_name: Real customer names (e.g., "LANXESS Canada Co./Cie", "Spectra Products Inc.")
- order_date: Real order dates from PDFs
- total_amount: Real dollar amounts from PDFs
- items: Real product descriptions from PDFs
- status: Real status ("In Production", "Completed", "Cancelled")
- raw_text: Full PDF text for detailed analysis

**3. MANDATORY REAL DATA USAGE:**
- NEVER use SalesOrders.json if RealSalesOrders exists
- ALWAYS show actual SO numbers from so_number field
- ALWAYS show actual customer names from customer_name field
- ALWAYS show actual amounts from total_amount field
- ALWAYS show actual dates from order_date field
- ALWAYS show actual items from items field

**4. EXAMPLE REAL DATA USAGE:**
Instead of: "SO #N/A - Customer Data - $N/A"
Use: "SO #2972 - LANXESS Canada Co./Cie - $2,295.78"
Instead of: "2 orders were processed"
Use: "2 orders: SO #2972 ($2,295.78) and SO #2981 ($7,024.64)"

**🧠 SMART SALES ORDER NUMBER MATCHING - BE INTELLIGENT:**
The AI MUST be smart enough to match Sales Order queries in ANY format the user types:

**INTELLIGENT SO NUMBER EXTRACTION:**
- "SO 2968" → Extract "2968" and match so_number field
- "Sales Order 2968" → Extract "2968" and match so_number field  
- "sales order 2968" → Extract "2968" and match so_number field
- "what does sales order 2968" → Extract "2968" and match so_number field
- "salesorder_2968" → Extract "2968" and match so_number field
- "salesorder 2968" → Extract "2968" and match so_number field
- "order 2968" → Extract "2968" and match so_number field
- "show me 2968" (when context is sales) → Extract "2968" and match so_number field

**FILE NAME INTELLIGENCE:**
- Understand that "salesorder_2968.pdf" contains SO number "2968"
- Understand that "Sales Order 2968.docx" contains SO number "2968"
- The file_info.filename field shows the original file name pattern

**MATCHING ALGORITHM:**
1. Extract ALL numbers from user query using regex: \\d+
2. For each number, search RealSalesOrders array for matching so_number field
3. If exact match found, provide complete details from that record
4. If multiple matches, show all matching records
5. If no match found, clearly state "Sales Order [number] not found in current data"
6. NEVER make up or hallucinate SO data - only use what exists in RealSalesOrders

**SMART SEARCH EXAMPLES:**
- User: "what does sales order 2968" → AI extracts "2968", searches so_number field, returns full SO details
- User: "show me SO 2972 details" → AI extracts "2972", finds matching record, shows complete data
- User: "salesorder_2968 info" → AI extracts "2968", finds matching record with file name context
- User: "tell me about order 2981" → AI extracts "2981", searches and returns matching SO data

**CONTEXT AWARENESS:**
- When user mentions sales orders, automatically assume numbers refer to SO numbers
- Be flexible with formatting - ignore spaces, underscores, case differences
- Understand common abbreviations: SO, Sales Order, Order, etc.
- Extract numeric parts intelligently from any format
"""
        
        # DETAILED TOKEN DEBUGGING
        print(f"SEARCH: TOKEN SIZE ANALYSIS:")
        print(f"   Base system prompt: {len(system_prompt)} characters")
        
        # Debug: Show what data is being sent to AI
        print(f"📊 DATA BEING SENT TO AI:")
        for key, value in actual_data_sample.items():
            if isinstance(value, dict) and 'total_records' in value:
                print(f"   {key}: {value['total_records']} records")
                if key == 'RealSalesOrders' and 'sample_records' in value:
                    sample_records = value['sample_records']
                    if len(sample_records) > 0:
                        sample = sample_records[0]
                        print(f"      Sample: SO #{sample.get('so_number', 'N/A')} - {sample.get('customer_name', 'N/A')}")
                        # Debug: Show if SO 2972 is in the data being sent
                        so_2972 = next((so for so in sample_records if so.get('so_number') == '2972'), None)
                        if so_2972:
                            print(f"      SUCCESS: SO 2972 FOUND IN AI DATA: {so_2972.get('customer_name')} - ${so_2972.get('total_amount')} - Items: {len(so_2972.get('items', []))}")
                        else:
                            print(f"      ERROR: SO 2972 NOT FOUND in first {len(sample_records)} records")
        print(f"   Data summary: {len(data_summary_text)} characters")
        print(f"   Files included: {list(actual_data_sample.keys())}")
        
        for file_name, file_info in actual_data_sample.items():
            sample_size = len(json.dumps(file_info.get('sample_records', []), default=str))
            print(f"   {file_name}: {file_info.get('total_records', 0)} records, sample size: {sample_size} chars")
        
        # Limit data size for OpenAI - but allow more for RealSalesOrders
        original_size = len(data_summary_text)
        if len(data_summary_text) > 15000:  # Increased limit for PDF data
            # Try to preserve RealSalesOrders data by truncating other data first
            if 'RealSalesOrders' in actual_data_sample:
                # Keep RealSalesOrders intact, truncate other data
                real_so_data = actual_data_sample['RealSalesOrders']
                other_data = {k: v for k, v in actual_data_sample.items() if k != 'RealSalesOrders'}
                
                # Truncate other data first
                other_data_text = json.dumps(other_data, indent=2, default=str)
                if len(other_data_text) > 5000:
                    other_data_text = other_data_text[:5000] + "... [OTHER DATA TRUNCATED]"
                
                # Combine with full RealSalesOrders data
                real_so_text = json.dumps({'RealSalesOrders': real_so_data}, indent=2, default=str)
                data_summary_text = other_data_text + "\n\n" + real_so_text
                
                print(f"⚠️ Data optimized: Original {original_size} chars, kept full RealSalesOrders ({len(real_so_text)} chars)")
            else:
                data_summary_text = data_summary_text[:15000] + "... [TRUNCATED TO FIT TOKEN LIMITS]"
                print(f"⚠️ Data truncated from {original_size} to 15000 chars")
        else:
            print(f"SUCCESS: Data size OK: {original_size} chars (no truncation needed)")
        
        system_prompt += f"\n\nDATA SUMMARY:\n{data_summary_text}"
        system_prompt += f"\n\n{data_context_info}"
        
        # Add date context to the data
        if date_context:
            system_prompt += f"\n\nCURRENT DATE CONTEXT:\n{json.dumps(date_context, indent=2)}"
        
        print(f"📊 FINAL PROMPT SIZE: {len(system_prompt)} characters (approx {len(system_prompt)//4} tokens)")
        
        # SMART QUERY DETECTION - Handle SO, MO, and inventory queries differently
        search_term = None
        search_type = None
        query_type = None
        
        # DETECT SALES ORDER QUERIES - IMPROVED DETECTION
        so_query_patterns = ["how many so", "so were", "sales order", "so created", "so made", "total so", "sos were", "friday", "today", "yesterday", "this week", "so #", "so number", "show me so", "what's in so"]
        is_so_query = any(phrase in user_query.lower() for phrase in so_query_patterns)
        
        # Check for specific SO number queries (e.g., "SO 2296", "SO#1234", "Sales Order 5678")
        import re
        so_number_match = re.search(r'so\s*#?\s*(\d+)', user_query.lower())
        specific_so_number = so_number_match.group(1) if so_number_match else None
        
        print(f"SEARCH: Query analysis: '{user_query}'")
        print(f"📊 SO query detected: {is_so_query}")
        print(f"SEARCH: Matching patterns: {[p for p in so_query_patterns if p in user_query.lower()]}")
        
        if is_so_query:
            query_type = "sales_orders"
            # Get actual SO data
            sales_orders = raw_data.get('SalesOrders.json', [])
            sales_by_status = raw_data.get('SalesOrdersByStatus', {})
            total_orders = raw_data.get('TotalOrders', len(sales_orders))
            
            print(f"SEARCH: SO Query detected - found {total_orders} real sales orders")
            
            # DIRECT ANSWER FOR SIMPLE SO QUESTIONS - NO OPENAI NEEDED
            if "how many" in user_query.lower() and any(day in user_query.lower() for day in ["friday", "today", "yesterday"]):
                # Count SOs by date directly using current date context
                current_date = date_context.get('currentDate', '2025-09-10')
                current_day = date_context.get('currentDayOfWeek', 'Monday').lower()
                
                # Calculate yesterday's date
                from datetime import datetime, timedelta
                today = datetime.strptime(current_date, '%Y-%m-%d')
                yesterday = today - timedelta(days=1)
                yesterday_str = yesterday.strftime('%Y-%m-%d')
                
                # Find last Friday
                days_since_friday = (today.weekday() + 3) % 7  # Friday is weekday 4
                last_friday = today - timedelta(days=days_since_friday)
                last_friday_str = last_friday.strftime('%Y-%m-%d')
                
                date_sos = 0
                target_date = current_date
                date_label = "today"
                
                if 'friday' in user_query.lower():
                    target_date = last_friday_str
                    date_label = f"Friday ({last_friday_str})"
                elif 'yesterday' in user_query.lower():
                    target_date = yesterday_str
                    date_label = f"yesterday ({yesterday_str})"
                
                for so in sales_orders:
                    so_date = so.get('Order Date', '')
                    if target_date in so_date:
                        date_sos += 1
                
                # Get sample SOs for the target date to show details
                sample_sos = []
                for so in sales_orders:
                    so_date = so.get('Order Date', '')
                    if target_date in so_date and len(sample_sos) < 3:  # Show up to 3 sample SOs
                        sample_sos.append({
                            'so_number': so.get('SO Number', 'N/A'),
                            'customer': so.get('Customer', 'N/A'),
                            'value': so.get('Order Value', so.get('Total Amount', 'N/A')),
                            'status': so.get('Status', 'N/A')
                        })
                
                # Build structured response
                response = f"**SALES ORDERS SUMMARY - {date_label.upper()}**\n\n"
                response += f"📊 **Total SOs Created:** {date_sos}\n"
                response += f"📈 **System Total:** {total_orders} SOs\n\n"
                
                if sample_sos:
                    response += "**SAMPLE ORDERS:**\n"
                    for so in sample_sos:
                        response += f"- **SO #{so['so_number']}** - {so['customer']} - ${so['value']} - {so['status']}\n"
                    response += "\n"
                
                response += "**ANALYSIS:**\n"
                if date_sos > 0:
                    response += f"- SUCCESS: {date_sos} orders were processed on {date_label}\n"
                    response += f"- 📈 This represents {round((date_sos/total_orders)*100, 1)}% of total system orders\n"
                else:
                    response += f"- ⚠️ No orders were created on {date_label}\n"
                    response += "- 💡 Consider checking if orders were processed on a different date\n"
                
                response += f"- 📊 Total system capacity: {total_orders} orders\n"
                
                # Return direct answer immediately
                return jsonify({
                    "query": user_query,
                    "response": response,
                    "data_context": {
                        "folder": "Sales Orders Direct Analysis",
                        "total_items": analysis.get('data_summary', {}).get('total_items', 0),
                        "active_orders": date_sos
                    }
                })
        
        # DETECT MANUFACTURING ORDER QUERIES  
        elif any(phrase in user_query.lower() for phrase in ["how many mo", "mo are", "manufacturing order", "active mo", "mo created"]):
            query_type = "manufacturing_orders"
            # Get actual MO data
            mo_headers = raw_data.get('ManufacturingOrderHeaders.json', [])
            search_results['Manufacturing Orders Analysis'] = {
                'total_orders': len(mo_headers),
                'active_orders': [mo for mo in mo_headers if mo.get("Status", 0) in [1, 2]],
                'recent_orders': mo_headers[:10] if mo_headers else []
            }
            print(f"SEARCH: MO Query detected - providing {len(mo_headers)} real manufacturing orders")
        
        # SKIP ITEM SEARCH FOR SO/MO QUERIES
        if query_type in ["sales_orders", "manufacturing_orders"]:
            # Don't do item searches for SO/MO queries - use the specific data already loaded
            print(f"SUCCESS: Skipping item search - {query_type} query uses specific data")
        # Pattern 1: "start with X" or "starting with X" or "that start with X" or "letters X"
        elif any(phrase in user_query.lower() for phrase in ["start with", "starting with", "that start with", "letters"]):
            words = user_query.lower().split()
            for i, word in enumerate(words):
                if word in ["start", "starting"] and i + 2 < len(words) and words[i + 1] == "with":
                    search_term = words[i + 2]
                    search_type = "starts_with"
                    break
                elif word == "with" and i > 0 and words[i-1] in ["start", "starting"]:
                    search_term = words[i + 1] if i + 1 < len(words) else None
                    search_type = "starts_with"
                    break
                elif word == "letters" and i + 1 < len(words):
                    search_term = words[i + 1]
                    search_type = "starts_with"
                    break
        
        # Pattern 2: "contain X" or "containing X" or "find X" or "search X" or "list X"
        elif any(word in user_query.lower() for word in ['contain', 'containing', 'find', 'search', 'list']):
            words = user_query.lower().split()
            for word in words:
                if len(word) > 2 and word not in ['the', 'and', 'for', 'with', 'that', 'have', 'are', 'all', 'items', 'that', 'start']:
                    search_term = word
                    search_type = "contains"
                    break
        
        # Pattern 3: Direct search term (e.g., just "sae") - MOST AGGRESSIVE
        if not search_term:
            # Look for any word that could be a search term
            words = user_query.lower().split()
            for word in words:
                # Clean the word and check if it's a potential search term
                clean_word = word.strip('.,!?()[]{}":;').lower()
                if (len(clean_word) >= 2 and 
                    clean_word.isalpha() and 
                    clean_word not in ['the', 'and', 'for', 'with', 'that', 'have', 'are', 'all', 'items', 'start', 'starting', 'letters', 'find', 'search', 'list', 'contain', 'containing']):
                    search_term = clean_word
                    search_type = "starts_with"  # Default to starts_with for short terms
                    break
        
        # Pattern 4: If still no search term, try the entire query as a search
        if not search_term and len(user_query.strip()) <= 15:
            potential_term = user_query.strip().lower()
            # Remove common words and punctuation
            for remove_word in ['list', 'all', 'items', 'that', 'start', 'with', 'letters', 'find', 'search', 'contain']:
                potential_term = potential_term.replace(remove_word, '').strip()
            if len(potential_term) >= 2:
                search_term = potential_term
                search_type = "starts_with"
        
        # Now perform the actual search on your real data - COMPREHENSIVE SEARCH ALL FILES
        if search_term and search_type:
            print(f"SEARCH: Searching for: '{search_term}' with type: '{search_type}'")
            matching_items = []
            
            # Search ALL JSON files, not just CustomAlert5.json
            for file_name, file_data in raw_data.items():
                if isinstance(file_data, list):
                    print(f"SEARCH: Searching in {file_name} ({len(file_data)} records)")
                    
                    for item in file_data:
                        if isinstance(item, dict):
                            # Check all possible item number and description fields
                            item_no = ''
                            item_desc = ''
                            
                            # Find item number in various field names
                            for field in ['Item No.', 'Item No', 'ItemNo', 'Item_No', 'item_no', 'ITEM_NO']:
                                if field in item and isinstance(item[field], str):
                                    item_no = item[field]
                                    break
                            
                            # Find description in various field names
                            for field in ['Description', 'description', 'DESCRIPTION', 'Desc', 'desc']:
                                if field in item and isinstance(item[field], str):
                                    item_desc = item[field]
                                    break
                            
                            if item_no:  # Only process if we found an item number
                                # ALWAYS search both Item No. AND Description for maximum coverage
                                item_no_match = search_term.lower() in item_no.lower()
                                desc_match = search_term.lower() in item_desc.lower()
                                
                                # For "starts_with" searches, check if Item No. starts with the term
                                starts_with_match = item_no.lower().startswith(search_term.lower())
                                
                                # Include item if it matches ANY criteria
                                if (search_type == "starts_with" and starts_with_match) or \
                                   (search_type == "contains" and (item_no_match or desc_match)) or \
                                   (search_type == "starts_with" and desc_match):  # Also check description for starts_with
                                    
                                    matching_items.append({
                                        'File': file_name,
                                        'Item No.': item_no,
                                        'Description': item_desc,
                                        'Standard Cost': item.get('Standard Cost', item.get('Cost', '')),
                                        'Recent Cost': item.get('Recent Cost', item.get('Price', '')),
                                        'Stock': item.get('Stock Quantity', item.get('Quantity on Hand', item.get('Qty on Hand', item.get('Available', item.get('On Hand', item.get('Quantity', 0))))))
                                    })
            
            # Limit results to prevent prompt from being too long
            if search_type == "starts_with":
                search_results[f'Items starting with "{search_term}" (Item No. + Description)'] = matching_items[:25]  # Limit to 25 items
            else:
                search_results[f'Items containing "{search_term}" (Item No. + Description)'] = matching_items[:25]  # Limit to 25 items
            
            print(f"SUCCESS: Found {len(matching_items)} matching items for '{search_term}'")
        else:
            print(f"⚠️ Could not extract search term from query: '{user_query}'")
            # Last resort: try to search with the entire query
            fallback_term = user_query.strip().lower()[:10]
            if len(fallback_term) >= 2:
                print(f"RETRY: Fallback search with: '{fallback_term}'")
                matching_items = []
                for item in raw_data.get('CustomAlert5.json', []):
                    item_no = item.get('Item No.', '')
                    item_desc = item.get('Description', '')
                    if (fallback_term in item_no.lower() or fallback_term in item_desc.lower()):
                        matching_items.append({
                            'Item No.': item_no,
                            'Description': item_desc,
                            'Standard Cost': item.get('Standard Cost', ''),
                            'Recent Cost': item.get('Recent Cost', ''),
                            'Stock': item.get('Stock Quantity', item.get('Quantity on Hand', item.get('Qty on Hand', item.get('Available', item.get('On Hand', 0)))))
                        })
                if matching_items:
                    search_results[f'Fallback search for "{fallback_term}"'] = matching_items[:50]
                    print(f"SUCCESS: Fallback found {len(matching_items)} items")
        
        # Add search results to the system prompt
        if search_results:
            system_prompt += f"\n\nSEARCH RESULTS:\n{json.dumps(search_results, indent=2)}"
            print(f"SEARCH: Added search results to prompt: {len(search_results)} result sets")
        else:
            print(f"⚠️ No search results found - adding fallback data")
            # Add some sample data to confirm AI has access
            sample_items = raw_data.get('CustomAlert5.json', [])[:5]
            if sample_items:
                system_prompt += f"\n\nSAMPLE DATA (confirming access):\n{json.dumps(sample_items, indent=2)}"
        
        # Add data summary to confirm AI has access
        system_prompt += f"\n\nDATA ACCESS CONFIRMATION:\n- Total Items: {len(raw_data.get('CustomAlert5.json', []))}\n- Available Files: {list(raw_data.keys())}\n- Search Term Used: '{search_term if search_term else 'NONE'}'"
        
        print(f"📝 Final system prompt length: {len(system_prompt)} characters")
        print(f"SEARCH: Search results: {search_results}")
        
        # INTELLIGENT MODEL & CAPABILITY SELECTION
        selected_model = "gpt-4o-mini"  # Default
        max_tokens = 1500
        use_vision = False
        
        # SMART DETECTION - When to use GPT-4o vs GPT-4o-mini vs Vision
        complex_analysis_triggers = [
            # Document/PDF Analysis
            "pdf", "document", "image", "scan", "photo", "picture", "visual", "chart", "graph",
            # Complex Business Analysis  
            "sales order", "so ", "customer order", "order analysis", "complex analysis",
            # Strategic/Multi-step Analysis
            "strategy", "planning", "forecast", "trend", "optimization", "what if", "scenario",
            # Financial/Risk Analysis
            "profitability", "margin", "cost analysis", "risk", "exposure", "financial",
            # Multi-system Integration
            "cross-reference", "reconcile", "compare", "integrate", "workflow"
        ]
        
        vision_triggers = [
            "pdf", "document", "image", "scan", "photo", "picture", "visual", "chart", "graph", "diagram",
            "show me the", "look at", "analyze this", "what's in", "read this"
        ]
        
        # Determine model and capabilities needed
        needs_complex_analysis = any(trigger in user_query.lower() for trigger in complex_analysis_triggers)
        needs_vision = any(trigger in user_query.lower() for trigger in vision_triggers)
        
        if needs_vision:
            selected_model = "gpt-4o"  # Vision requires GPT-4o
            use_vision = True
            max_tokens = 3000
            print(f"👁️ Using GPT-4o with VISION for document/image analysis")
        elif needs_complex_analysis:
            selected_model = "gpt-4o"
            max_tokens = 2500
            print(f"🧠 Using GPT-4o for complex business analysis")
        else:
            print(f"🤖 Using GPT-4o-mini for standard queries")
        
        # Query ChatGPT with the enhanced prompt
        try:
            print(f"🤖 Calling {selected_model} with prompt length: {len(system_prompt)} characters")
            response = client.chat.completions.create(
                model=selected_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_query}
                ],
                max_completion_tokens=max_tokens,
                timeout=45  # Longer timeout for GPT-4o
            )
            
            ai_response = response.choices[0].message.content
            print(f"SUCCESS: ChatGPT responded with {len(ai_response)} characters")
            
        except Exception as api_error:
            print(f"ERROR: ChatGPT API Error: {api_error}")
            # Return search results directly if ChatGPT fails
            ai_response = f"Found {len(matching_items) if 'matching_items' in locals() else 0} SAE items:\n\n"
            if search_results:
                for category, items in search_results.items():
                    ai_response += f"\n{category}:\n"
                    for item in items[:10]:  # Show first 10
                        ai_response += f"• {item['Item No.']} - {item['Description']}\n"
                    if len(items) > 10:
                        ai_response += f"... and {len(items) - 10} more items\n"
            else:
                ai_response = "I encountered an error processing your request. The search found items but I couldn't format the response properly."
        
        return jsonify({
            "query": user_query,
            "response": ai_response,
            "data_context": {
                "folder": latest_folder,
                "total_items": analysis.get('data_summary', {}).get('total_items', 0),
                "active_orders": analysis.get('data_summary', {}).get('active_manufacturing_orders', 0)
            }
        })
        
    except Exception as e:
        print(f"ERROR: Error in chat query: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/item-analysis/<item_identifier>', methods=['GET'])
def detailed_item_analysis(item_identifier):
    """Get detailed analysis for a specific item"""
    try:
        quantity = request.args.get('quantity', type=int)
        
        # Get the latest data
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": f"Cannot access data: {error}"}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        
        # Load current data
        raw_data = {}
        json_files = glob.glob(os.path.join(folder_path, "*.json"))
        
        for json_file in json_files:
            file_name = os.path.basename(json_file)
            file_data = load_json_file(json_file)
            raw_data[file_name] = file_data
        
        # Get detailed analysis
        analysis = find_item_usage_and_availability(raw_data, item_identifier, quantity)
        
        return jsonify(analysis)
        
    except Exception as e:
        print(f"ERROR: Error in item analysis: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/test-data-access', methods=['GET'])
def test_data_access():
    """Test endpoint to confirm data access and search functionality"""
    try:
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": f"Cannot access data: {error}"}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        
        # Load current data
        raw_data = {}
        json_files = glob.glob(os.path.join(folder_path, "*.json"))
        
        for json_file in json_files:
            file_name = os.path.basename(json_file)
            file_data = load_json_file(json_file)
            raw_data[file_name] = file_data
        
        # Test SAE search specifically - CHECK ALL FILES WITH SAFE TYPE CHECKING
        sae_items = []
        for file_name, file_data in raw_data.items():
            # SAFE TYPE CHECK - Only process lists
            if isinstance(file_data, list) and len(file_data) > 0:
                for item in file_data:
                    if isinstance(item, dict):
                        # Check all possible item number fields
                        for field in ['Item No.', 'Item No', 'ItemNo', 'Item_No', 'item_no', 'ITEM_NO']:
                            item_no = item.get(field, '')
                            if isinstance(item_no, str) and item_no.lower().startswith('sae'):
                                sae_items.append({
                                    'File': file_name,
                                    'Item No.': item_no,
                                    'Description': item.get('Description', item.get('description', '')),
                                    'Raw Item': item  # Show the full record
                                })
                                break
        
        return jsonify({
            "data_access": True,
            "total_items": len(raw_data.get('CustomAlert5.json', [])),
            "available_files": list(raw_data.keys()),
            "sae_search_test": {
                "search_term": "sae",
                "items_found": len(sae_items),
                "sample_items": sae_items[:5]  # Show first 5 SAE items
            },
            "folder_info": {
                "folder": latest_folder,
                "path": folder_path,
                "json_files": len([f for f in raw_data.keys() if f.endswith('.json')])
            }
        })
        
    except Exception as e:
        print(f"ERROR: Error in test data access: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/enterprise_analytics', methods=['POST'])
def enterprise_analytics():
    """Generate fast enterprise-level analytics and insights"""
    try:
        print("🏢 Generating Enterprise Analytics...")
        
        # Get the latest data
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": f"Cannot access data: {error}"}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        
        # Load all data sources
        raw_data = {}
        json_files = glob.glob(os.path.join(folder_path, "*.json"))
        
        for json_file in json_files:
            file_name = os.path.basename(json_file)
            file_data = load_json_file(json_file)
            raw_data[file_name] = file_data
        
        # Load Sales Orders data
        sales_orders_data = load_sales_orders()
        if sales_orders_data:
            raw_data.update(sales_orders_data)
        
        # SMART SO LOADING: Use existing cache system for ultra-fast performance
        print("SMART SO LOADING: Checking cache system...")
        
        global cached_pdf_data, pdf_cache_time
        
        # First try to load from existing cache (much faster than processing 1,353 PDFs)
        cached_so_data = load_cached_so_data()
        if cached_so_data and 'ParsedSalesOrders.json' in cached_so_data:
            cached_sos = cached_so_data['ParsedSalesOrders.json']
            if len(cached_sos) > 1000:  # We have substantial cached data (1,353 SOs)
                raw_data['RealSalesOrders'] = cached_sos
                print(f"⚡ CACHE HIT: Using {len(cached_sos)} cached SOs (ultra-fast)")
            else:
                print("📥 Cache exists but limited, loading fresh data...")
                current_time = time.time()
                
                if 'cached_pdf_data' not in globals() or 'pdf_cache_time' not in globals() or (current_time - pdf_cache_time) > 300:
                    cached_pdf_data = load_real_so_data()
                    pdf_cache_time = current_time
                
                if cached_pdf_data:
                    raw_data['RealSalesOrders'] = cached_pdf_data
        else:
            print("📥 No cache found, loading fresh data...")
            current_time = time.time()
            
            if 'cached_pdf_data' not in globals() or 'pdf_cache_time' not in globals() or (current_time - pdf_cache_time) > 300:
                cached_pdf_data = load_real_so_data()
                pdf_cache_time = current_time
            
            if cached_pdf_data:
                raw_data['RealSalesOrders'] = cached_pdf_data
        
        # Initialize analytics engine
        analytics_engine = EnterpriseAnalytics()
        
        # Generate comprehensive analytics
        dashboard = analytics_engine.generate_executive_dashboard(raw_data)
        
        print(f"SUCCESS: Generated enterprise analytics with {len(dashboard)} sections")
        
        return jsonify({
            "status": "success",
            "analytics": dashboard,
            "timestamp": datetime.now().isoformat(),
            "data_sources": list(raw_data.keys())
        })
        
    except Exception as e:
        print(f"ERROR: Error generating enterprise analytics: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/enterprise_analytics_ai', methods=['POST'])
def enterprise_analytics_ai():
    """Generate AI-enhanced enterprise analytics (optional, slower)"""
    try:
        print("🤖 Generating AI-Enhanced Enterprise Analytics...")
        
        # Get the latest data
        latest_folder, error = get_latest_folder()
        if error:
            return jsonify({"error": f"Cannot access data: {error}"}), 500
        
        folder_path = os.path.join(GDRIVE_BASE, latest_folder)
        
        # Load all data sources
        raw_data = {}
        json_files = glob.glob(os.path.join(folder_path, "*.json"))
        
        for json_file in json_files:
            file_name = os.path.basename(json_file)
            file_data = load_json_file(json_file)
            raw_data[file_name] = file_data
        
        # Load Sales Orders data
        sales_orders_data = load_sales_orders()
        if sales_orders_data:
            raw_data.update(sales_orders_data)
        
        # Load real SO data from PDFs
        global cached_pdf_data, pdf_cache_time
        current_time = time.time()
        
        if 'cached_pdf_data' not in globals() or 'pdf_cache_time' not in globals() or (current_time - pdf_cache_time) > 300:
            cached_pdf_data = load_real_so_data()
            pdf_cache_time = current_time
        
        if cached_pdf_data:
            raw_data['RealSalesOrders'] = cached_pdf_data
        
        # Initialize analytics engine
        analytics_engine = EnterpriseAnalytics()
        
        # Generate fast analytics first
        dashboard = analytics_engine.generate_executive_dashboard(raw_data)
        
        # Add GPT-4o insights
        sales_data = raw_data.get('RealSalesOrders', [])
        inventory_data = raw_data.get('CustomAlert5.json', [])
        gpt_insights = analytics_engine.generate_gpt4o_insights(sales_data, inventory_data)
        
        dashboard['ai_enhanced_insights'] = gpt_insights
        
        print(f"SUCCESS: Generated AI-enhanced enterprise analytics")
        
        return jsonify({
            "status": "success",
            "analytics": dashboard,
            "timestamp": datetime.now().isoformat(),
            "data_sources": list(raw_data.keys()),
            "processing_type": "AI-Enhanced (GPT-4o)"
        })
        
    except Exception as e:
        print(f"ERROR: Error generating AI-enhanced enterprise analytics: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/ai/so-question', methods=['POST'])
def ai_so_question():
    """AI-powered Sales Order question answering"""
    try:
        if not openai_available or not client:
            return jsonify({
                'success': False,
                'error': 'OpenAI service not available'
            }), 503
        
        data = request.get_json()
        question = data.get('question', '')
        so_data = data.get('soData', [])
        conversation_history = data.get('conversationHistory', [])
        
        if not question:
            return jsonify({
                'success': False,
                'error': 'Question is required'
            }), 400
        
        # Prepare context from SO data
        context = "Sales Order Data:\n"
        for i, so in enumerate(so_data):
            context += f"\n--- Order {i+1} ---\n"
            context += f"Order Number: {so.get('orderNumber', 'N/A')}\n"
            context += f"Customer: {so.get('customerName', 'N/A')}\n"
            context += f"Order Date: {so.get('orderDate', 'N/A')}\n"
            context += f"Due Date: {so.get('dueDate', 'N/A')}\n"
            context += f"Total Amount: {so.get('totalAmount', 'N/A')}\n"
            context += f"Status: {so.get('status', 'N/A')}\n"
            if so.get('items'):
                context += "Items:\n"
                for item in so['items']:
                    context += f"  - {item.get('itemNumber', 'N/A')}: {item.get('description', 'N/A')} (Qty: {item.get('quantity', 'N/A')})\n"
            context += f"Raw Text: {so.get('rawText', '')[:500]}...\n"
        
        # Prepare conversation history
        history_context = ""
        if conversation_history:
            history_context = "\n\nPrevious conversation:\n"
            for msg in conversation_history[-3:]:  # Last 3 messages
                history_context += f"{msg['type']}: {msg['content']}\n"
        
        # Create the prompt
        prompt = f"""You are an AI assistant specialized in analyzing Sales Order data. 
        You have access to detailed Sales Order information including order numbers, customers, dates, amounts, items, and full text content.
        
        {context}
        
        {history_context}
        
        Question: {question}
        
        Please provide a helpful, accurate answer based on the Sales Order data provided. 
        If you need to reference specific orders, mention their order numbers.
        If the data doesn't contain enough information to answer the question, say so clearly.
        Be specific and cite relevant details from the data when possible.
        
        Answer:"""
        
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful AI assistant that analyzes Sales Order data to answer business questions. Provide accurate, specific answers based on the data provided."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=1000,
            temperature=0.3
        )
        
        answer = response.choices[0].message.content
        
        # Extract sources (order numbers mentioned in the answer)
        sources = []
        for so in so_data:
            if so.get('orderNumber') and so['orderNumber'] in answer:
                sources.append(f"SO-{so['orderNumber']}")
        
        return jsonify({
            'success': True,
            'answer': answer,
            'sources': sources
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error processing AI question: {str(e)}'
        }), 500

@app.route('/api/smart-so-search', methods=['POST'])
def smart_so_search_endpoint():
    """Smart Sales Order search endpoint"""
    try:
        if not openai_available or not client:
            return jsonify({
                'success': False,
                'error': 'OpenAI not available'
            }), 503
        
        data = request.get_json()
        query = data.get('query', '')
        
        if not query:
            return jsonify({
                'success': False,
                'error': 'Query is required'
            }), 400
        
        from smart_so_search import get_smart_search
        smart_search = get_smart_search(client)
        result = smart_search.smart_so_search(query)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Smart search failed: {str(e)}'
        }), 500

@app.route('/api/vision-analyze-so', methods=['POST'])
def vision_analyze_so():
    """GPT-4 Vision analysis of Sales Order files"""
    try:
        if not openai_available or not client:
            return jsonify({
                'success': False,
                'error': 'OpenAI not available'
            }), 503
        
        data = request.get_json()
        so_number = data.get('so_number', '')
        
        if not so_number:
            return jsonify({
                'success': False,
                'error': 'SO number is required'
            }), 400
        
        from smart_so_search import get_smart_search
        smart_search = get_smart_search(client)
        
        # Find the SO file
        found_files = smart_search.find_so_files_by_number(so_number)
        
        if not found_files:
            return jsonify({
                'success': False,
                'error': f'No files found for SO {so_number}'
            }), 404
        
        # Analyze the first found file with Vision
        file_path = found_files[0][0]
        vision_result = smart_search.analyze_so_with_vision(file_path)
        
        return jsonify(vision_result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Vision analysis failed: {str(e)}'
        }), 500

if __name__ == '__main__':
    print("Starting Flask backend...")
    print(f"G: Drive path: {GDRIVE_BASE}")
    
    # Test G: Drive access on startup - OPTIMIZED for speed
    print("RETRY: Testing G: Drive access on startup...")
    try:
        latest_folder, error = get_latest_folder()
        if latest_folder:
            print(f"SUCCESS: G: Drive accessible. Latest folder: {latest_folder}")
        else:
            print(f"⚠️ G: Drive issue: {error}")
    except Exception as e:
        print(f"ERROR: G: Drive test failed: {e}")
    
    # Run Flask app - bind to 0.0.0.0 to accept connections from localhost and 127.0.0.1
    print("🌐 Starting Flask server on http://0.0.0.0:5002 (accessible via localhost:5002)")
    app.run(host='0.0.0.0', port=5002, debug=False)
