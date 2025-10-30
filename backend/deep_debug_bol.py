#!/usr/bin/env python3
"""
DEEP DEBUG BOL - Investigate exactly what's happening in BOL generation
"""

from docx import Document
import os
import tempfile

def deep_debug_bol_generation():
    """Deep debug of BOL generation process"""
    
    print("🔍 DEEP DEBUG: BOL GENERATION INVESTIGATION")
    print("=" * 60)
    
    # Step 1: Check if template exists and can be loaded
    template_path = r"G:\Shared drives\IT_Automation\Automating Roles\Logistics\Bill Of Landing\Canoil Skeletal BOL Template.docx"
    
    print(f"📄 STEP 1: TEMPLATE VERIFICATION")
    print(f"   Template path: {template_path}")
    print(f"   Template exists: {os.path.exists(template_path)}")
    
    if not os.path.exists(template_path):
        print(f"❌ CRITICAL: Template not found!")
        return
    
    try:
        doc = Document(template_path)
        print(f"✅ Template loaded successfully")
        print(f"   Tables found: {len(doc.tables)}")
    except Exception as e:
        print(f"❌ CRITICAL: Template loading failed: {e}")
        return
    
    # Step 2: Check Table 4 structure (the items table)
    print(f"\n📋 STEP 2: TABLE 4 STRUCTURE ANALYSIS")
    if len(doc.tables) > 4:
        items_table = doc.tables[4]
        print(f"   Table 4 dimensions: {len(items_table.rows)} rows × {len(items_table.columns)} columns")
        
        # Check specific cells that should be replaced
        critical_cells = [
            (1, 1, "Product description placeholder"),
            (1, 2, "Weight placeholder"),
            (2, 1, "Batch number placeholder"),
            (5, 1, "PO number placeholder"),
            (9, 1, "Skid sizing placeholder")
        ]
        
        print(f"   Critical cells content:")
        for row_idx, col_idx, description in critical_cells:
            if row_idx < len(items_table.rows) and col_idx < len(items_table.columns):
                cell_text = items_table.rows[row_idx].cells[col_idx].text.strip()
                print(f"      Row {row_idx}, Col {col_idx} ({description}): {repr(cell_text[:50])}...")
            else:
                print(f"      Row {row_idx}, Col {col_idx} ({description}): CELL NOT FOUND")
    else:
        print(f"❌ CRITICAL: Table 4 not found! Only {len(doc.tables)} tables exist")
        return
    
    # Step 3: Test cell replacement function
    print(f"\n🔧 STEP 3: CELL REPLACEMENT FUNCTION TEST")
    
    # Test the replace_cell_content_completely function
    test_cell = items_table.rows[1].cells[1]  # Product description cell
    original_text = test_cell.text.strip()
    print(f"   Original cell text: {repr(original_text[:50])}...")
    
    # Test replacement
    try:
        # Clear the cell
        test_cell.clear()
        print(f"   ✅ Cell cleared successfully")
        
        # Add new content
        test_cell.add_paragraph("TEST REPLACEMENT TEXT")
        new_text = test_cell.text.strip()
        print(f"   New cell text: {repr(new_text)}")
        
        if "TEST REPLACEMENT TEXT" in new_text:
            print(f"   ✅ Cell replacement WORKS")
        else:
            print(f"   ❌ Cell replacement FAILED - text not found")
            
    except Exception as e:
        print(f"   ❌ Cell replacement ERROR: {e}")
    
    # Step 4: Test with simulated form_data
    print(f"\n📊 STEP 4: SIMULATED BOL GENERATION TEST")
    
    # Simulate the form_data that should be passed to BOL generation
    form_data = {
        'total_weight': '2645.5 lb',
        'batch_number': 'CC-09-06-24',
        'skid_sizing': '2 pallet (48x40 standard size)',
        'po_number': '325033',
        'items': [
            {'description': 'MOV Long Life0 30 Pack 400g Tube Case', 'quantity': '4', 'unit': 'case'},
            {'description': 'MOV LONG LIFE1 30x400g PACK TUBE', 'quantity': '4', 'unit': 'case'},
            {'description': 'MOV Long Life 1 - PAILS', 'quantity': '4', 'unit': 'pail'}
        ]
    }
    
    print(f"   Form data prepared:")
    for key, value in form_data.items():
        if key != 'items':
            print(f"      {key}: {value}")
        else:
            print(f"      items: {len(value)} products")
    
    # Step 5: Apply replacements manually to test
    print(f"\n🎯 STEP 5: MANUAL REPLACEMENT TEST")
    
    try:
        # Test Row 1, Col 2 - Weight
        weight_cell = items_table.rows[1].cells[2]
        print(f"   Testing weight replacement...")
        print(f"      Before: {repr(weight_cell.text.strip()[:30])}...")
        weight_cell.clear()
        weight_cell.add_paragraph(form_data['total_weight'])
        print(f"      After: {repr(weight_cell.text.strip())}")
        
        # Test Row 2, Col 1 - Batch
        batch_cell = items_table.rows[2].cells[1]
        print(f"   Testing batch replacement...")
        print(f"      Before: {repr(batch_cell.text.strip()[:30])}...")
        batch_cell.clear()
        batch_cell.add_paragraph(f"Batch: {form_data['batch_number']}")
        print(f"      After: {repr(batch_cell.text.strip())}")
        
        # Test Row 5, Col 1 - PO Number
        po_cell = items_table.rows[5].cells[1]
        print(f"   Testing PO replacement...")
        print(f"      Before: {repr(po_cell.text.strip()[:30])}...")
        po_cell.clear()
        po_cell.add_paragraph(f"PO#: {form_data['po_number']}")
        print(f"      After: {repr(po_cell.text.strip())}")
        
        # Test Row 9, Col 1 - Skid sizing
        skid_cell = items_table.rows[9].cells[1]
        print(f"   Testing skid replacement...")
        print(f"      Before: {repr(skid_cell.text.strip()[:30])}...")
        skid_cell.clear()
        skid_cell.add_paragraph(form_data['skid_sizing'])
        print(f"      After: {repr(skid_cell.text.strip())}")
        
        print(f"   ✅ All manual replacements completed")
        
    except Exception as e:
        print(f"   ❌ Manual replacement failed: {e}")
        import traceback
        traceback.print_exc()
    
    # Step 6: Save test document
    print(f"\n💾 STEP 6: SAVING TEST DOCUMENT")
    try:
        temp_dir = tempfile.gettempdir()
        test_output = os.path.join(temp_dir, "DEBUG_BOL_TEST.docx")
        doc.save(test_output)
        print(f"   ✅ Test document saved: {test_output}")
        print(f"   📁 Open this file to see if replacements worked")
        
    except Exception as e:
        print(f"   ❌ Save failed: {e}")
    
    print(f"\n🎯 INVESTIGATION COMPLETE")
    print(f"   Check the saved document to see if manual replacements worked")
    print(f"   If they did, the issue is in the BOL generation function logic")
    print(f"   If they didn't, the issue is in the cell replacement method")

if __name__ == "__main__":
    deep_debug_bol_generation()

